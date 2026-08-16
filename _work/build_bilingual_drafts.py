from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from PIL import Image, ImageDraw, ImageFont
import shutil, math

ROOT = Path(r"C:\Users\wilki\OneDrive\Desktop\Trabajo de Curso")
IMG_SRC = Path(r"C:\Users\wilki\.codex\generated_images\01a00634-44ce-7b31-928a-aa3428ef7495\exec-d6c3f8a0-632a-4ff9-aa18-cd35d48ca8f3.png")
IMG = ROOT / "_work" / "anexo_casa_assados_sofia.png"
shutil.copy2(IMG_SRC, IMG)
CHART_DIR = ROOT / "_work" / "charts"
CHART_DIR.mkdir(exist_ok=True)

mix = [
    ("O Clássico da Sofia", 70, 69.90, 26.50),
    ("Costela Suprema", 35, 119.90, 48.00),
    ("Dueto Sofia", 35, 94.90, 36.00),
    ("Kit Churrasco Família", 20, 169.90, 68.00),
]
revenue = sum(q*p for _,q,p,c in mix)
cmv = sum(q*c for _,q,p,c in mix)
tax = revenue*.04
fees = revenue*.02
fixed = 6690.00
profit = revenue-cmv-tax-fees-fixed
cm_ratio = (revenue-cmv-tax-fees)/revenue
breakeven = fixed/cm_ratio
payback = 20000/profit

def make_charts():
    colors=['#8B2F2F','#D79A3B','#4D6A4D','#567A9D']
    W,H=1500,820; margin=(150,100,60,150)
    try: font=ImageFont.truetype('arial.ttf',26); small=ImageFont.truetype('arial.ttf',21); titlef=ImageFont.truetype('arialbd.ttf',34)
    except: font=small=titlef=ImageFont.load_default()
    def base(title):
        im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im); d.text((W//2,35),title,font=titlef,fill='#222',anchor='ma')
        x0,y0,x1,y1=margin[0],margin[1],W-margin[2],H-margin[3]; d.line((x0,y1,x1,y1),fill='#444',width=3); d.line((x0,y0,x0,y1),fill='#444',width=3)
        return im,d,(x0,y0,x1,y1)
    def bars(path,title,labs,vals,cols):
        im,d,(x0,y0,x1,y1)=base(title); vmax=max(vals)*1.12; bw=(x1-x0)/(len(vals)*1.7); gap=((x1-x0)-bw*len(vals))/(len(vals)+1)
        for i,(lab,val,col) in enumerate(zip(labs,vals,cols)):
            x=x0+gap*(i+1)+bw*i; h=(y1-y0)*val/vmax; d.rectangle((x,y1-h,x+bw,y1),fill=col)
            d.text((x+bw/2,y1-h-12),f'{val:,.0f}'.replace(',','.'),font=small,fill='#222',anchor='ms')
            parts=lab.split(' '); d.multiline_text((x+bw/2,y1+18),'\n'.join(parts),font=small,fill='#222',anchor='ma',align='center',spacing=2)
        im.save(path)
    bars(CHART_DIR/'mix.png','Mix mensal de vendas - cenário-base',[x[0] for x in mix],[x[1] for x in mix],colors)
    bars(CHART_DIR/'mix_es.png','Mezcla mensual de ventas - escenario base',[x[0] for x in mix],[x[1] for x in mix],colors)
    bars(CHART_DIR/'dre.png','Composição do resultado mensal',['Receita','CMV','Tributos e taxas','Custos fixos','Lucro'],[revenue,cmv,tax+fees,fixed,profit],['#567A9D','#D79A3B','#B9A37C','#8B2F2F','#4D6A4D'])
    bars(CHART_DIR/'dre_es.png','Composición del resultado mensual',['Ingresos','CMV','Tributos y tasas','Costos fijos','Utilidad'],[revenue,cmv,tax+fees,fixed,profit],['#567A9D','#D79A3B','#B9A37C','#8B2F2F','#4D6A4D'])
    def line_chart(path,title,xvals,series,xlabel):
        im,d,(x0,y0,x1,y1)=base(title); allv=[v for _,vals,_ in series for v in vals]; vmin=min(0,min(allv)); vmax=max(allv)*1.08
        for k in range(5):
            yy=y1-(y1-y0)*k/4; d.line((x0,yy,x1,yy),fill='#DDDDDD',width=1); val=vmin+(vmax-vmin)*k/4; d.text((x0-15,yy),f'{val:,.0f}'.replace(',','.'),font=small,fill='#444',anchor='rm')
        for label,vals,col in series:
            pts=[]
            for i,(x,v) in enumerate(zip(xvals,vals)):
                px=x0+(x1-x0)*i/(len(xvals)-1); py=y1-(y1-y0)*(v-vmin)/(vmax-vmin); pts.append((px,py))
            d.line(pts,fill=col,width=5)
            for px,py in pts: d.ellipse((px-5,py-5,px+5,py+5),fill=col)
        d.text(((x0+x1)//2,H-35),xlabel,font=font,fill='#333',anchor='ma')
        for idx in range(0,len(xvals),max(1,len(xvals)//10)): d.text((x0+(x1-x0)*idx/(len(xvals)-1),y1+15),str(xvals[idx]),font=small,fill='#333',anchor='ma')
        lx=x1-360; ly=y0+15
        for label,vals,col in series: d.line((lx,ly+12,lx+50,ly+12),fill=col,width=5); d.text((lx+65,ly),label,font=small,fill='#222'); ly+=35
        im.save(path)
    qs=list(range(0,261,10)); avg=revenue/160; var_ratio=1-cm_ratio; revs=[q*avg for q in qs]; costs=[fixed+r*var_ratio for r in revs]
    line_chart(CHART_DIR/'breakeven.png','Ponto de equilíbrio',qs,[('Receita',revs,'#567A9D'),('Custo total',costs,'#8B2F2F')],'Combos equivalentes / mês')
    line_chart(CHART_DIR/'breakeven_es.png','Punto de equilibrio',qs,[('Ingresos',revs,'#567A9D'),('Costo total',costs,'#8B2F2F')],'Combos equivalentes / mes')
    months=list(range(1,13)); quantities=list(range(160,249,8)); results=[q*avg*cm_ratio-fixed for q in quantities]
    line_chart(CHART_DIR/'result12.png','Evolução projetada do resultado',months,[('Resultado operacional',results,'#4D6A4D')],'Mês')
    line_chart(CHART_DIR/'result12_es.png','Evolución proyectada del resultado',months,[('Resultado operativo',results,'#4D6A4D')],'Mes')

make_charts()

SOURCES = [
    ("IBGE", "Curitiba: panorama municipal, população estimada de 1.829.225 habitantes (2024) e PIB per capita de R$ 67.691,30 (2023).", "https://www.ibge.gov.br/cidades-e-estados/pr/curitiba.html"),
    ("Prefeitura de Curitiba", "Administração Regional Bairro Novo, que atende Umbará, Ganchinho e Sítio Cercado.", "https://www.curitiba.pr.gov.br/noticias/veja-como-entrar-em-contato-com-as-regionais-e-qual-delas-cuida-do-seu-bairro/40912"),
    ("SEBRAE", "Orientações sobre alimentação fora do lar, implantação de delivery e adequação ao perfil regional.", "https://meuatendimento.sebrae.com.br/sites/PortalSebrae/artigos/como-implantar-delivery-na-era-digital,039214266f1e2710VgnVCM1000004c00210aRCRD"),
    ("SEBRAE", "CRM e automação de marketing para restaurantes e deliveries como instrumentos de recompra e fidelização.", "https://meuatendimento.sebrae.com.br/sites/PortalSebrae/ufs/sp/programas/sebrae-conecta-repediu,801a0f9160f5a910VgnVCM1000001b00320aRCRD"),
    ("BRASIL", "Portal oficial do empreendedor e regras de enquadramento empresarial; consulta realizada em 15 ago. 2026.", "https://www.gov.br/empresas-e-negocios/pt-br/empreendedor"),
    ("CURITIBA", "Informações municipais sobre alvará, vigilância sanitária e licenciamento.", "https://www.curitiba.pr.gov.br/servicos/"),
    ("CEASA PARANÁ", "Referência institucional para abastecimento hortigranjeiro em Curitiba.", "https://www.ceasa.pr.gov.br/"),
]

def money(x):
    return f"R$ {x:,.2f}".replace(",","X").replace(".",",").replace("X",".")

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_cell_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa')

def table(doc, headers, rows, widths=None, font=8.5):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); set_cell_shading(c,'D9EAD3'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True; r.font.name='Arial'; r.font.size=Pt(font)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs: r.font.name='Arial'; r.font.size=Pt(font)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): set_cell_width(row.cells[i],w)
        tblPr=t._tbl.tblPr
        tblW=tblPr.find(qn('w:tblW'))
        if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
        tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
        grid=t._tbl.tblGrid
        for child in list(grid): grid.remove(child)
        for w in widths:
            gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); grid.append(gc)
    for row_idx,row in enumerate(t.rows):
        for cell in row.cells:
            tcPr=cell._tc.get_or_add_tcPr(); mar=tcPr.find(qn('w:tcMar'))
            if mar is None: mar=OxmlElement('w:tcMar'); tcPr.append(mar)
            for side,val in [('top',90),('bottom',90),('start',110),('end',110)]:
                e=mar.find(qn('w:'+side))
                if e is None: e=OxmlElement('w:'+side); mar.append(e)
                e.set(qn('w:w'),str(val)); e.set(qn('w:type'),'dxa')
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.line_spacing=1.0
                if row_idx==0: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # repeating header
    trPr=t.rows[0]._tr.get_or_add_trPr(); rep=OxmlElement('w:tblHeader'); rep.set(qn('w:val'),'true'); trPr.append(rep)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def add_figure(doc,path,caption,pt,source=None):
    source = source or ('Elaboração própria a partir das premissas do plano (2026).' if pt else 'Elaboración propia a partir de los supuestos del plan (2026).')
    doc.add_picture(str(path),width=Cm(15.5)); doc.inline_shapes[-1]._inline.docPr.set('descr', caption); doc.inline_shapes[-1]._inline.docPr.set('title', caption); p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0)
    cap=doc.add_paragraph(caption); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.first_line_indent=Cm(0); cap.runs[0].font.size=Pt(10); cap.runs[0].bold=True
    src=doc.add_paragraph(('Fonte: ' if pt else 'Fuente: ')+source); src.alignment=WD_ALIGN_PARAGRAPH.CENTER; src.paragraph_format.first_line_indent=Cm(0); src.runs[0].font.size=Pt(9); src.runs[0].italic=True

def page_field(paragraph):
    run=paragraph.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.addnext(fld)

def toc(doc, title, pt):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.bold=True; r.font.name='Arial'; r.font.size=Pt(12)
    entries=([('INTRODUÇÃO',5),('1 SUMÁRIO EXECUTIVO',7),('2 ANÁLISE DE MERCADO',9),('3 PLANO DE MARKETING',11),('4 PLANO OPERACIONAL',13),('5 PLANO FINANCEIRO',16),('6 ANÁLISE DE VIABILIDADE',20),('7 ANEXOS E INSTRUMENTOS DE IMPLANTAÇÃO',23),('CONCLUSÃO',26),('REFERÊNCIAS',27)] if pt else [('INTRODUCCIÓN',5),('1 RESUMEN EJECUTIVO',7),('2 ANÁLISIS DE MERCADO',9),('3 PLAN DE MARKETING',11),('4 PLAN OPERATIVO',13),('5 PLAN FINANCIERO',16),('6 ANÁLISIS DE VIABILIDAD',20),('7 ANEXOS E INSTRUMENTOS DE IMPLANTACIÓN',23),('CONCLUSIÓN',26),('REFERENCIAS',27)])
    for label,page in entries:
        p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(4)
        p.add_run(label); p.add_run('\t'+str(page)); p.paragraph_format.tab_stops.add_tab_stop(Cm(15),2,1)
    note=doc.add_paragraph("Paginação provisória do rascunho; atualizar após a revisão institucional final." if pt else "Paginación provisional del borrador; actualizar tras la revisión institucional final.")
    note.alignment=WD_ALIGN_PARAGRAPH.CENTER; note.runs[0].italic=True; note.runs[0].font.size=Pt(9)

def configure(doc):
    sec=doc.sections[0]; sec.page_height=Cm(29.7); sec.page_width=Cm(21); sec.top_margin=Cm(3); sec.left_margin=Cm(3); sec.right_margin=Cm(2); sec.bottom_margin=Cm(2)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(12); normal.paragraph_format.line_spacing=1.5; normal.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.first_line_indent=Cm(1.25)
    for name,size in [('Title',16),('Heading 1',14),('Heading 2',12),('Heading 3',12)]:
        s=styles[name]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor(0,0,0); s.paragraph_format.first_line_indent=Cm(0); s.paragraph_format.space_before=Pt(12); s.paragraph_format.space_after=Pt(6); s.paragraph_format.keep_with_next=True
    styles['Heading 1'].font.all_caps=True
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run('Casa de Assados Sofia | '); page_field(footer)
    footer.style=styles['Normal']; footer.paragraph_format.first_line_indent=Cm(0)

def add_cover(doc, lang):
    pt=lang=='pt'
    for text,space,bold,size in [
        ("FACULDADE DE CIÊNCIAS E EMPREENDEDORISMO - FACEMP",42,True,12),
        ("BACHARELADO EM ADMINISTRAÇÃO" if pt else "BACHILLERATO EN ADMINISTRACIÓN",70,True,12),
        ("WILKIN BARBAN ROSABAL",90,True,12),
        ("CASA DE ASSADOS SOFIA",8,True,16),
        (("PLANO DE NEGÓCIO PARA IMPLANTAÇÃO DE UMA MICROEMPRESA DE ASSADOS COM CRM EM UMBARÁ, CURITIBA - PR" if pt else "PLAN DE NEGOCIOS PARA LA IMPLANTACIÓN DE UNA MICROEMPRESA DE ASADOS CON CRM EN UMBARÁ, CURITIBA - PR"),120,True,13),
        ("CURITIBA - PR\n2026",0,True,12)]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(space)
        r=p.add_run(text); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(size)
    doc.add_page_break()
    p=doc.add_paragraph('WILKIN BARBAN ROSABAL'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True; p.paragraph_format.space_after=Pt(90); p.paragraph_format.first_line_indent=Cm(0)
    title=("CASA DE ASSADOS SOFIA: PLANO DE NEGÓCIO PARA UMA OPERAÇÃO DE FIM DE SEMANA APOIADA POR CRM" if pt else "CASA DE ASSADOS SOFIA: PLAN DE NEGOCIOS PARA UNA OPERACIÓN DE FIN DE SEMANA APOYADA POR CRM")
    p=doc.add_paragraph(title); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold=True; p.paragraph_format.space_after=Pt(70); p.paragraph_format.first_line_indent=Cm(0)
    note=("Trabalho de Conclusão de Curso apresentado ao Bacharelado em Administração da Faculdade de Ciências e Empreendedorismo - FACEMP, como requisito parcial para obtenção do grau de Bacharel em Administração.\n\nOrientador(a): ______________________________" if pt else "Trabajo de Conclusión de Curso presentado al Bachillerato en Administración de la Facultad de Ciencias y Emprendimiento - FACEMP, como requisito parcial para la obtención del título de Bachiller en Administración.\n\nTutor(a): ______________________________")
    p=doc.add_paragraph(note); p.paragraph_format.left_indent=Cm(8); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.line_spacing=1; p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p=doc.add_paragraph('CURITIBA - PR\n2026'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(110); p.paragraph_format.first_line_indent=Cm(0); p.runs[0].bold=True
    doc.add_page_break()

def add_p(doc,text,boldlead=None):
    p=doc.add_paragraph();
    if boldlead and text.startswith(boldlead):
        p.add_run(boldlead).bold=True; p.add_run(text[len(boldlead):])
    else: p.add_run(text)
    return p

def add_bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.first_line_indent=Cm(0); p.add_run(item)

def chapter(doc,n,title):
    doc.add_page_break(); doc.add_heading(f"{n} {title}",level=1)

def build(lang, output):
    pt=lang=='pt'; doc=Document(); configure(doc); add_cover(doc,lang)
    # Resumos
    doc.add_heading('RESUMO' if pt else 'RESUMEN',level=1)
    summary_pt=("Este trabalho apresenta o plano de negócio da Casa de Assados Sofia, microempresa projetada para operar aos sábados, domingos e feriados em Umbará, Curitiba. O modelo combina retirada no local e entrega em raio controlado, cardápio enxuto de quatro combos familiares e um sistema próprio de gestão do relacionamento com clientes (CRM). A pesquisa é aplicada, exploratória e documental, apoiada em manuais de plano de negócio, normas acadêmicas, dados públicos e simulação econômico-financeira. No cenário-base de 160 combos mensais, a receita estimada é de R$ 15.809,00, o lucro operacional de R$ 2.015,46, o ponto de equilíbrio de R$ 12.148,95 e o prazo de retorno de 9,92 meses. Conclui-se que o empreendimento é viável sob controle rigoroso de custos, validação prévia da demanda, regularização sanitária e uso responsável do CRM, com consentimento e proteção de dados.")
    summary_es=("Este trabajo presenta el plan de negocios de Casa de Assados Sofia, microempresa proyectada para operar los sábados, domingos y días feriados en Umbará, Curitiba. El modelo combina retiro en el local y entrega dentro de un radio controlado, un menú reducido de cuatro combos familiares y un sistema propio de gestión de relaciones con clientes (CRM). La investigación es aplicada, exploratoria y documental, apoyada en manuales de plan de negocios, normas académicas, datos públicos y simulación económico-financiera. En el escenario base de 160 combos mensuales, los ingresos estimados son R$ 15.809,00, la utilidad operativa R$ 2.015,46, el punto de equilibrio R$ 12.148,95 y el plazo de recuperación 9,92 meses. Se concluye que el emprendimiento es viable bajo control riguroso de costos, validación previa de la demanda, regularización sanitaria y uso responsable del CRM, con consentimiento y protección de datos.")
    add_p(doc,summary_pt if pt else summary_es)
    add_p(doc,("Palavras-chave: plano de negócio; alimentação; delivery; CRM; Curitiba." if pt else "Palabras clave: plan de negocios; alimentación; entrega a domicilio; CRM; Curitiba."))
    doc.add_page_break(); toc(doc,'SUMÁRIO' if pt else 'ÍNDICE',pt);
    # Intro and methodology
    doc.add_page_break(); doc.add_heading('INTRODUÇÃO' if pt else 'INTRODUCCIÓN',level=1)
    add_p(doc,("O projeto nasce da oportunidade de atender famílias da região sul de Curitiba que buscam almoço de fim de semana pronto, com sabor caseiro, previsibilidade de horário e pedido simples pelo celular. A proposta não tenta competir por variedade infinita: concentra capacidade em poucos produtos de giro elevado e usa dados de relacionamento para planejar produção, reduzir desperdícios e estimular recompra." if pt else "El proyecto nace de la oportunidad de atender a familias de la región sur de Curitiba que buscan un almuerzo de fin de semana listo, con sabor casero, horario previsible y pedido sencillo por el celular. La propuesta no intenta competir mediante una variedad infinita: concentra la capacidad en pocos productos de alta rotación y utiliza datos de relación para planificar la producción, reducir desperdicios y estimular la recompra."))
    add_p(doc,("Problema de pesquisa: em que condições mercadológicas, operacionais, financeiras e tecnológicas a Casa de Assados Sofia pode iniciar suas atividades de modo sustentável em Umbará? O objetivo geral é elaborar um plano de negócio completo. Os objetivos específicos são analisar o mercado local; definir marketing, operação e CRM; projetar investimento, receitas e custos; calcular indicadores de viabilidade; e organizar riscos e ações de implantação." if pt else "Problema de investigación: ¿bajo qué condiciones de mercado, operación, finanzas y tecnología puede Casa de Assados Sofia iniciar sus actividades de manera sostenible en Umbará? El objetivo general es elaborar un plan de negocios completo. Los objetivos específicos son analizar el mercado local; definir marketing, operación y CRM; proyectar inversión, ingresos y costos; calcular indicadores de viabilidad; y organizar riesgos y acciones de implantación."))
    doc.add_heading('METODOLOGIA E LIMITAÇÕES' if pt else 'METODOLOGÍA Y LIMITACIONES',level=2)
    add_p(doc,("Adota-se abordagem aplicada e exploratória, com pesquisa documental em materiais acadêmicos, manuais de plano de negócio e fontes institucionais. Os valores de preços, custos e demanda são premissas de planejamento em reais de 2026 e devem ser confirmados por três cotações e por um piloto comercial de quatro fins de semana. Não houve pesquisa de campo auditada. Questionários, personas e imagens dos anexos são instrumentos simulados/ilustrativos; não constituem prova de entrevistas, clientes, instalações ou vendas reais." if pt else "Se adopta un enfoque aplicado y exploratorio, con investigación documental en materiales académicos, manuales de plan de negocios y fuentes institucionales. Los precios, costos y demanda son supuestos de planificación en reales de 2026 y deben confirmarse mediante tres cotizaciones y un piloto comercial de cuatro fines de semana. No se realizó una investigación de campo auditada. Los cuestionarios, perfiles e imágenes de los anexos son instrumentos simulados/ilustrativos; no constituyen prueba de entrevistas, clientes, instalaciones o ventas reales."))
    doc.add_heading('FUNDAMENTAÇÃO TEÓRICA' if pt else 'FUNDAMENTACIÓN TEÓRICA',level=2)
    add_p(doc,("O plano de negócio transforma hipóteses em decisões verificáveis sobre clientes, proposta de valor, recursos, operação e retorno. Na perspectiva de Schumpeter, o empreendedor combina recursos de forma inovadora; em Dornelas e Dolabela, planejamento e aprendizagem reduzem incerteza sem eliminar o risco. O CRM, entendido como estratégia e processo apoiado por tecnologia, integra aquisição, atendimento, histórico de compra, segmentação e fidelização. Para um pequeno negócio, sua utilidade não está em acumular dados, mas em gerar decisões: quanto produzir, a quem comunicar, quando recuperar um cliente e quais ofertas aumentam valor sem destruir margem." if pt else "El plan de negocios transforma hipótesis en decisiones verificables sobre clientes, propuesta de valor, recursos, operación y retorno. Desde la perspectiva de Schumpeter, el emprendedor combina recursos de forma innovadora; en Dornelas y Dolabela, la planificación y el aprendizaje reducen la incertidumbre sin eliminar el riesgo. El CRM, entendido como estrategia y proceso apoyado por tecnología, integra adquisición, atención, historial de compra, segmentación y fidelización. Para una pequeña empresa, su utilidad no está en acumular datos, sino en generar decisiones: cuánto producir, a quién comunicar, cuándo recuperar un cliente y qué ofertas aumentan valor sin destruir margen."))

    chapter(doc,1,'SUMÁRIO EXECUTIVO' if pt else 'RESUMEN EJECUTIVO')
    add_p(doc,("A Casa de Assados Sofia será instalada na Rua Deputado Pinheiro Júnior, 1380, Umbará, Curitiba - PR, com funcionamento concentrado em fins de semana e feriados. O negócio oferecerá frango assado, costela bovina, costelinha suína, linguiça e acompanhamentos em quatro combos. Os pedidos serão recebidos por canal digital e confirmados pelo CRM, com janelas de retirada e entrega que protegem a qualidade e a capacidade da cozinha." if pt else "Casa de Assados Sofia se instalará en la Rua Deputado Pinheiro Júnior, 1380, Umbará, Curitiba - PR, con funcionamiento concentrado en fines de semana y feriados. El negocio ofrecerá pollo asado, costilla bovina, costilla de cerdo, salchicha y acompañamientos en cuatro combos. Los pedidos se recibirán por un canal digital y serán confirmados por el CRM, con franjas de retiro y entrega que protegen la calidad y la capacidad de la cocina."))
    table(doc,['Item' if pt else 'Elemento','Definição' if pt else 'Definición'],[
        [('Razão/forma' if pt else 'Razón/forma'),('SLU enquadrada como microempresa no Simples Nacional; validação com contador.' if pt else 'SLU encuadrada como microempresa en Simples Nacional; validación con contador.')],
        [('Proprietário' if pt else 'Propietario'),'Wilkin Barban Rosabal'],
        [('Missão' if pt else 'Misión'),('Entregar refeições de fim de semana saborosas, seguras e pontuais, criando vínculo com as famílias do bairro.' if pt else 'Entregar comidas de fin de semana sabrosas, seguras y puntuales, creando vínculo con las familias del barrio.')],
        [('Visão' if pt else 'Visión'),('Ser referência de assados por encomenda na região do Bairro Novo até 2029.' if pt else 'Ser referencia de asados por encargo en la región de Bairro Novo hasta 2029.')],
        [('Diferencial' if pt else 'Diferencial'),('Operação enxuta + produção orientada por reservas + CRM de recompra.' if pt else 'Operación ajustada + producción orientada por reservas + CRM de recompra.')],
    ],[2200,7160])
    add_p(doc,("A opção por SLU/ME corrige um risco do rascunho inicial: a receita anual projetada e a necessidade de mais de uma pessoa na operação tornam o MEI inadequado como premissa central. O enquadramento definitivo depende da atividade econômica, do faturamento efetivo e da orientação contábil." if pt else "La opción por SLU/ME corrige un riesgo del borrador inicial: los ingresos anuales proyectados y la necesidad de más de una persona en la operación vuelven inadecuado al MEI como supuesto central. El encuadre definitivo depende de la actividad económica, la facturación efectiva y la orientación contable."))
    doc.add_heading('1.1 Perfil do empreendedor e objetivos estratégicos' if pt else '1.1 Perfil del emprendedor y objetivos estratégicos',level=2)
    add_p(doc,("Wilkin Barban Rosabal exercerá a administração geral, com responsabilidade por compras, caixa, fornecedores, indicadores e governança do CRM. A concentração inicial é coerente com uma microempresa, mas exige rotinas documentadas e delegação gradual para evitar que toda decisão dependa de uma única pessoa. Nos primeiros doze meses, os objetivos são validar a demanda recorrente, atingir pontualidade mínima de 90%, limitar desperdício a 5% da produção e formar uma base consentida de clientes ativos." if pt else "Wilkin Barban Rosabal ejercerá la administración general, con responsabilidad sobre compras, caja, proveedores, indicadores y gobierno del CRM. La concentración inicial es coherente con una microempresa, pero exige rutinas documentadas y delegación gradual para evitar que toda decisión dependa de una sola persona. En los primeros doce meses, los objetivos son validar la demanda recurrente, alcanzar una puntualidad mínima del 90%, limitar el desperdicio al 5% de la producción y formar una base consentida de clientes activos."))
    doc.add_page_break(); doc.add_heading('1.2 Metas por horizonte' if pt else '1.2 Metas por horizonte',level=2)
    table(doc,[('Horizonte' if pt else 'Horizonte'),('Meta operacional' if pt else 'Meta operativa'),('Evidência no CRM' if pt else 'Evidencia en CRM')],[
        [('0-3 meses'),('Piloto e estabilização em 160 combos/mês' if pt else 'Piloto y estabilización en 160 combos/mes'),('Pedidos, capacidade, atrasos e perdas' if pt else 'Pedidos, capacidad, retrasos y pérdidas')],
        [('4-12 meses'),('Elevar recompra e chegar a 220-248 combos/mês' if pt else 'Elevar recompra y llegar a 220-248 combos/mes'),('Coortes, frequência, ticket e margem' if pt else 'Cohortes, frecuencia, ticket y margen')],
        [('13-24 meses'),('Avaliar compra de mais duas máquinas' if pt else 'Evaluar compra de dos máquinas adicionales'),('Fila perdida, ocupação e retorno incremental' if pt else 'Cola perdida, ocupación y retorno incremental')],
    ],[1700,3830,3830])

    chapter(doc,2,'ANÁLISE DE MERCADO' if pt else 'ANÁLISIS DE MERCADO')
    add_p(doc,("Curitiba possuía população estimada de 1.829.225 habitantes em 2024 e PIB per capita de R$ 67.691,30 em 2023, segundo o IBGE. Umbará integra a Administração Regional Bairro Novo, junto de Ganchinho e Sítio Cercado. Para a empresa, o mercado relevante não é toda a cidade, mas famílias em raio operacional aproximado de 5 km, com maior demanda entre 11h e 14h nos fins de semana." if pt else "Curitiba tenía una población estimada de 1.829.225 habitantes en 2024 y un PIB per cápita de R$ 67.691,30 en 2023, según el IBGE. Umbará integra la Administración Regional Bairro Novo, junto con Ganchinho y Sítio Cercado. Para la empresa, el mercado relevante no es toda la ciudad, sino las familias dentro de un radio operativo aproximado de 5 km, con mayor demanda entre las 11:00 y las 14:00 durante los fines de semana."))
    doc.add_heading('2.1 Delimitação e dimensionamento do mercado' if pt else '2.1 Delimitación y dimensionamiento del mercado',level=2)
    add_p(doc,("O dimensionamento utiliza três níveis. O mercado total é a população de Curitiba; o mercado atendível corresponde aos domicílios do entorno que aceitam retirada ou entrega curta; o mercado-alvo inicial é limitado pela capacidade de produção, não pelo tamanho da cidade. Com 160 combos por mês, o negócio precisa de 40 pedidos por fim de semana ou aproximadamente 20 por dia regular. Essa meta é operacionalmente observável e deve ser testada, enquanto qualquer estimativa de participação municipal seria artificialmente ampla." if pt else "El dimensionamiento utiliza tres niveles. El mercado total es la población de Curitiba; el mercado atendible corresponde a los hogares del entorno que aceptan retiro o entrega corta; el mercado objetivo inicial está limitado por la capacidad productiva, no por el tamaño de la ciudad. Con 160 combos al mes, el negocio necesita 40 pedidos por fin de semana o aproximadamente 20 por día regular. Esta meta es operativamente observable y debe probarse, mientras que cualquier estimación de participación municipal sería artificialmente amplia."))
    doc.add_heading('2.2 Público-alvo e necessidade' if pt else '2.2 Público objetivo y necesidad',level=2)
    add_bullets(doc,[
        ('Famílias de 3 a 6 pessoas que desejam reduzir tempo de preparo e limpeza.' if pt else 'Familias de 3 a 6 personas que desean reducir el tiempo de preparación y limpieza.'),
        ('Moradores que valorizam comida conhecida, porção previsível e compra por WhatsApp.' if pt else 'Residentes que valoran comida conocida, porción previsible y compra por WhatsApp.'),
        ('Clientes de retirada no bairro e delivery curto, evitando perda de temperatura.' if pt else 'Clientes de retiro en el barrio y entrega corta, evitando pérdida de temperatura.'),
    ])
    doc.add_heading('2.3 Concorrência e posicionamento' if pt else '2.3 Competencia y posicionamiento',level=2)
    table(doc,[('Tipo' if pt else 'Tipo'),('Força concorrente' if pt else 'Fortaleza competitiva'),('Resposta Sofia' if pt else 'Respuesta Sofia')],[
        [('Assadores locais' if pt else 'Asadores locales'),('Tradição e proximidade' if pt else 'Tradición y proximidad'),('Reserva, pontualidade e combos completos' if pt else 'Reserva, puntualidad y combos completos')],
        [('Supermercados' if pt else 'Supermercados'),('Preço e fluxo de clientes' if pt else 'Precio y flujo de clientes'),('Especialização, frescor e relacionamento' if pt else 'Especialización, frescura y relación')],
        [('Apps/marketplaces' if pt else 'Apps/marketplaces'),('Variedade e conveniência' if pt else 'Variedad y conveniencia'),('Canal direto, menor dependência de comissão' if pt else 'Canal directo, menor dependencia de comisión')],
    ],[1800,3300,4260])
    doc.add_heading('2.4 Fornecedores e critérios de homologação' if pt else '2.4 Proveedores y criterios de homologación',level=2)
    add_p(doc,("A proximidade da CEASA Curitiba favorece compras de hortifrutigranjeiros às sextas-feiras, mas preço não deve ser o único critério. Carnes e bebidas serão cotadas com distribuidores da região sul de Curitiba e de São José dos Pinhais. Cada fornecedor receberá nota para regularidade fiscal e sanitária, temperatura de entrega, prazo, padrão do corte, preço, substituição de lote e capacidade de atender picos. Pelo menos dois fornecedores devem permanecer homologados para cada insumo crítico." if pt else "La proximidad de CEASA Curitiba favorece las compras de frutas y verduras los viernes, pero el precio no debe ser el único criterio. Las carnes y bebidas se cotizarán con distribuidores de la región sur de Curitiba y de São José dos Pinhais. Cada proveedor recibirá una puntuación por regularidad fiscal y sanitaria, temperatura de entrega, plazo, patrón del corte, precio, sustitución de lote y capacidad para atender picos. Al menos dos proveedores deben permanecer homologados para cada insumo crítico."))
    table(doc,[('Grupo' if pt else 'Grupo'),('Fonte primária' if pt else 'Fuente primaria'),('Plano B'),('Controle' if pt else 'Control')],[
        [('Hortifruti'),('CEASA Curitiba'),('Atacadista local' if pt else 'Mayorista local'),('Preço, lote, qualidade e perda' if pt else 'Precio, lote, calidad y merma')],
        [('Carnes' if pt else 'Carnes'),('Distribuidor homologado' if pt else 'Distribuidor homologado'),('Segundo frigorífico' if pt else 'Segundo frigorífico'),('Temperatura, corte, validade' if pt else 'Temperatura, corte, validez')],
        [('Bebidas/embalagens' if pt else 'Bebidas/envases'),('Atacarejo/Linha Verde'),('Distribuidor regional' if pt else 'Distribuidor regional'),('Prazo, volume e ruptura' if pt else 'Plazo, volumen y faltante')],
    ],[1700,2600,2300,2760])
    doc.add_heading('2.5 Validação proposta' if pt else '2.5 Validación propuesta',level=2)
    add_p(doc,("Antes da abertura definitiva, recomenda-se um teste de quatro fins de semana com limite de 25, 35, 45 e 55 combos. O CRM registrará origem do pedido, combo, horário, bairro, canal, atraso, avaliação e intenção de recompra. A decisão de avançar exige venda mínima de 75% da capacidade, desperdício inferior a 5%, pontualidade acima de 90% e margem de contribuição compatível com o plano." if pt else "Antes de la apertura definitiva, se recomienda una prueba de cuatro fines de semana con límite de 25, 35, 45 y 55 combos. El CRM registrará origen del pedido, combo, horario, barrio, canal, retraso, evaluación e intención de recompra. La decisión de avanzar exige vender al menos el 75% de la capacidad, desperdicio inferior al 5%, puntualidad superior al 90% y margen de contribución compatible con el plan."))

    chapter(doc,3,'PLANO DE MARKETING' if pt else 'PLAN DE MARKETING')
    add_p(doc,("O posicionamento proposto é 'o almoço de domingo resolvido': comida familiar, reserva fácil e entrega pontual. A identidade deve comunicar calor, confiança e simplicidade; preço baixo isoladamente não será a promessa, pois compromete qualidade e margem." if pt else "El posicionamiento propuesto es 'el almuerzo del domingo resuelto': comida familiar, reserva fácil y entrega puntual. La identidad debe comunicar calidez, confianza y simplicidad; el precio bajo por sí solo no será la promesa, porque compromete calidad y margen."))
    doc.add_heading('3.1 Composto de marketing' if pt else '3.1 Mezcla de marketing',level=2)
    table(doc,[('P' if pt else 'P'),('Aplicação' if pt else 'Aplicación')],[
        [('Produto' if pt else 'Producto'),('Quatro combos, porções padronizadas, embalagem segura e adicionais controlados.' if pt else 'Cuatro combos, porciones estandarizadas, envase seguro y adicionales controlados.')],
        [('Preço' if pt else 'Precio'),('Custo por ficha técnica, margem de contribuição e comparação local; revisão mensal.' if pt else 'Costo por ficha técnica, margen de contribución y comparación local; revisión mensual.')],
        [('Praça' if pt else 'Plaza'),('Retirada no Umbará e delivery em raio inicial de 5 km por janelas.' if pt else 'Retiro en Umbará y entrega en radio inicial de 5 km por franjas.')],
        [('Promoção' if pt else 'Promoción'),('Conteúdo local, pré-venda, indicação, pós-venda e reativação consentida.' if pt else 'Contenido local, preventa, recomendación, posventa y reactivación consentida.')],
    ],[1300,8060])
    table(doc,[('Combo' if pt else 'Combo'),('Conteúdo resumido' if pt else 'Contenido resumido'),('Preço' if pt else 'Precio'),('CMV')],[[n,('Conforme ficha técnica do Apêndice A' if pt else 'Según ficha técnica del Apéndice A'),money(p),money(c)] for n,q,p,c in mix],[2500,3760,1550,1550])
    doc.add_heading('3.2 Canais, promoção e CRM' if pt else '3.2 Canales, promoción y CRM',level=2)
    add_bullets(doc,[
        ('Google Business Profile e Instagram para descoberta; WhatsApp para conversão e atendimento.' if pt else 'Google Business Profile e Instagram para descubrimiento; WhatsApp para conversión y atención.'),
        ('Pré-venda na sexta-feira, lembrete no sábado e recuperação apenas de clientes com consentimento.' if pt else 'Preventas el viernes, recordatorio el sábado y recuperación solo de clientes con consentimiento.'),
        ('Cupom de segunda compra com validade curta; programa simples de indicação, medido por código.' if pt else 'Cupón de segunda compra con validez corta; programa simple de recomendación, medido por código.'),
        ('Indicadores: taxa de conversão, ticket médio, recompra em 30 dias, custo de aquisição, opt-out e satisfação.' if pt else 'Indicadores: tasa de conversión, ticket promedio, recompra en 30 días, costo de adquisición, bajas y satisfacción.'),
    ])
    add_p(doc,("O CRM Sofia organiza contatos, histórico, preferências, restrições alimentares informadas, consentimento, pedidos, pagamentos, horários e ocorrências. Segmentos iniciais: novos clientes, recorrentes, inativos há 45 dias e compradores de alto ticket. Mensagens devem ter finalidade clara e opção de saída. Dados desnecessários não serão coletados." if pt else "El CRM Sofia organiza contactos, historial, preferencias, restricciones alimentarias informadas, consentimiento, pedidos, pagos, horarios e incidencias. Segmentos iniciales: clientes nuevos, recurrentes, inactivos durante 45 días y compradores de alto ticket. Los mensajes deben tener finalidad clara y opción de baja. No se recogerán datos innecesarios."))
    doc.add_heading('3.3 Jornada do cliente e indicadores' if pt else '3.3 Recorrido del cliente e indicadores',level=2)
    table(doc,[('Etapa' if pt else 'Etapa'),('Experiência desejada' if pt else 'Experiencia deseada'),('Indicador CRM' if pt else 'Indicador CRM')],[
        [('Descoberta' if pt else 'Descubrimiento'),('Ver cardápio e prova social local' if pt else 'Ver menú y prueba social local'),('Origem e custo por contato' if pt else 'Origen y costo por contacto')],
        [('Pedido' if pt else 'Pedido'),('Escolher, pagar e reservar horário em poucos passos' if pt else 'Elegir, pagar y reservar horario en pocos pasos'),('Conversão e abandono' if pt else 'Conversión y abandono')],
        [('Recebimento' if pt else 'Recepción'),('Pedido completo, quente e pontual' if pt else 'Pedido completo, caliente y puntual'),('Atraso, erro e ocorrência' if pt else 'Retraso, error e incidencia')],
        [('Pós-venda' if pt else 'Posventa'),('Avaliação curta e resolução rápida' if pt else 'Evaluación breve y resolución rápida'),('Satisfação e tempo de solução' if pt else 'Satisfacción y tiempo de solución')],
        [('Recompra' if pt else 'Recompra'),('Oferta relevante, sem excesso de mensagens' if pt else 'Oferta relevante, sin exceso de mensajes'),('Recompra 30/60 dias e opt-out' if pt else 'Recompra 30/60 días y bajas')],
    ],[1700,4200,3460])
    add_figure(doc,CHART_DIR/('mix.png' if pt else 'mix_es.png'),('Figura 1 - Mix de vendas do cenário-base.' if pt else 'Figura 1 - Mezcla de ventas del escenario base.'),pt)

    chapter(doc,4,'PLANO OPERACIONAL' if pt else 'PLAN OPERATIVO')
    add_p(doc,("A operação será dimensionada para oito dias regulares por mês, com feriados tratados como capacidade adicional. Sexta-feira concentra compras, recebimento, pré-preparo autorizado e confirmação de reservas; sábado e domingo concentram cocção, montagem, expedição e higienização." if pt else "La operación se dimensionará para ocho días regulares por mes, con los feriados tratados como capacidad adicional. El viernes concentra compras, recepción, preparación previa autorizada y confirmación de reservas; sábado y domingo concentran cocción, montaje, despacho e higienización."))
    doc.add_heading('4.1 Arranjo físico e fluxo sanitário' if pt else '4.1 Distribución física y flujo sanitario',level=2)
    add_p(doc,("O arranjo físico deve impedir cruzamentos entre recebimento de matéria-prima, manipulação de alimentos crus, cocção, montagem de itens prontos e saída de pedidos. A sequência recomendada é: recebimento e inspeção; armazenamento seco/refrigerado; pré-preparo; cocção; bancada limpa de montagem; conferência; expedição. A lavagem de utensílios e o descarte seguem rota separada. O computador do CRM fica protegido de calor, gordura e respingos, sem substituir registros sanitários obrigatórios." if pt else "La distribución física debe impedir cruces entre recepción de materia prima, manipulación de alimentos crudos, cocción, montaje de productos listos y salida de pedidos. La secuencia recomendada es: recepción e inspección; almacenamiento seco/refrigerado; preparación previa; cocción; mesa limpia de montaje; verificación; despacho. El lavado de utensilios y la eliminación de residuos siguen una ruta separada. El computador del CRM permanece protegido del calor, grasa y salpicaduras, sin sustituir los registros sanitarios obligatorios."))
    doc.add_heading('4.2 Capacidade instalada e gargalos' if pt else '4.2 Capacidad instalada y cuellos de botella',level=2)
    add_p(doc,("Para fins de planejamento, cada máquina giratória é considerada capaz de processar aproximadamente 10 frangos por ciclo. Com duas máquinas e dois ciclos úteis, a capacidade teórica é de 40 frangos por dia, antes de ajustes por peso, tempo e especificação do fabricante. A churrasqueira atende costelas, linguiças e pão de alho. Entretanto, o gargalo provável não é apenas a cocção: montagem, conferência e concentração de retiradas entre 11h30 e 13h podem limitar o serviço. Por isso, o CRM distribui pedidos em janelas de 15 minutos e bloqueia horários quando a capacidade é atingida." if pt else "Para fines de planificación, cada máquina giratoria se considera capaz de procesar aproximadamente 10 pollos por ciclo. Con dos máquinas y dos ciclos útiles, la capacidad teórica es de 40 pollos por día, antes de ajustes por peso, tiempo y especificación del fabricante. La parrilla atiende costillas, salchichas y pan de ajo. Sin embargo, el cuello de botella probable no es solo la cocción: montaje, verificación y concentración de retiros entre 11:30 y 13:00 pueden limitar el servicio. Por ello, el CRM distribuye pedidos en franjas de 15 minutos y bloquea horarios cuando se alcanza la capacidad."))
    table(doc,[('Etapa' if pt else 'Etapa'),('Horário/regras' if pt else 'Horario/reglas'),('Controle no CRM' if pt else 'Control en CRM')],[
        [('Planejamento' if pt else 'Planificación'),('Quinta: previsão por reservas e histórico' if pt else 'Jueves: previsión por reservas e historial'),('Demanda por combo e janela' if pt else 'Demanda por combo y franja')],
        [('Compras' if pt else 'Compras'),('Sexta: CEASA/atacadistas e carnes homologadas' if pt else 'Viernes: CEASA/mayoristas y carnes homologadas'),('Lote, fornecedor, custo e validade' if pt else 'Lote, proveedor, costo y validez')],
        [('Produção' if pt else 'Producción'),('Sáb/dom: sequência por tempo de cocção' if pt else 'Sáb/dom: secuencia por tiempo de cocción'),('Status e hora prometida' if pt else 'Estado y hora prometida')],
        [('Expedição' if pt else 'Despacho'),('Conferência dupla; retirada ou rota curta' if pt else 'Doble verificación; retiro o ruta corta'),('Comanda, pagamento e prova de entrega' if pt else 'Comanda, pago y prueba de entrega')],
        [('Pós-venda' if pt else 'Posventa'),('Pesquisa curta após entrega' if pt else 'Encuesta breve tras la entrega'),('NPS/nota, ocorrência e recuperação' if pt else 'NPS/nota, incidencia y recuperación')],
    ],[1700,4260,3400])
    doc.add_heading('4.3 Pessoas e conformidade' if pt else '4.3 Personas y cumplimiento',level=2)
    add_p(doc,("A equipe operacional prevista é composta pelo proprietário/gerente, um churrasqueiro, dois auxiliares e um entregador nos dias de funcionamento. A contratação por diária não elimina risco trabalhista quando existem pessoalidade, habitualidade, subordinação e onerosidade. Portanto, contador e advogado devem definir contratos e jornada adequados; o plano não usa a informalidade como fonte artificial de viabilidade." if pt else "El equipo operativo previsto está compuesto por el propietario/gerente, un asador, dos auxiliares y un repartidor en los días de funcionamiento. La contratación por jornada no elimina el riesgo laboral cuando existen prestación personal, habitualidad, subordinación y remuneración. Por ello, contador y abogado deben definir contratos y jornada adecuados; el plan no utiliza la informalidad como fuente artificial de viabilidad."))
    add_p(doc,("A abertura depende de consulta prévia de viabilidade do endereço, CNPJ, inscrições aplicáveis, alvará/licenciamento municipal, exigências sanitárias, prevenção contra incêndio e procedimentos de boas práticas. Temperatura, higiene, separação de alimentos crus e prontos, água potável, rastreabilidade e descarte devem ser documentados." if pt else "La apertura depende de consulta previa de viabilidad de la dirección, CNPJ, inscripciones aplicables, licencia municipal, exigencias sanitarias, prevención contra incendios y procedimientos de buenas prácticas. Temperatura, higiene, separación de alimentos crudos y listos, agua potable, trazabilidad y descarte deben documentarse."))
    doc.add_heading('4.4 Responsabilidades e rotinas' if pt else '4.4 Responsabilidades y rutinas',level=2)
    table(doc,[('Função' if pt else 'Función'),('Responsabilidade principal' if pt else 'Responsabilidad principal'),('Controle' if pt else 'Control')],[
        [('Gerente'),('Compras, caixa, capacidade, CRM e ocorrências' if pt else 'Compras, caja, capacidad, CRM e incidencias'),('Fechamento diário e painel' if pt else 'Cierre diario y panel')],
        [('Churrasqueiro' if pt else 'Asador'),('Cocção, temperatura e padrão' if pt else 'Cocción, temperatura y estándar'),('Ficha de produção' if pt else 'Ficha de producción')],
        [('Auxiliar 1'),('Pré-preparo e acompanhamentos' if pt else 'Preparación previa y acompañamientos'),('Lotes e checklist' if pt else 'Lotes y lista de control')],
        [('Auxiliar 2'),('Montagem, conferência e higienização' if pt else 'Montaje, verificación e higienización'),('Comanda e checklist' if pt else 'Comanda y lista de control')],
        [('Motoboy' if pt else 'Repartidor'),('Rota, conservação e prova de entrega' if pt else 'Ruta, conservación y prueba de entrega'),('Status e horário' if pt else 'Estado y horario')],
    ],[1900,4700,2760])
    doc.add_heading('4.5 Arquitetura do CRM Sofia' if pt else '4.5 Arquitectura del CRM Sofia',level=2)
    add_p(doc,("A arquitetura mínima prevê VPS, banco de dados, aplicação web, proxy reverso, certificado TLS, rotina de backup e painel no computador local. O bot coleta o pedido, valida disponibilidade e oferece adicional; a cozinha recebe uma comanda única; o gerente acompanha fila, produção e entrega. O uso de integração não oficial com WhatsApp pode gerar bloqueio e descontinuidade. Para produção, recomenda-se canal oficial ou provedor autorizado; a Evolution API pode permanecer apenas em protótipo controlado, sem ser tratada como garantia de custo zero." if pt else "La arquitectura mínima prevé VPS, base de datos, aplicación web, proxy inverso, certificado TLS, rutina de respaldo y panel en el computador local. El bot recoge el pedido, valida disponibilidad y ofrece un adicional; la cocina recibe una comanda única; el gerente acompaña cola, producción y entrega. El uso de una integración no oficial con WhatsApp puede provocar bloqueo y discontinuidad. Para producción, se recomienda un canal oficial o proveedor autorizado; Evolution API puede permanecer solo como prototipo controlado, sin tratarla como garantía de costo cero."))
    doc.add_heading('4.6 Estoque, qualidade e sustentabilidade' if pt else '4.6 Inventario, calidad y sostenibilidad',level=2)
    add_p(doc,("O estoque seguirá o princípio primeiro que vence, primeiro que sai, com registros de lote e validade. Reservas permitem comprar e produzir perto da demanda, reduzindo sobra. A empresa adotará embalagens dimensionadas à porção, separação de óleo e resíduos, consumo racional de água e avaliação futura de embalagens recicláveis compatíveis com calor e segurança do alimento. Sustentabilidade, neste plano, significa reduzir perda e risco sem promessas ambientais não comprovadas." if pt else "El inventario seguirá el principio primero en vencer, primero en salir, con registros de lote y validez. Las reservas permiten comprar y producir cerca de la demanda, reduciendo sobrantes. La empresa adoptará envases dimensionados a la porción, separación de aceite y residuos, consumo racional de agua y evaluación futura de envases reciclables compatibles con calor y seguridad alimentaria. Sostenibilidad, en este plan, significa reducir pérdidas y riesgos sin promesas ambientales no comprobadas."))

    chapter(doc,5,'PLANO FINANCEIRO' if pt else 'PLAN FINANCIERO')
    add_p(doc,("A projeção usa moeda corrente de 2026 e separa investimento, custos variáveis e custos fixos. Embalagens estão incluídas no CMV dos combos. Tributos foram estimados em 4% da receita e meios de pagamento em 2%; ambos devem ser substituídos por alíquotas e contratos reais." if pt else "La proyección usa moneda corriente de 2026 y separa inversión, costos variables y costos fijos. Los envases están incluidos en el CMV de los combos. Los tributos se estimaron en 4% de los ingresos y los medios de pago en 2%; ambos deben sustituirse por las alícuotas y contratos reales."))
    assets=[('2 máquinas giratórias usadas' if pt else '2 máquinas giratorias usadas',3600),('Churrasqueira tradicional' if pt else 'Parrilla tradicional',1200),('Freezer',2200),('Computador',1000),('Lavadora de pressão' if pt else 'Hidrolavadora',700),('Mesa inox',900),('Balança' if pt else 'Balanza',300),('Caixas térmicas' if pt else 'Cajas térmicas',400),('Utensílios' if pt else 'Utensilios',1300)]
    work=[('Caução + primeiro aluguel' if pt else 'Depósito + primer alquiler',2000),('Estoque inicial' if pt else 'Inventario inicial',2300),('Embalagens' if pt else 'Envases',700),('Licenças/contabilidade' if pt else 'Licencias/contabilidad',500),('Comunicação visual' if pt else 'Comunicación visual',600),('Marketing de lançamento' if pt else 'Marketing de lanzamiento',500),('Caixa operacional' if pt else 'Caja operativa',1800)]
    table(doc,[('Investimento fixo' if pt else 'Inversión fija'),'Valor'],[[n,money(v)] for n,v in assets]+[['TOTAL',money(sum(v for n,v in assets))]],[6760,2600])
    table(doc,[('Capital de giro e implantação' if pt else 'Capital de trabajo e implantación'),'Valor'],[[n,money(v)] for n,v in work]+[['TOTAL',money(sum(v for n,v in work))]],[6760,2600])
    add_p(doc,("Fontes: R$ 10.000,00 de capital próprio e R$ 10.000,00 de crédito. Para a simulação, o crédito é pago em 24 parcelas de R$ 500,00, totalizando R$ 12.000,00; a parcela está incluída nos custos fixos. A proposta real de Fomento Paraná deve substituir esta hipótese antes da entrega final." if pt else "Fuentes: R$ 10.000,00 de capital propio y R$ 10.000,00 de crédito. Para la simulación, el crédito se paga en 24 cuotas de R$ 500,00, totalizando R$ 12.000,00; la cuota está incluida en los costos fijos. La propuesta real de Fomento Paraná debe sustituir este supuesto antes de la entrega final."))
    doc.add_page_break()
    table(doc,[('Produto' if pt else 'Producto'),('Qtde.' if pt else 'Cant.'),('Preço unit.' if pt else 'Precio unit.'),('Receita' if pt else 'Ingreso')],[[n,q,money(p),money(q*p)] for n,q,p,c in mix]+[['TOTAL',160,'-',money(revenue)]],[3600,1100,2000,2660])
    table(doc,[('Produto' if pt else 'Producto'),('CMV unit.'),('CMV mensal' if pt else 'CMV mensual'),('Margem unit.' if pt else 'Margen unit.')],[[n,money(c),money(q*c),money(p-c)] for n,q,p,c in mix]+[['TOTAL','-',money(cmv),money((revenue-cmv)/160)]],[3600,1900,1960,1900])
    fixed_rows=[('Diárias da equipe (8 dias)' if pt else 'Jornadas del equipo (8 días)',3840),('Aluguel' if pt else 'Alquiler',1000),('Água, energia e gás' if pt else 'Agua, energía y gas',350),('Internet/telefone' if pt else 'Internet/teléfono',120),('VPS',50),('Contabilidade' if pt else 'Contabilidad',250),('Marketing',200),('Limpeza/manutenção' if pt else 'Limpieza/mantenimiento',180),('Parcela do crédito' if pt else 'Cuota del crédito',500),('Reserva/diversos' if pt else 'Reserva/varios',200)]
    table(doc,[('Custo fixo mensal' if pt else 'Costo fijo mensual'),'Valor'],[[n,money(v)] for n,v in fixed_rows]+[['TOTAL',money(sum(v for n,v in fixed_rows))]],[6760,2600])
    table(doc,[('DRE mensal simplificada' if pt else 'Estado mensual simplificado'),'Valor'],[
        [('Receita bruta' if pt else 'Ingresos brutos'),money(revenue)],
        [('(-) CMV'),money(cmv)],
        [('(-) Tributos estimados (4%)' if pt else '(-) Tributos estimados (4%)'),money(tax)],
        [('(-) Taxas de pagamento (2%)' if pt else '(-) Comisiones de pago (2%)'),money(fees)],
        [('= Margem de contribuição' if pt else '= Margen de contribución'),money(revenue-cmv-tax-fees)],
        [('(-) Custos fixos' if pt else '(-) Costos fijos'),money(fixed)],
        [('= Lucro operacional' if pt else '= Utilidad operativa'),money(profit)],
    ],[6760,2600])
    add_figure(doc,CHART_DIR/('dre.png' if pt else 'dre_es.png'),('Figura 2 - Composição do resultado mensal.' if pt else 'Figura 2 - Composición del resultado mensual.'),pt)
    doc.add_heading('5.1 Fluxo de caixa projetado' if pt else '5.1 Flujo de caja proyectado',level=2)
    cash=[]
    cum=-20000
    for m,q in enumerate(range(160,249,8),1):
        rev=q*(revenue/160); op=rev*cm_ratio-fixed; cum+=op; cash.append([m,q,money(rev),money(op),money(cum)])
    table(doc,[('Mês' if pt else 'Mes'),('Combos'),('Receita' if pt else 'Ingresos'),('Resultado' if pt else 'Resultado'),('Saldo após investimento' if pt else 'Saldo tras inversión')],cash,[900,1200,2200,2200,2860],8)
    add_figure(doc,CHART_DIR/('result12.png' if pt else 'result12_es.png'),('Figura 3 - Evolução projetada do resultado operacional.' if pt else 'Figura 3 - Evolución proyectada del resultado operativo.'),pt)

    chapter(doc,6,'ANÁLISE DE VIABILIDADE' if pt else 'ANÁLISIS DE VIABILIDAD')
    indicators=[
        [('Margem de contribuição' if pt else 'Margen de contribución'),f"{cm_ratio*100:.2f}%".replace('.',',')],
        [('Ponto de equilíbrio mensal' if pt else 'Punto de equilibrio mensual'),money(breakeven)],
        [('Combos equivalentes no equilíbrio' if pt else 'Combos equivalentes en equilibrio'),str(math.ceil(breakeven/(revenue/160)))],
        [('Lucratividade operacional' if pt else 'Rentabilidad operativa'),f"{profit/revenue*100:.2f}%".replace('.',',')],
        [('Payback simples' if pt else 'Recuperación simple'),f"{payback:.2f} meses".replace('.',',')],
    ]
    table(doc,[('Indicador' if pt else 'Indicador'),('Resultado' if pt else 'Resultado')],indicators,[6760,2600])
    add_p(doc,("Os indicadores foram calculados pelas relações: margem de contribuição = receita - CMV - tributos - taxas; índice de margem = margem de contribuição ÷ receita; ponto de equilíbrio = custos fixos ÷ índice de margem; lucratividade = lucro operacional ÷ receita; e payback simples = investimento inicial ÷ lucro operacional mensal. As fórmulas tornam as premissas auditáveis e facilitam a atualização após cotações reais." if pt else "Los indicadores se calcularon mediante las relaciones: margen de contribución = ingresos - CMV - tributos - comisiones; índice de margen = margen de contribución ÷ ingresos; punto de equilibrio = costos fijos ÷ índice de margen; rentabilidad sobre ventas = utilidad operativa ÷ ingresos; y recuperación simple = inversión inicial ÷ utilidad operativa mensual. Las fórmulas vuelven auditables los supuestos y facilitan su actualización después de cotizaciones reales."))
    add_figure(doc,CHART_DIR/('breakeven.png' if pt else 'breakeven_es.png'),('Figura 4 - Ponto de equilíbrio em combos equivalentes.' if pt else 'Figura 4 - Punto de equilibrio en combos equivalentes.'),pt)
    add_p(doc,("O cenário-base apresenta folga de 37 combos sobre o ponto de equilíbrio, mas não suporta perda prolongada de volume. Se as vendas caírem 20% para 128 combos, o resultado mensal aproxima-se de R$ 275,00; se subirem 20% para 192 combos, aproxima-se de R$ 3.756,00, mantidas as proporções. Portanto, reservas e previsão no CRM são controles de viabilidade, não meros recursos promocionais." if pt else "El escenario base presenta una holgura de 37 combos sobre el punto de equilibrio, pero no soporta una pérdida prolongada de volumen. Si las ventas caen un 20% hasta 128 combos, el resultado mensual se aproxima a R$ 275,00; si aumentan un 20% hasta 192 combos, se aproxima a R$ 3.756,00, manteniendo las proporciones. Por tanto, las reservas y la previsión en el CRM son controles de viabilidad, no simples recursos promocionales."))
    doc.add_heading('6.1 Construção de cenários' if pt else '6.1 Construcción de escenarios',level=2)
    scenarios=[]
    for label,factor in [(('Pessimista' if pt else 'Pesimista'),.8),(('Base'),1.0),(('Otimista' if pt else 'Optimista'),1.2)]:
        q=round(160*factor); rev=q*(revenue/160); res=rev*cm_ratio-fixed; scenarios.append([label,q,money(rev),money(res),f"{res/rev*100:.1f}%".replace('.',',')])
    table(doc,[('Cenário' if pt else 'Escenario'),('Combos'),('Receita' if pt else 'Ingresos'),('Resultado' if pt else 'Resultado'),('Margem líquida' if pt else 'Margen neto')],scenarios,[1800,1200,2200,2200,1960])
    add_p(doc,("No cenário pessimista, a empresa permanece ligeiramente positiva, porém sem folga suficiente para imprevistos ou remuneração adicional do proprietário. A gestão deve acionar um gatilho: duas semanas abaixo de 75% da meta exigem redução de produção, revisão de campanhas e análise dos motivos de abandono. O cenário otimista não autoriza expansão automática; primeiro deve-se comprovar que atrasos, reclamações e desperdício permanecem controlados." if pt else "En el escenario pesimista, la empresa permanece ligeramente positiva, pero sin holgura suficiente para imprevistos o remuneración adicional del propietario. La gestión debe activar un gatillo: dos semanas por debajo del 75% de la meta exigen reducción de producción, revisión de campañas y análisis de los motivos de abandono. El escenario optimista no autoriza una expansión automática; primero debe comprobarse que retrasos, reclamaciones y desperdicio permanecen controlados."))
    doc.add_heading('6.2 Matriz SWOT' if pt else '6.2 Matriz FODA',level=2)
    table(doc,[('Forças' if pt else 'Fortalezas'),('Fraquezas' if pt else 'Debilidades')],[
        [('Cardápio enxuto; canal direto; CRM próprio; proximidade.' if pt else 'Menú reducido; canal directo; CRM propio; proximidad.'),('Marca nova; capacidade limitada; dependência do proprietário.' if pt else 'Marca nueva; capacidad limitada; dependencia del propietario.')],
    ],[4680,4680])
    table(doc,[('Oportunidades' if pt else 'Oportunidades'),('Ameaças' if pt else 'Amenazas')],[
        [('Cultura de almoço familiar; delivery; recompra orientada por dados.' if pt else 'Cultura del almuerzo familiar; entrega; recompra orientada por datos.'),('Inflação de alimentos; fiscalização; falha do canal; concorrência.' if pt else 'Inflación de alimentos; fiscalización; falla del canal; competencia.')],
    ],[4680,4680])
    doc.add_page_break(); doc.add_heading('6.3 Riscos e respostas' if pt else '6.3 Riesgos y respuestas',level=2)
    table(doc,[('Risco' if pt else 'Riesgo'),('Prob./impacto' if pt else 'Prob./impacto'),('Resposta' if pt else 'Respuesta')],[
        [('Demanda abaixo de 123 combos' if pt else 'Demanda inferior a 123 combos'),'M/A',('Piloto, reservas, reduzir produção e custo variável' if pt else 'Piloto, reservas, reducir producción y costo variable')],
        [('Aumento de carnes' if pt else 'Aumento de carnes'),'A/A',('Ficha técnica semanal, fornecedores alternativos, reajuste seletivo' if pt else 'Ficha técnica semanal, proveedores alternativos, ajuste selectivo')],
        [('Incidente sanitário' if pt else 'Incidente sanitario'),'B/Muito A' if pt else 'B/Muy A',('Boas práticas, temperatura, rastreabilidade, seguro' if pt else 'Buenas prácticas, temperatura, trazabilidad, seguro')],
        [('Bloqueio de mensageria' if pt else 'Bloqueo de mensajería'),'M/A',('Canal oficial, exportação de contatos consentidos, contingência' if pt else 'Canal oficial, exportación de contactos consentidos, contingencia')],
        [('Vazamento de dados' if pt else 'Fuga de datos'),'B/A',('Mínimo de dados, acesso por função, backup cifrado, resposta LGPD' if pt else 'Mínimo de datos, acceso por función, respaldo cifrado, respuesta LGPD')],
    ],[2600,1500,5260])

    chapter(doc,7,'ANEXOS E INSTRUMENTOS DE IMPLANTAÇÃO' if pt else 'ANEXOS E INSTRUMENTOS DE IMPLANTACIÓN')
    add_p(doc,("Este capítulo reúne instrumentos produzidos para o plano. Por terem sido elaborados pelo autor, são tecnicamente apêndices; documentos externos futuros, como licenças e cotações assinadas, deverão ser apresentados como anexos. As imagens abaixo são ilustrações geradas por IA e não comprovam a existência do imóvel, dos produtos ou do sistema." if pt else "Este capítulo reúne instrumentos producidos para el plan. Por haber sido elaborados por el autor, técnicamente son apéndices; los documentos externos futuros, como licencias y cotizaciones firmadas, deberán presentarse como anexos. Las imágenes siguientes son ilustraciones generadas por IA y no prueban la existencia del inmueble, los productos o el sistema."))
    doc.add_picture(str(IMG),width=Cm(16)); p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0)
    inline = doc.inline_shapes[-1]._inline
    inline.docPr.set('descr', ('Tríptico ilustrativo com fachada, combos de assados e estação do CRM Sofia.' if pt else 'Tríptico ilustrativo con fachada, combos de asados y estación del CRM Sofia.'))
    cap=doc.add_paragraph(('Figura 5 - Conceito ilustrativo: fachada, produtos e estação do CRM Sofia.' if pt else 'Figura 5 - Concepto ilustrativo: fachada, productos y estación del CRM Sofia.')); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.first_line_indent=Cm(0); cap.runs[0].font.size=Pt(10)
    src=doc.add_paragraph(('Fonte: imagem gerada por IA para este rascunho (2026).' if pt else 'Fuente: imagen generada por IA para este borrador (2026).')); src.alignment=WD_ALIGN_PARAGRAPH.CENTER; src.paragraph_format.first_line_indent=Cm(0); src.runs[0].font.size=Pt(9); src.runs[0].italic=True
    doc.add_page_break()
    doc.add_heading('7.1 Plano 5W2H dos primeiros 30 dias' if pt else '7.1 Plan 5W2H de los primeros 30 días',level=2)
    actions=[
        [('Validar endereço' if pt else 'Validar dirección'),('Evitar impedimento' if pt else 'Evitar impedimento'),'Wilkin','D1-D3',('Consulta municipal' if pt else 'Consulta municipal'),money(0)],
        [('Formalizar empresa' if pt else 'Formalizar empresa'),('Regularidade' if pt else 'Regularidad'),'Wilkin/contador','D3-D10',('SLU/ME + licenças' if pt else 'SLU/ME + licencias'),money(500)],
        [('Homologar fornecedores' if pt else 'Homologar proveedores'),('Preço e segurança' if pt else 'Precio y seguridad'),'Wilkin','D5-D12',('3 cotações/fichas' if pt else '3 cotizaciones/fichas'),money(0)],
        [('Configurar CRM' if pt else 'Configurar CRM'),('Pedido e dados' if pt else 'Pedido y datos'),'Wilkin','D7-D18',('MVP, backup, consentimento' if pt else 'MVP, respaldo, consentimiento'),money(50)],
        [('Treinar operação' if pt else 'Capacitar operación'),('Padronizar qualidade' if pt else 'Estandarizar calidad'),'Equipe','D15-D22',('Simulação de serviço' if pt else 'Simulación de servicio'),money(300)],
        [('Executar piloto' if pt else 'Ejecutar piloto'),('Validar demanda' if pt else 'Validar demanda'),'Equipe','D23-D30',('Pré-venda e métricas' if pt else 'Preventa y métricas'),money(500)],
    ]
    table(doc,['O quê/Qué','Por quê/Por qué','Quem/Quién','Quando/Cuándo'],[[a,b,c,d] for a,b,c,d,e,f in actions],[2600,3200,1800,1760],8)
    table(doc,['Ação/Acción','Como/Cómo','Quanto/Cuánto'],[[a,e,f] for a,b,c,d,e,f in actions],[2800,4360,2200],8)
    doc.add_heading('7.2 Questionário de validação (não aplicado)' if pt else '7.2 Cuestionario de validación (no aplicado)',level=2)
    questions=[
        ('Quantas pessoas costumam participar do almoço de fim de semana?' if pt else '¿Cuántas personas suelen participar del almuerzo de fin de semana?'),
        ('Com que frequência compra refeição pronta aos sábados ou domingos?' if pt else '¿Con qué frecuencia compra comida preparada los sábados o domingos?'),
        ('Prefere retirada ou entrega? Qual distância/tempo aceita?' if pt else '¿Prefiere retiro o entrega? ¿Qué distancia/tiempo acepta?'),
        ('Qual combo e faixa de preço parecem mais adequados?' if pt else '¿Qué combo y rango de precio parecen más adecuados?'),
        ('Aceitaria receber o cardápio semanal por mensagem? [ ] Sim [ ] Não' if pt else '¿Aceptaría recibir el menú semanal por mensaje? [ ] Sí [ ] No'),
        ('Quais fatores determinam a recompra: sabor, porção, preço, pontualidade, atendimento?' if pt else '¿Qué factores determinan la recompra: sabor, porción, precio, puntualidad, atención?'),
    ];
    for i,q in enumerate(questions,1): add_p(doc,f"{i}. {q}")
    doc.add_heading('7.3 Dicionário mínimo de dados do CRM' if pt else '7.3 Diccionario mínimo de datos del CRM',level=2)
    table(doc,[('Campo' if pt else 'Campo'),('Finalidade' if pt else 'Finalidad'),('Retenção/controle' if pt else 'Retención/control')],[
        [('Nome e telefone' if pt else 'Nombre y teléfono'),('Atender e identificar pedido' if pt else 'Atender e identificar pedido'),('Acesso restrito; exclusão sob solicitação' if pt else 'Acceso restringido; eliminación a solicitud')],
        [('Endereço' if pt else 'Dirección'),('Entrega' if pt else 'Entrega'),('Somente quando necessário' if pt else 'Solo cuando sea necesario')],
        [('Consentimento' if pt else 'Consentimiento'),('Comprovar autorização de marketing' if pt else 'Probar autorización de marketing'),('Data, canal e finalidade' if pt else 'Fecha, canal y finalidad')],
        [('Itens/valores' if pt else 'Artículos/valores'),('Produção, financeiro e recomendação' if pt else 'Producción, finanzas y recomendación'),('Prazo fiscal e analítico definido' if pt else 'Plazo fiscal y analítico definido')],
        [('Ocorrência/avaliação' if pt else 'Incidencia/evaluación'),('Qualidade e recuperação' if pt else 'Calidad y recuperación'),('Sem dados sensíveis desnecessários' if pt else 'Sin datos sensibles innecesarios')],
    ],[2200,3560,3600])

    doc.add_page_break(); doc.add_heading('CONCLUSÃO' if pt else 'CONCLUSIÓN',level=1)
    add_p(doc,("O estudo indica viabilidade condicionada, não garantia de sucesso. O modelo torna-se interessante porque combina foco operacional e relacionamento direto, mas a margem depende de vender ao menos 123 combos equivalentes por mês, manter CMV por ficha técnica e evitar informalidade trabalhista, sanitária e digital. O CRM é central porque transforma reserva, produção, expedição e pós-venda em um único fluxo mensurável. A recomendação é avançar por etapas: validar endereço e licenças, obter cotações, executar piloto, revisar números e somente então comprometer todo o investimento." if pt else "El estudio indica viabilidad condicionada, no una garantía de éxito. El modelo resulta interesante porque combina enfoque operativo y relación directa, pero el margen depende de vender al menos 123 combos equivalentes al mes, mantener el costo de alimentos mediante fichas técnicas y evitar informalidad laboral, sanitaria y digital. El CRM es central porque transforma reserva, producción, despacho y posventa en un único flujo medible. La recomendación es avanzar por etapas: validar dirección y licencias, obtener cotizaciones, ejecutar el piloto, revisar los números y solo entonces comprometer toda la inversión."))
    doc.add_page_break(); doc.add_heading('REFERÊNCIAS' if pt else 'REFERENCIAS',level=1)
    refs=[
        "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. NBR 14724: informação e documentação - trabalhos acadêmicos - apresentação. Rio de Janeiro: ABNT, 2011.",
        "DOLABELA, Fernando. O segredo de Luísa. São Paulo: Cultura, 2008.",
        "DORNELAS, José Carlos Assis. Empreendedorismo: transformando ideias em negócios. São Paulo: Empreende, 2021.",
        "SCHUMPETER, Joseph A. Teoria do desenvolvimento econômico. São Paulo: Nova Cultural, 1997.",
        "SWIFT, Ronald. CRM: customer relationship management. Rio de Janeiro: Campus, 2001.",
    ]
    for r in refs: add_p(doc,r)
    for org,desc,url in SOURCES: add_p(doc,f"{org}. {desc} Disponível em/Disponible en: {url}. Acesso/Acceso em: 15 ago. 2026.")
    add_p(doc,("FACEMP. Materiais fornecidos no projeto: Estrutura_Plano_de_Negocio_Passo_a_Passo.pptx; manuais de plano de negócio; manuais e modelos de normas ABNT; exemplos de planos de negócio. Consulta em agosto de 2026." if pt else "FACEMP. Materiales suministrados en el proyecto: Estrutura_Plano_de_Negocio_Passo_a_Passo.pptx; manuales de plan de negocios; manuales y modelos de normas ABNT; ejemplos de planes de negocios. Consulta en agosto de 2026."))
    # metadata
    doc.core_properties.title='Casa de Assados Sofia - Plano de Negócio' if pt else 'Casa de Assados Sofia - Plan de Negocios'
    doc.core_properties.author='Wilkin Barban Rosabal'
    doc.core_properties.subject='Trabalho final de graduação - rascunho bilíngue' if pt else 'Trabajo final de graduación - borrador bilingüe'
    doc.save(output)

build('pt',ROOT/'Borrador_Casa_de_Assados_Sofia_Portugues.docx')
build('es',ROOT/'Borrador_Casa_de_Assados_Sofia_Espanol.docx')
print('created', money(revenue), money(profit), money(breakeven), payback)
