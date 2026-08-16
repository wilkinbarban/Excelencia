import os
import shutil
import math
from pathlib import Path
import docx
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT = Path(r"c:\Users\wilki\OneDrive\Documentos\Trabajo de Curso")
CHART_DIR = ROOT / "_work" / "charts"
IMG_ANEXO = ROOT / "_work" / "anexo_casa_assados_sofia.png"

# Financial parameters
mix = [
    ("O Clássico da Sofia", "El Clásico de Sofia", 70, 69.90, 26.50),
    ("Costela Suprema", "Costilla Suprema", 35, 119.90, 48.00),
    ("Dueto Sofia", "Dueto Sofia", 35, 94.90, 36.00),
    ("Kit Churrasco Família", "Kit Parrillero Familia", 20, 169.90, 68.00),
]

revenue = sum(q * p for _, _, q, p, c in mix)
cmv = sum(q * c for _, _, q, p, c in mix)
tax = revenue * 0.04
fees = revenue * 0.02
fixed = 6690.00
profit = revenue - cmv - tax - fees - fixed
cm_ratio = (revenue - cmv - tax - fees) / revenue
breakeven = fixed / cm_ratio
payback = 20000.0 / profit

def money(x):
    return f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="B0B0B0", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxa}" w:type="dxa"/>')
    tcPr.append(tcW)

def build_styled_table(doc, headers, rows, widths=None, font_size=8.5, align_right_cols=None):
    if align_right_cols is None:
        align_right_cols = []
    
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t, color="B0B0B0", sz="5")
    
    # Header Row
    header_tr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
    header_tr.append(tblHeader)
    
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = str(h)
        set_cell_shading(c, "E2ECF6")
        set_cell_margins(c, top=100, bottom=100, left=130, right=130)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else (WD_ALIGN_PARAGRAPH.CENTER if len(str(h)) < 18 else WD_ALIGN_PARAGRAPH.LEFT)
        for r in p.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            
    # Data Rows
    for row_idx, row in enumerate(rows):
        row_cells = t.add_row().cells
        bg_color = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row):
            c = row_cells[i]
            c.text = str(val)
            if bg_color != "FFFFFF":
                set_cell_shading(c, bg_color)
            set_cell_margins(c, top=80, bottom=80, left=130, right=130)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(font_size)
                if row[0] in ("TOTAL", "Total", "TOTAL MENSUAL", "TOTAL MENSAL") or (isinstance(val, str) and "TOTAL" in val):
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
                    
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                set_cell_width(row.cells[i], w)
        tblPr = t._tbl.tblPr
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{sum(widths)}" w:type="dxa"/>')
        tblPr.append(tblW)
        
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)
    p_spacer.paragraph_format.line_spacing = 1.0
    return t

def add_figure_with_caption(doc, img_path, fig_num, title, source, is_pt, width_cm=15.5):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(3)
    
    run = p_img.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    
    label_prefix = f"Figura {fig_num} – "
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Cm(0)
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(2)
    r_lbl = p_cap.add_run(label_prefix)
    r_lbl.bold = True
    r_lbl.font.size = Pt(10)
    r_lbl.font.name = "Arial"
    r_title = p_cap.add_run(title)
    r_title.font.size = Pt(10)
    r_title.font.name = "Arial"
    
    src_prefix = "Fonte: " if is_pt else "Fuente: "
    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_src.paragraph_format.first_line_indent = Cm(0)
    p_src.paragraph_format.space_before = Pt(0)
    p_src.paragraph_format.space_after = Pt(8)
    r_src = p_src.add_run(src_prefix + source)
    r_src.italic = True
    r_src.font.size = Pt(9)
    r_src.font.name = "Arial"

def configure_doc_styles(doc):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    
    for name, size, space_b, space_a in [
        ('Heading 1', 14, 18, 8),
        ('Heading 2', 12, 14, 6),
        ('Heading 3', 12, 10, 4)
    ]:
        s = styles[name]
        s.font.name = 'Arial'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        s.paragraph_format.first_line_indent = Cm(0)
        s.paragraph_format.space_before = Pt(space_b)
        s.paragraph_format.space_after = Pt(space_a)
        s.paragraph_format.keep_with_next = True
        
    styles['Heading 1'].font.all_caps = True
    
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.first_line_indent = Cm(0)
    r_foot = footer.add_run("Casa de Assados Sofia | ")
    r_foot.font.size = Pt(9)
    r_foot.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    
    fld = parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>')
    footer._p.append(fld)

def add_p(doc, text, boldlead=None):
    p = doc.add_paragraph()
    if boldlead and text.startswith(boldlead):
        r_lead = p.add_run(boldlead)
        r_lead.bold = True
        p.add_run(text[len(boldlead):])
    else:
        p.add_run(text)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item)

def add_pretextual_elements(doc, is_pt):
    # Capa
    for text, space, bold, size in [
        ("FACULDADE DE CIÊNCIAS E EMPREENDEDORISMO - FACEMP", 36, True, 12),
        ("CURSO DE BACHARELADO EM ADMINISTRAÇÃO" if is_pt else "CARRERA DE BACHILLERATO EN ADMINISTRACIÓN", 65, True, 12),
        ("WILKIN BARBAN ROSABAL", 80, True, 12),
        ("CASA DE ASSADOS SOFIA", 8, True, 16),
        (("PLANO DE NEGÓCIO PARA IMPLANTAÇÃO DE UMA MICROEMPRESA DE ASSADOS COM GESTÃO POR CRM EM CURITIBA - PR" if is_pt else "PLAN DE NEGOCIOS PARA LA IMPLANTACIÓN DE UNA MICROEMPRESA DE ASADOS CON GESTIÓN POR CRM EN CURITIBA - PR"), 110, True, 12.5),
        ("CURITIBA - PR\n2026", 0, True, 12)
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(space)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = 'Arial'
        r.font.size = Pt(size)
    doc.add_page_break()

    # Folha de Rosto
    p = doc.add_paragraph('WILKIN BARBAN ROSABAL')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(75)
    p.paragraph_format.first_line_indent = Cm(0)
    
    title = ("CASA DE ASSADOS SOFIA: PLANO DE NEGÓCIO PARA UMA OPERAÇÃO DE FIM DE SEMANA APOIADA POR CRM" if is_pt else "CASA DE ASSADOS SOFIA: PLAN DE NEGOCIOS PARA UNA OPERACIÓN DE FIN DE SEMANA APOYADA POR CRM")
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(60)
    p.paragraph_format.first_line_indent = Cm(0)
    
    note = (
        "Trabalho de Conclusão de Curso apresentado ao Curso de Bacharelado em Administração da Faculdade de Ciências e Empreendedorismo - FACEMP, como requisito parcial para obtenção do título de Bacharel em Administração.\n\nÁrea de Concentração: Administração Geral e Empreendedorismo.\n\nOrientador(a): Prof(a). Me./Dr(a). ______________________________"
        if is_pt else
        "Trabajo de Conclusión de Curso presentado a la Carrera de Bachillerato en Administración de la Facultad de Ciencias y Emprendimiento - FACEMP, como requisito parcial para la obtención del título de Bachiller en Administración.\n\nÁrea de Concentración: Administración General y Emprendimiento.\n\nTutor(a): Prof(a). Mg./Dr(a). ______________________________"
    )
    p = doc.add_paragraph(note)
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph('CURITIBA - PR\n2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    doc.add_page_break()

    # Folha de Aprovação
    p = doc.add_paragraph('WILKIN BARBAN ROSABAL')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(40)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(40)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph(note.split('\n\nOrientador')[0])
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    aprov_text = "Aprovado em: _____ / _____ / 2026\n\nBANCA EXAMINADORA:" if is_pt else "Aprobado el: _____ / _____ / 2026\n\nTRIBUNAL EVALUADOR:"
    p = doc.add_paragraph(aprov_text)
    p.runs[0].bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(30)
    
    for member in [
        ("Prof(a). Orientador(a) - FACEMP", "Presidente"),
        ("Prof(a). Avaliador(a) 1 - FACEMP", "Membro 1" if is_pt else "Miembro 1"),
        ("Prof(a). Avaliador(a) 2 - FACEMP", "Membro 2" if is_pt else "Miembro 2")
    ]:
        p = doc.add_paragraph("____________________________________________________\n" + member[0] + " (" + member[1] + ")")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(25)
        p.paragraph_format.line_spacing = 1.0
    doc.add_page_break()

    # Dedicatória & Agradecimentos & Epígrafe
    doc.add_heading("DEDICATÓRIA" if is_pt else "DEDICATORIA", level=1)
    ded_text = (
        "Dedico este trabalho à minha família, cuja paciência, incentivo constante e apoio incondicional foram os alicerces para a superação de cada desafio desta jornada acadêmica. Aos amigos que compreenderam as ausências necessárias e compartilharam os sonhos de empreender com propósito e excelência."
        if is_pt else
        "Dedico este trabajo a mi familia, cuya paciencia, estímulo constante y apoyo incondicional fueron los pilares para superar cada desafío de esta travesía académica. A los amigos que comprendieron las ausencias necesarias y compartieron los sueños de emprender con propósito y excelencia."
    )
    p = doc.add_paragraph(ded_text)
    p.paragraph_format.left_indent = Cm(7)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.line_spacing = 1.2
    doc.add_page_break()

    doc.add_heading("AGRADECIMENTOS" if is_pt else "AGRADECIMIENTOS", level=1)
    agr_text = (
        "Agradeço, em primeiro lugar, a Deus, por iluminar meus passos, conceder-me saúde e perseverança ao longo de toda a formação universitária.\n\n"
        "Aos professores e à coordenação do Curso de Bacharelado em Administração da FACEMP, por compartilharem seus conhecimentos, rigor acadêmico e entusiasmo pela gestão empresarial de excelência.\n\n"
        "Ao meu orientador, pela dedicação, leitura atenta, correções precisas e valiosas diretrizes metodológicas que permitiram transformar uma ideia de negócio em um plano estruturado e viável.\n\n"
        "A todos os colegas de curso, com quem tive a honra de debater, aprender e construir uma visão crítica sobre os desafios do empreendedorismo contemporâneo."
        if is_pt else
        "Agradezco, en primer lugar, a Dios, por iluminar mis pasos, concederme salud y perseverancia a lo largo de toda la formación universitaria.\n\n"
        "A los profesores y a la coordinación de la Carrera de Bachillerato en Administración de la FACEMP, por compartir sus conocimientos, rigor académico y entusiasmo por la gestión empresarial de excelencia.\n\n"
        "A mi tutor docente, por la dedicación, lectura atenta, correcciones precisas y valiosas directrices metodológicas que permitieron transformar una idea de negocio en un plan estructurado y viable.\n\n"
        "A todos los compañeros de carrera, con quienes tuve el honor de debatir, aprender y construir una visión crítica sobre los desafíos del emprendimiento contemporáneo."
    )
    for par in agr_text.split('\n\n'):
        add_p(doc, par)
    doc.add_page_break()

    # Epígrafe
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(7)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(140)
    p.paragraph_format.line_spacing = 1.2
    epi_text = (
        "“O melhor modo de prever o futuro é criá-lo. O planejamento não diz respeito a decisões futuras, mas às implicações futuras de decisões presentes.”\n\n(Peter F. Drucker)"
        if is_pt else
        "“La mejor manera de predecir el futuro es crearlo. La planificación no se refiere a decisiones futuras, sino a las implicaciones futuras de las decisiones presentes.”\n\n(Peter F. Drucker)"
    )
    r = p.add_run(epi_text)
    r.italic = True
    doc.add_page_break()

    # Resumo / Resumen
    doc.add_heading("RESUMO" if is_pt else "RESUMEN", level=1)
    res_text = (
        "Este Trabalho de Conclusão de Curso apresenta o plano de negócio para a implantação da Casa de Assados Sofia, uma microempresa gastronômica projetada para operar exclusivamente aos sábados, domingos e feriados no bairro Umbará, em Curitiba - PR. O modelo de negócio fundamenta-se na oferta enxuta de quatro combos familiares de assados tradicionais (frango recheado, costela bovina ao bafo, costelinha suína e cortes combinados com acompanhamentos artesanais), comercializados via retirada programada no balcão e entrega em domicílio em raio restrito de 5 km. O diferencial estratégico repousa na integração transversal de um sistema próprio de Gestão do Relacionamento com Clientes (CRM Sofia), que gerencia campanhas de pré-venda às sextas-feiras, nivela a capacidade produtiva das churrasqueiras em janelas de 15 minutos, mitiga perdas de matérias-primas perecíveis e potencializa a fidelização e a recompra recorrente com total conformidade à Lei Geral de Proteção de Dados (LGPD). A metodologia adotada possui natureza aplicada, descritiva e documental, combinando análise mercadológica com dados oficiais do IBGE, Prefeitura de Curitiba e SEBRAE, além de modelagem econômico-financeira de custos e receitas. No cenário-base dimensionado para 160 combos mensais, o empreendimento requer um investimento inicial de R$ 20.000,00, gerando receita bruta mensal de R$ 15.809,00, margem de contribuição de 55,16%, lucro operacional líquido de R$ 2.015,46 (lucratividade de 12,75%), ponto de equilíbrio em R$ 12.148,95 (~123 combos equivalentes) e prazo de retorno do investimento (payback simples) de 9,92 meses. Conclui-se que o empreendimento é comercialmente e financeiramente viável, condicionando seu sucesso à rigorosa padronização de fichas técnicas, homologação de fornecedores regionais (CEASA Curitiba) e execução disciplinada dos processos operacionais orientados por dados."
        if is_pt else
        "Este Trabajo de Conclusión de Curso presenta el plan de negocios para la implantación de Casa de Assados Sofia, una microempresa gastronómica proyectada para operar exclusivamente los sábados, domingos y días feriados en el barrio Umbará, en Curitiba - PR. El modelo de negocio se fundamenta en una oferta reducida de cuatro combos familiares de asados tradicionales (pollo relleno, costilla vacuna braseada, costilla de cerdo y cortes combinados con guarniciones artesanales), comercializados mediante retiro programado en mostrador y entrega a domicilio dentro de un radio restringido de 5 km. El diferencial estratégico radica en la integración transversal de un sistema propio de Gestión de Relaciones con Clientes (CRM Sofia), el cual gestiona campañas de preventa los viernes, nivela la capacidad productiva de las parrillas en franjas horarias de 15 minutos, mitiga pérdidas de insumos perecederos y potencia la fidelización y la recompra recurrente con estricta conformidad con la Ley General de Protección de Datos (LGPD). La metodología adoptada posee naturaleza aplicada, descriptiva y documental, combinando análisis de mercado con datos oficiales del IBGE, Municipio de Curitiba y SEBRAE, además de una modelización económico-financiera de costos e ingresos. En el escenario base proyectado para 160 combos mensuales, el emprendimiento requiere una inversión inicial de R$ 20.000,00, generando ingresos brutos mensuales de R$ 15.809,00, un margen de contribución del 55,16%, utilidad operativa neta de R$ 2.015,46 (rentabilidad sobre ventas del 12,75%), punto de equilibrio en R$ 12.148,95 (~123 combos equivalentes) y un plazo de recuperación de la inversión (payback simple) de 9,92 meses. Se concluye que el emprendimiento es comercial y financieramente viable, condicionando su éxito a la rigurosa estandarización de fichas técnicas, homologación de proveedores regionales (CEASA Curitiba) y una ejecución disciplinada de los procesos operativos orientados por datos."
    )
    add_p(doc, res_text)
    kw_text = (
        "Palavras-chave: Plano de Negócio; Gastronomia de Conveniência; Gestão do Relacionamento com Clientes (CRM); Viabilidade Econômico-Financeira; Curitiba."
        if is_pt else
        "Palabras clave: Plan de Negocios; Gastronomía de Conveniencia; Gestión de Relaciones con Clientes (CRM); Viabilidad Económico-Financiera; Curitiba."
    )
    add_p(doc, kw_text, boldlead=("Palavras-chave:" if is_pt else "Palabras clave:"))
    doc.add_page_break()

    # Listas de Figuras e Tabelas
    doc.add_heading("LISTA DE ILUSTRAÇÕES" if is_pt else "LISTA DE ILUSTRACIONES", level=1)
    figs = [
        ("Figura 1 – Mix Mensal de Vendas (Cenário-Base: 160 Combos)", "Figura 1 – Mezcla Mensual de Ventas (Escenario Base: 160 Combos)", 12),
        ("Figura 2 – Composição do Resultado Mensal (DRE Cenário-Base)", "Figura 2 – Composición del Resultado Mensual (Estado de Resultados Base)", 17),
        ("Figura 3 – Gráfico do Ponto de Equilíbrio Operacional", "Figura 3 – Gráfico del Punto de Equilibrio Operativo", 18),
        ("Figura 4 – Projeção do Resultado Operacional em 12 Meses", "Figura 4 – Proyección del Resultado Operativo en 12 Meses", 19),
        ("Figura 5 – Análise de Sensibilidade e Comparação de Cenários", "Figura 5 – Análisis de Sensibilidad y Comparación de Escenarios", 22),
        ("Figura 6 – Conceito Ilustrativo: Fachada, Embalagens, Produtos e Estação CRM Sofia", "Figura 6 – Concepto Ilustrativo: Fachada, Envases, Productos y Estación CRM Sofia", 25)
    ]
    for pt_title, es_title, p_num in figs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(pt_title if is_pt else es_title)
        p.add_run(f"\t{p_num}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    doc.add_page_break()

    doc.add_heading("LISTA DE TABELAS" if is_pt else "LISTA DE TABLAS", level=1)
    tabs = [
        ("Tabela 1 – Síntese da Identificação Institucional da Empresa", "Tabla 1 – Síntesis de la Identificación Institucional de la Empresa", 8),
        ("Tabela 2 – Matriz de Objetivos Estratégicos e Metas por Horizonte", "Tabla 2 – Matriz de Objetivos Estratégicos y Metas por Horizonte", 9),
        ("Tabela 3 – Matriz Comparativa de Concorrentes Diretos e Indiretos", "Tabla 3 – Matriz Comparativa de Competidores Directos e Indirectos", 10),
        ("Tabela 4 – Matriz de Homologação de Fornecedores Estratégicos", "Tabla 4 – Matriz de Homologación de Proveedores Estratégicos", 11),
        ("Tabela 5 – Composto de Marketing (4Ps de Serviços)", "Tabla 5 – Mezcla de Marketing (4Ps de Servicios)", 12),
        ("Tabela 6 – Engenharia de Cardápio, Preços de Venda e CMV Unitário", "Tabla 6 – Ingeniería de Menú, Precios de Venta y CMV Unitario", 13),
        ("Tabela 7 – Mapeamento da Jornada do Cliente e Indicadores do CRM Sofia", "Tabla 7 – Mapeo del Recorrido del Cliente e Indicadores del CRM Sofia", 14),
        ("Tabela 8 – Cronograma de Rotinas Operacionais Semanais e Controle no CRM", "Tabla 8 – Cronograma de Rutinas Operativas Semanales y Control en CRM", 15),
        ("Tabela 9 – Quadro de Funções, Responsabilidades e Rotinas da Equipe", "Tabla 9 – Cuadro de Funciones, Responsabilidades y Rutinas del Equipo", 16),
        ("Tabela 10 – Estimativa do Investimento Fixo Inicial", "Tabla 10 – Estimación de la Inversión Fija Inicial", 17),
        ("Tabela 11 – Estimativa do Capital de Giro e Despesas Pré-Operacionais", "Tabla 11 – Estimación del Capital de Trabajo y Gastos Preoperativos", 17),
        ("Tabela 12 – Projeção de Faturamento Mensal no Cenário-Base", "Tabla 12 – Proyección de Facturación Mensual en el Escenario Base", 18),
        ("Tabela 13 – Estimativa Detalhada dos Custos Fixos Mensais", "Tabla 13 – Estimación Detallada de los Costos Fijos Mensuales", 18),
        ("Tabela 14 – Demonstração do Resultado do Exercício Mensal (DRE Projetada)", "Tabla 14 – Estado de Resultados Mensual (DRE Proyectada)", 19),
        ("Tabela 15 – Projeção de Fluxo de Caixa para 12 Meses de Operação", "Tabla 15 – Proyección de Flujo de Caja para 12 Meses de Operación", 19),
        ("Tabela 16 – Síntese dos Indicadores de Viabilidade Econômico-Financeira", "Tabla 16 – Síntesis de los Indicadores de Viabilidad Económico-Financiera", 20),
        ("Tabela 17 – Análise de Sensibilidade em Três Cenários de Demanda", "Tabla 17 – Análisis de Sensibilidad en Tres Escenarios de Demanda", 21),
        ("Tabela 18 – Matriz Estratégica SWOT / FODA", "Tabla 18 – Matriz Estratégica FODA / SWOT", 22),
        ("Tabela 19 – Matriz de Gestão de Riscos e Planos de Contingência", "Tabla 19 – Matriz de Gestión de Riesgos y Planes de Contingencia", 23),
        ("Tabela 20 – Plano de Ação 5W2H para os Primeiros 30 Dias de Implantação", "Tabla 20 – Plan de Acción 5W2H para los Primeros 30 Días de Implantación", 24),
        ("Tabela 21 – Dicionário de Dados do Sistema CRM Sofia", "Tabla 21 – Diccionario de Datos del Sistema CRM Sofia", 25)
    ]
    for pt_t, es_t, p_num in tabs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(pt_t if is_pt else es_t)
        p.add_run(f"\t{p_num}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    doc.add_page_break()

    # Sumário
    doc.add_heading("SUMÁRIO" if is_pt else "ÍNDICE GENERAL", level=1)
    toc_items = [
        ("INTRODUÇÃO", "INTRODUCCIÓN", 5),
        ("1.1 Contextualização e Justificativa", "1.1 Contextualización y Justificación", 5),
        ("1.2 Problema de Pesquisa", "1.2 Problema de Investigación", 6),
        ("1.3 Objetivos da Pesquisa", "1.3 Objetivos de la Investigación", 6),
        ("1.4 Aspectos Metodológicos", "1.4 Aspectos Metodológicos", 6),
        ("1.5 Fundamentação Teórica", "1.5 Fundamentación Teórica", 7),
        ("1 SUMÁRIO EXECUTIVO", "1 RESUMEN EJECUTIVO", 8),
        ("1.1 Conceito do Negócio e Proposta de Valor", "1.1 Concepto del Negocio y Propuesta de Valor", 8),
        ("1.2 Perfil do Empreendedor e Competências", "1.2 Perfil del Emprendedor y Competencias", 8),
        ("1.3 Missão, Visão e Valores", "1.3 Misión, Visión y Valores", 9),
        ("1.4 Estrutura Jurídica e Enquadramento Tributário", "1.4 Estructura Jurídica y Encuadre Tributario", 9),
        ("1.5 Localização e Instalações", "1.5 Ubicación e Instalaciones", 9),
        ("1.6 Metas e Objetivos por Horizontes", "1.6 Metas y Objetivos por Horizontes", 10),
        ("2 ANÁLISE DE MERCADO", "2 ANÁLISIS DE MERCADO", 10),
        ("2.1 Contexto Econômico e Demográfico Local", "2.1 Contexto Económico y Demográfico Local", 10),
        ("2.2 Dimensionamento do Mercado (TAM, SAM, SOM)", "2.2 Dimensionamiento del Mercado (TAM, SAM, SOM)", 11),
        ("2.3 Segmentação e Comportamento do Consumidor", "2.3 Segmentación y Comportamiento del Consumidor", 11),
        ("2.4 Mapeamento e Análise da Concorrência", "2.4 Mapeo y Análisis de la Competencia", 11),
        ("2.5 Fornecedores e Matriz de Homologação", "2.5 Proveedores y Matriz de Homologación", 12),
        ("2.6 Protocolo de Validação Empírica Preliminar", "2.6 Protocolo de Validación Empírica Preliminar", 12),
        ("3 PLANO DE MARKETING E CRM", "3 PLAN DE MARKETING Y CRM", 13),
        ("3.1 Posicionamento e os 4Ps de Serviços", "3.1 Posicionamiento y los 4Ps de Servicios", 13),
        ("3.2 Engenharia do Cardápio e Fichas Técnicas", "3.2 Ingeniería del Menú y Fichas Técnicas", 13),
        ("3.3 Estratégia de Preço e Margem de Contribuição", "3.3 Estrategia de Precio y Margen de Contribución", 14),
        ("3.4 Praça e Canais de Distribuição", "3.4 Plaza y Canales de Distribución", 14),
        ("3.5 Promoção, Comunicação e Presença Digital", "3.5 Promoción, Comunicación y Presencia Digital", 14),
        ("3.6 Sistema CRM Sofia: Estratégia e Retenção", "3.6 Sistema CRM Sofia: Estrategia y Retención", 15),
        ("3.7 Jornada do Cliente e Indicadores Comerciais", "3.7 Recorrido del Cliente e Indicadores Comerciales", 15),
        ("4 PLANO OPERACIONAL E TECNOLÓGICO", "4 PLAN OPERATIVO Y TECNOLÓGICO", 16),
        ("4.1 Arranjo Físico e Fluxo Sanitário (RDC 216)", "4.1 Distribución Física y Flujo Sanitario (RDC 216)", 16),
        ("4.2 Capacidade Instalada e Dimensionamento", "4.2 Capacidad Instalada y Dimensionamiento", 16),
        ("4.3 Gestão de Gargalos e Balanceamento com CRM", "4.3 Gestión de Cuellos de Botella y Balanceo con CRM", 17),
        ("4.4 Mapeamento do Processo Produtivo Semanal", "4.4 Mapeo del Proceso Productivo Semanal", 17),
        ("4.5 Estrutura Organizacional e Funções", "4.5 Estructura Organizacional y Funciones", 17),
        ("4.6 Requisitos Sanitários e Licenciamento", "4.6 Requisitos Sanitarios y Licenciamiento", 18),
        ("4.7 Arquitetura Tecnológica do CRM Sofia", "4.7 Arquitectura Tecnológica del CRM Sofia", 18),
        ("4.8 Gestão de Estoques (PEPS) e Sustentabilidade", "4.8 Gestión de Inventarios (PEPS) y Sostenibilidad", 19),
        ("5 PLANO FINANCEIRO", "5 PLAN FINANCIERO", 19),
        ("5.1 Investimento Inicial Total", "5.1 Inversión Inicial Total", 19),
        ("5.2 Estrutura de Financiamento", "5.2 Estructura de Financiamiento", 20),
        ("5.3 Custos Variáveis e CMV dos Combos", "5.3 Costos Variables y CMV de los Combos", 20),
        ("5.4 Custos Fixos Mensais Detalhados", "5.4 Costos Fijos Mensuales Detallados", 20),
        ("5.5 Demonstração do Resultado do Exercício (DRE)", "5.5 Estado de Resultados del Ejercicio (DRE)", 21),
        ("5.6 Fluxo de Caixa Projetado para 12 Meses", "5.6 Flujo de Caja Proyectado a 12 Meses", 21),
        ("5.7 Indicadores Financeiros e Ponto de Equilíbrio", "5.7 Indicadores Financieros y Punto de Equilibrio", 22),
        ("6 ANÁLISE DE VIABILIDADE E GESTÃO DE RISCOS", "6 ANÁLISIS DE VIABILIDAD Y GESTIÓN DE RIESGOS", 23),
        ("6.1 Matriz SWOT / FODA Estratégica", "6.1 Matriz FODA / SWOT Estratégica", 23),
        ("6.2 Análise de Sensibilidade em Três Cenários", "6.2 Análisis de Sensibilidad en Tres Escenarios", 23),
        ("6.3 Avaliação dos Indicadores de Viabilidade", "6.3 Evaluación de Indicadores de Viabilidad", 24),
        ("6.4 Matriz de Riscos e Planos de Contingência", "6.4 Matriz de Riesgos y Planes de Contingencia", 24),
        ("7 ANEXOS E INSTRUMENTOS DE IMPLANTAÇÃO", "7 ANEXOS E INSTRUMENTOS DE IMPLANTACIÓN", 25),
        ("7.1 Plano de Ação 5W2H de 30 Dias", "7.1 Plan de Acción 5W2H de 30 Días", 25),
        ("7.2 Questionário Estruturado de Pesquisa", "7.2 Cuestionario Estructurado de Encuesta", 25),
        ("7.3 Dicionário de Dados do CRM Sofia", "7.3 Diccionario de Datos del CRM Sofia", 26),
        ("7.4 Renders e Documentação Fotográfica", "7.4 Renders y Documentación Fotográfica", 26),
        ("CONCLUSÃO", "CONCLUSIÓN", 27),
        ("REFERÊNCIAS", "REFERENCIAS", 28),
    ]
    for pt_item, es_item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(pt_item if is_pt else es_item)
        if pt_item.startswith(('INTRODUÇÃO', '1 ', '2 ', '3 ', '4 ', '5 ', '6 ', '7 ', 'CONCLUSÃO', 'REFERÊNCIAS')):
            r.bold = True
        p.add_run(f"\t{page}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    
    note_p = doc.add_paragraph("Paginação preliminar de referência formatada conforme normas ABNT." if is_pt else "Paginación preliminar de referencia formateada conforme a normas ABNT.")
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_p.paragraph_format.first_line_indent = Cm(0)
    note_p.runs[0].italic = True
    note_p.runs[0].font.size = Pt(9)
    doc.add_page_break()

print('Pretextual module loaded.')
