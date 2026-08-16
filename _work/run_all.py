import subprocess
import sys
from pathlib import Path
import docx

ROOT = Path(r"c:\Users\wilki\OneDrive\Documentos\Trabajo de Curso")

def run_script(script_name):
    print(f"Running {script_name}...")
    res = subprocess.run([sys.executable, str(ROOT / "_work" / script_name)], capture_output=True, text=True, cwd=str(ROOT))
    print(res.stdout)
    if res.stderr:
        print("STDERR:", res.stderr)
    if res.returncode != 0:
        raise RuntimeError(f"Script {script_name} failed with code {res.returncode}")

run_script("build_pt.py")
run_script("build_es.py")

def verify_docx(filename):
    doc_path = ROOT / filename
    doc = docx.Document(doc_path)
    print(f"\n================ Verification for {filename} ================")
    print(f"File exists: {doc_path.exists()} ({doc_path.stat().st_size:,} bytes)")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    print(f"Total Images / Inline Shapes: {len(doc.inline_shapes)}")
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                words += sum(len(p.text.split()) for p in cell.paragraphs)
    print(f"Total Approximate Words (including tables): {words:,}")

verify_docx("Borrador_Casa_de_Assados_Sofia_Portugues.docx")
verify_docx("Borrador_Casa_de_Assados_Sofia_Espanol.docx")
