import os
import shutil
import math
from pathlib import Path
import docx
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT = Path(r"c:\Users\wilki\OneDrive\Documentos\Trabajo de Curso")
CHART_DIR = ROOT / "_work" / "charts"
IMG_ANEXO = ROOT / "_work" / "anexo_casa_assados_sofia.png"
IMG_PLANTA_ES = CHART_DIR / "planta_baixa_sofia_es.png"
IMG_BRAND = CHART_DIR / "brand_mockup_sofia.jpg"
IMG_MENU_PRINT = CHART_DIR / "cardapio_impresso_sofia.jpg"
IMG_MENU_WA = CHART_DIR / "cardapio_whatsapp_sofia.jpg"

# 6 Equipos Individuales
IMG_EQUIP1 = CHART_DIR / "equip1_asadora_gas.jpg"
IMG_EQUIP2 = CHART_DIR / "equip2_churrasqueira_carvao.jpg"
IMG_EQUIP3 = CHART_DIR / "equip3_coifa_industrial.jpg"
IMG_EQUIP4 = CHART_DIR / "equip4_freezer_horizontal.jpg"
IMG_EQUIP5 = CHART_DIR / "equip5_refrigerador_inox.jpg"
IMG_EQUIP6 = CHART_DIR / "equip6_bancada_balanca.jpg"

# 4 Combos Individuales
IMG_COMBO1 = CHART_DIR / "combo1_classico_sofia.jpg"
IMG_COMBO2 = CHART_DIR / "combo2_costela_sofia.jpg"
IMG_COMBO3 = CHART_DIR / "combo3_dueto_sofia.jpg"
IMG_COMBO4 = CHART_DIR / "combo4_familia_sofia.jpg"

mix_es = [
    ("El Clásico de Sofia", 70, 69.90, 26.50, "1 Pollo relleno entero (~1,4kg asado), farofa artesanal crocante (250g), mayonesa casera tradicional de patata (300g). Rinde 3 a 4 personas."),
    ("Costilla Suprema", 35, 119.90, 48.00, "1kg de Costilla vacuna premium asada lentamente al vapor por 6 horas, mandioca a la manteca de botella (300g), vinagreta fresca y farofa de la casa (250g). Rinde 4 personas."),
    ("Dueto Sofia", 35, 94.90, 36.00, "Medio pollo asado dorado + 500g de Costilla de cerdo marinada en hierbas, patatas rústicas doradas (300g) y farofa de la casa (200g). Rinde 3 a 4 personas."),
    ("Kit Parrillero Familia", 20, 169.90, 68.00, "1 Pollo relleno entero + 700g de Costilla vacuna braseada + 4 Chorizos criollos artesanales a la parrilla, mayonesa grande (500g), farofa grande (400g) y panes de ajo (4 un). Rinde 5 a 6 personas."),
]

revenue = sum(q * p for _, q, p, c, _ in mix_es)
cmv = sum(q * c for _, q, p, c, _ in mix_es)
tax = revenue * 0.04
fees = revenue * 0.02
fixed = 6870.00
profit = revenue - cmv - tax - fees - fixed
cm_ratio = (revenue - cmv - tax - fees) / revenue
breakeven = fixed / cm_ratio

def money(x):
    return f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="B0B0B0", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxa}" w:type="dxa"/>')
    tcPr.append(tcW)

def build_styled_table(doc, headers, rows, widths=None, font_size=8.5, align_right_cols=None):
    if align_right_cols is None:
        align_right_cols = []
    
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t, color="B0B0B0", sz="5")
    
    header_tr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
    header_tr.append(tblHeader)
    
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = str(h)
        set_cell_shading(c, "E2ECF6")
        set_cell_margins(c, top=100, bottom=100, left=130, right=130)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else (WD_ALIGN_PARAGRAPH.CENTER if len(str(h)) < 18 else WD_ALIGN_PARAGRAPH.LEFT)
        for r in p.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            
    for row_idx, row in enumerate(rows):
        row_cells = t.add_row().cells
        bg_color = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row):
            c = row_cells[i]
            c.text = str(val)
            if bg_color != "FFFFFF":
                set_cell_shading(c, bg_color)
            set_cell_margins(c, top=80, bottom=80, left=130, right=130)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(font_size)
                if row[0] in ("TOTAL", "Total", "TOTAL MENSUAL", "TOTAL BIENIO (2026-2028)") or (isinstance(val, str) and "TOTAL" in val):
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
                    
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                set_cell_width(row.cells[i], w)
        tblPr = t._tbl.tblPr
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{sum(widths)}" w:type="dxa"/>')
        tblPr.append(tblW)
        
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)
    p_spacer.paragraph_format.line_spacing = 1.0
    return t

def add_figure_with_caption(doc, img_path, fig_num, title, source, width_cm=15.5):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(3)
    
    run = p_img.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Cm(0)
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(2)
    r_lbl = p_cap.add_run(f"Figura {fig_num} – ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(10)
    r_lbl.font.name = "Arial"
    r_title = p_cap.add_run(title)
    r_title.font.size = Pt(10)
    r_title.font.name = "Arial"
    
    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_src.paragraph_format.first_line_indent = Cm(0)
    p_src.paragraph_format.space_before = Pt(0)
    p_src.paragraph_format.space_after = Pt(8)
    r_src = p_src.add_run("Fuente: " + source)
    r_src.italic = True
    r_src.font.size = Pt(9)
    r_src.font.name = "Arial"

def configure_doc_styles(doc):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    
    for name, size, space_b, space_a in [
        ('Heading 1', 14, 18, 8),
        ('Heading 2', 12, 14, 6),
        ('Heading 3', 12, 10, 4)
    ]:
        s = styles[name]
        s.font.name = 'Arial'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        s.paragraph_format.first_line_indent = Cm(0)
        s.paragraph_format.space_before = Pt(space_b)
        s.paragraph_format.space_after = Pt(space_a)
        s.paragraph_format.keep_with_next = True
        
    styles['Heading 1'].font.all_caps = True
    
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.first_line_indent = Cm(0)
    r_foot = footer.add_run("Casa de Assados Sofia | ")
    r_foot.font.size = Pt(9)
    r_foot.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    
    fld = parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>')
    footer._p.append(fld)

def add_p(doc, text, boldlead=None):
    p = doc.add_paragraph()
    if boldlead and text.startswith(boldlead):
        r_lead = p.add_run(boldlead)
        r_lead.bold = True
        p.add_run(text[len(boldlead):])
    else:
        p.add_run(text)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item)

def generate_spanish_thesis(output_path):
    doc = Document()
    configure_doc_styles(doc)
    
    # ------------------ ELEMENTOS PRE-TEXTUALES ------------------
    # Capa
    for text, space, bold, size in [
        ("COLÉGIO EXCELÊNCIA", 36, True, 12),
        ("CARRERA TÉCNICA EN ADMINISTRACIÓN E INFORMÁTICA", 65, True, 12),
        ("WILKIN BARBAN ROSABAL", 80, True, 12),
        ("CASA DE ASSADOS SOFIA", 8, True, 16),
        ("PLAN DE NEGOCIOS PARA LA IMPLANTACIÓN DE UNA MICROEMPRESA DE ASADOS CON GESTIÓN POR CRM EN CURITIBA - PR", 110, True, 12.5),
        ("CURITIBA - PR\n2026", 0, True, 12)
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(space)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = 'Arial'
        r.font.size = Pt(size)
    doc.add_page_break()

    # Portadilla
    p = doc.add_paragraph('WILKIN BARBAN ROSABAL')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(75)
    p.paragraph_format.first_line_indent = Cm(0)
    
    title = "CASA DE ASSADOS SOFIA: PLAN DE NEGOCIOS PARA UNA OPERACIÓN DE FIN DE SEMANA APOYADA POR CRM"
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(60)
    p.paragraph_format.first_line_indent = Cm(0)
    
    note = "Trabajo de Conclusión de Curso presentado a la Carrera Técnica en Administración e Informática del Colégio Excelência, como requisito parcial para la obtención del título de Técnico en Administración e Informática.\n\nÁrea de Concentración: Gestión Empresarial, Emprendimiento y Tecnología de la Información.\n\nTutor(a): Prof(a). Mg./Dr(a). ______________________________"
    p = doc.add_paragraph(note)
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph('CURITIBA - PR\n2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    doc.add_page_break()

    # Hoja de Aprobación
    p = doc.add_paragraph('WILKIN BARBAN ROSABAL')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(40)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(40)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph(note.split('\n\nTutor')[0])
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("Aprobado el: _____ / _____ / 2026\n\nTRIBUNAL EVALUADOR:")
    p.runs[0].bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(30)
    
    for member in [
        ("Prof(a). Tutor(a) - Colégio Excelência", "Presidente"),
        ("Prof(a). Evaluador(a) 1 - Colégio Excelência", "Miembro 1"),
        ("Prof(a). Evaluador(a) 2 - Colégio Excelência", "Miembro 2")
    ]:
        p = doc.add_paragraph("____________________________________________________\n" + member[0] + " (" + member[1] + ")")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(25)
        p.paragraph_format.line_spacing = 1.0
    doc.add_page_break()

    # Dedicatoria
    doc.add_heading("DEDICATORIA", level=1)
    p = doc.add_paragraph("Dedico este trabajo a mi familia, cuya paciencia, estímulo constante y apoyo incondicional fueron los pilares para superar cada desafío de esta travesía académica. A los amigos que comprendieron las ausencias necesarias y compartieron los sueños de emprender con propósito y excelencia.")
    p.paragraph_format.left_indent = Cm(7)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.line_spacing = 1.2
    doc.add_page_break()

    # Agradecimientos
    doc.add_heading("AGRADECIMIENTOS", level=1)
    agr_text = (
        "Agradezco, en primer lugar, a Dios, por iluminar mis pasos, concederme salud y perseverancia a lo largo de toda la formación académica.\n\n"
        "A los profesores y a la coordinación de la Carrera Técnica en Administración e Informática del Colégio Excelência, por compartir sus conocimientos teóricos y prácticos, rigor técnico y entusiasmo por la gestión empresarial orientada por datos.\n\n"
        "A mi tutor docente, por la dedicación, lectura atenta, correcciones precisas y valiosas directrices metodológicas que permitieron transformar una idea de negocio en un plan estructurado, innovador y viable.\n\n"
        "A todos los compañeros de carrera, con quienes tuve el honor de debatir, aprender y construir una visión crítica y moderna sobre los desafíos del emprendimiento contemporáneo."
    )
    for par in agr_text.split('\n\n'):
        add_p(doc, par)
    doc.add_page_break()

    # Epígrafe
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(7)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(140)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run("“La mejor manera de predecir el futuro es crearlo. La planificación no se refiere a decisiones futuras, sino a las implicaciones futuras de las decisiones presentes.”\n\n(Peter F. Drucker)")
    r.italic = True
    doc.add_page_break()

    # Resumen
    doc.add_heading("RESUMEN", level=1)
    res_text = (
        "Este Trabajo de Conclusión de Curso presenta el plan de negocios para la implantación de Casa de Assados Sofia, una microempresa gastronómica proyectada para operar los sábados, domingos y en todos los días feriados nacionales, provinciales y municipales en el barrio Umbará, en Curitiba - PR. El modelo de negocio se fundamenta en una oferta reducida de cuatro combos familiares de asados tradicionales (pollo relleno, costilla vacuna braseada, costilla de cerdo y cortes combinados con guarniciones artesanales), comercializados mediante retiro programado en mostrador y entrega a domicilio dentro de un radio restringido de 5 km. El diferencial estratégico radica en la integración transversal de un sistema propio de Gestión de Relaciones con Clientes (CRM Sofia), el cual gestiona campañas de preventa los viernes y vísperas de feriados, nivela la capacidad productiva de las parrillas en franjas horarias de 15 minutos, mitiga pérdidas de insumos perecederos y potencia la fidelización y la recompra recurrente con estricta conformidad con la Ley General de Protección de Datos (LGPD). La metodología adoptada posee naturaleza aplicada, descriptiva y documental, combinando análisis de mercado con datos oficiales actualizados del IBGE (población estimada en 1.830.795 habitantes y PIB municipal de R$ 120,06 mil millones en Curitiba), Municipio de Curitiba y SEBRAE, además de una modelización económico-financiera desarrollada en el Colégio Excelência. En el escenario base proyectado estrictamente para los fines de semana regulares (160 combos mensuales), el emprendimiento adopta un presupuesto confortable y óptimo de implantación de R$ 38.000,00 (compuesto por R$ 18.000,00 de capital propio y R$ 20.000,00 financiados mediante microcrédito de Fomento Paraná), generando ingresos brutos mensuales de R$ 15.809,00, un margen de contribución del 55,16%, utilidad operativa neta de R$ 1.850,46 (rentabilidad sobre ventas del 11,71%), punto de equilibrio en R$ 12.454,37 (~126 combos equivalentes) y amortización total del capital invertido en la curva dinámica de 12 meses entre el 11º y el 12º mes. Como factor de conservadurismo contable, los ingresos de las operaciones en días feriados hábiles (estimados en 10 a 12 días anuales adicionales, generando R$ 20.000 a R$ 24.000 en facturación extra y R$ 6.250 a R$ 7.500 de utilidad neta líquida) se reservaron como colchón de apalancamiento y liquidez. Se concluye que el emprendimiento es comercial y financieramente viable, condicionando su éxito a la rigurosa estandarización de fichas técnicas, identidad visual atractiva, menús optimizados para salón y WhatsApp, homologación de proveedores regionales (CEASA Curitiba), seguridad jurídica en la contratación intermitente de personal y una ejecución disciplinada de los procesos operativos orientados por datos."
    )
    add_p(doc, res_text)
    add_p(doc, "Palabras clave: Plan de Negocios; Gastronomía de Conveniencia; Gestión de Relaciones con Clientes (CRM); Administración e Informática; Curitiba; Días Feriados.", boldlead="Palabras clave:")
    doc.add_page_break()

    # Listas
    doc.add_heading("LISTA DE ILUSTRACIONES", level=1)
    figs = [
        ("Figura 1 – Mezcla Mensual de Ventas (Escenario Base: 160 Combos)", 12),
        ("Figura 2 – Identidad Visual Fotográfica, Fachada y Comunicación de Punto de Venta", 15),
        ("Figura 3 – Diseño Gráfico del Menú Comercial Impreso para Salón y Mostrador", 16),
        ("Figura 4 – Diseño e Interfaz del Menú Digital Interactivo para WhatsApp y Mobile", 17),
        ("Figura 5 – Composición del Resultado Mensual (Estado de Resultados Base)", 21),
        ("Figura 6 – Gráfico del Punto de Equilibrio Operativo", 22),
        ("Figura 7 – Proyección del Resultado Operativo en 12 Meses", 23),
        ("Figura 8 – Análisis de Sensibilidad y Comparación de Escenarios", 26),
        ("Figura 9 – Máquinas Asadoras Giratorias a Gas GLP con Quemadores Infrarrojos", 29),
        ("Figura 10 – Parrilla Tradicional a Carbón con Tapa Articulada al Vapor y Elevador", 30),
        ("Figura 11 – Sistema de Campana Industrial en Acero Inoxidable con Extracción Mecánica", 30),
        ("Figura 12 – Congelador Horizontal Comercial Doble Función de 510 Litros", 31),
        ("Figura 13 – Refrigerador Comercial Vertical de Acero Inoxidable de 4 Puertas", 31),
        ("Figura 14 – Mesada Central de Manipulación Inox AISI 304 con Balanza Digital", 32),
        ("Figura 15 – Plano Arquitectónico y Flujo Sanitario Unidireccional (60,0 m²)", 33),
        ("Figura 16 – Documentación Fotográfica del Combo 1: El Clásico de Sofia", 37),
        ("Figura 17 – Documentación Fotográfica del Combo 2: Costilla Suprema al Vapor", 38),
        ("Figura 18 – Documentación Fotográfica del Combo 3: Dueto Sofia (Pollo & Costilla de Cerdo)", 38),
        ("Figura 19 – Documentación Fotográfica del Combo 4: Kit Parrillero Familia", 39),
        ("Figura 20 – Concepto Ilustrativo: Fachada, Envases, Productos y Estación CRM Sofia", 40)
    ]
    for es_title, p_num in figs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.add_run(es_title)
        p.add_run(f"\t{p_num}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    doc.add_page_break()

    doc.add_heading("LISTA DE TABLAS", level=1)
    tabs = [
        ("Tabla 1 – Síntesis de la Identificación Institucional de la Empresa", 8),
        ("Tabla 2 – Matriz de Objetivos Estratégicos y Metas por Horizonte", 9),
        ("Tabla 3 – Matriz Comparativa de Competidores Directos e Indirectos", 10),
        ("Tabla 4 – Matriz de Homologación de Proveedores Estratégicos", 11),
        ("Tabla 5 – Mezcla de Marketing (4Ps de Servicios)", 12),
        ("Tabla 6 – Ingeniería de Menú, Precios de Venta y CMV Unitario", 13),
        ("Tabla 7 – Mapeo del Recorrido del Cliente e Indicadores del CRM Sofia", 14),
        ("Tabla 8 – Cronograma de Rutinas Operativas Semanales y Control en CRM", 18),
        ("Tabla 9 – Cuadro de Funciones, Responsabilidades y Rutinas del Equipo", 19),
        ("Tabla 10 – Estrategia de Contratación y Régimen de Trabajo Intermitente (CLT 452-A)", 19),
        ("Tabla 11 – Estimación de la Inversión Fija Inicial (Maquinaria y Exhaustación)", 20),
        ("Tabla 12 – Estimación del Capital de Trabajo y Gastos Preoperativos", 20),
        ("Tabla 13 – Proyección de Facturación Mensual en el Escenario Base", 21),
        ("Tabla 14 – Estimación Detallada de los Costos Fijos Mensuales", 21),
        ("Tabla 15 – Estado de Resultados Mensual (DRE Proyectada)", 22),
        ("Tabla 16 – Proyección de Flujo de Caja para 12 Meses de Operación", 23),
        ("Tabla 17 – Síntesis de los Indicadores de Viabilidad Econômico-Financeira", 24),
        ("Tabla 18 – Cronograma y Proyección de Ingresos Incrementales en Feriados (2026-2028)", 25),
        ("Tabla 19 – Análisis de Sensibilidad en Tres Escenarios de Demanda", 26),
        ("Tabla 20 – Matriz Estratégica FODA / SWOT", 27),
        ("Tabla 21 – Matriz de Gestión de Riesgos y Planes de Contingencia", 28),
        ("Tabla 22 – Plan de Acción 5W2H para los Primeros 30 Días de Implantación", 29),
        ("Tabla 23 – Simulación de Factura Electrónica de Insumos y Carnes (NF-e)", 34),
        ("Tabla 24 – Simulación de Factura Electrónica de Maquinaria y Equipos (NF-e)", 35),
        ("Tabla 25 – Cuadro Síntesis de Licencias, Habilitación y Permisos Sanitarios de Curitiba", 35),
        ("Tabla 26 – Minuta Estructurada del Contrato de Trabajo Intermitente (CLT Art. 452-A)", 36),
        ("Tabla 27 – Modelo de Recibo de Pago por Jornal / Holerite con Desglose Legal e INSS", 36),
        ("Tabla 28 – Diccionario de Datos del Sistema CRM Sofia", 37)
    ]
    for es_t, p_num in tabs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.add_run(es_t)
        p.add_run(f"\t{p_num}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    doc.add_page_break()

    # Sumario
    doc.add_heading("ÍNDICE GENERAL", level=1)
    toc_items = [
        ("INTRODUCCIÓN", 5),
        ("1.1 Contextualización y Justificación", 5),
        ("1.2 Problema de Investigación", 6),
        ("1.3 Objetivos de la Investigación", 6),
        ("1.4 Aspectos Metodológicos", 6),
        ("1.5 Fundamentación Teórica", 7),
        ("1 RESUMEN EJECUTIVO", 8),
        ("1.1 Concepto del Negocio y Propuesta de Valor", 8),
        ("1.2 Perfil del Emprendedor y Competencias", 8),
        ("1.3 Misión, Visión y Valores Organizacionales", 9),
        ("1.4 Estructura Jurídica y Encuadre Tributario", 9),
        ("1.5 Ubicación Estratégica e Instalaciones", 9),
        ("1.6 Metas y Objetivos por Horizontes", 10),
        ("2 ANÁLISIS DE MERCADO", 10),
        ("2.1 Contexto Económico y Demográfico Actualizado de Curitiba y Umbará", 10),
        ("2.2 Dimensionamiento del Mercado (TAM, SAM, SOM)", 11),
        ("2.3 Segmentación y Comportamiento del Consumidor Objetivo", 11),
        ("2.4 Mapeo y Análisis de la Competencia", 11),
        ("2.5 Proveedores Estratégicos y Matriz de Homologación", 12),
        ("2.6 Protocolo de Validación Empírica Preliminar", 12),
        ("3 PLAN DE MARKETING, IDENTIDAD VISUAL Y CRM", 13),
        ("3.1 Posicionamiento Estratégico y los 4Ps de Servicios", 13),
        ("3.2 Ingeniería del Menú y Fichas Técnicas de los Combos", 13),
        ("3.3 Identidad Visual, Slogan y Comunicación de Marca", 14),
        ("3.4 Diseño Gráfico de los Menús: Versión Impresa y Versión WhatsApp", 15),
        ("3.5 Plaza y Canales de Distribución", 17),
        ("3.6 Promoción, Comunicación y Presencia Digital Local", 17),
        ("3.7 Sistema CRM Sofia: Estrategia Transversal y Retención", 18),
        ("3.8 Recorrido del Cliente, Embudo de Conversión e Indicadores", 18),
        ("4 PLAN OPERATIVO Y TECNOLÓGICO", 19),
        ("4.1 Distribución Física y Flujo Sanitario Unidireccional (RDC 216)", 19),
        ("4.2 Capacidad Instalada y Dimensionamiento de Equipos", 19),
        ("4.3 Gestión de Cuellos de Botella y Balanceo con CRM", 20),
        ("4.4 Mapeo del Proceso Productivo Semanal", 20),
        ("4.5 Estructura Organizacional y Estrategia de Contratación de Jornaleros", 20),
        ("4.6 Requisitos Sanitarios y Licenciamiento Municipal", 21),
        ("4.7 Arquitectura Tecnológica del CRM Sofia", 21),
        ("4.8 Gestión de Inventarios (PEPS) y Sostenibilidad", 22),
        ("5 PLAN FINANCIERO", 22),
        ("5.1 Inversión Inicial Total", 22),
        ("5.2 Estructura de Financiamiento y Fuentes de Capital", 23),
        ("5.3 Costos Variables Unitarios y CMV de los Combos", 23),
        ("5.4 Costos Fijos Mensuales Detallados", 23),
        ("5.5 Estado de Resultados del Ejercicio (DRE Proyectada)", 24),
        ("5.6 Flujo de Caja Proyectado a 12 Meses", 24),
        ("5.7 Indicadores Financieros y Punto de Equilibrio", 25),
        ("5.8 Impacto Operativo y Financiero de los Días Feriados como Apalancamiento Adicional (2026-2028)", 25),
        ("6 ANÁLISIS DE VIABILIDAD Y GESTIÓN DE RIESGOS", 26),
        ("6.1 Matriz FODA / SWOT Estratégica", 26),
        ("6.2 Análisis de Sensibilidad en Tres Escenarios", 26),
        ("6.3 Evaluación de Indicadores de Viabilidad", 27),
        ("6.4 Matriz de Riesgos y Planes de Contingencia", 27),
        ("7 ANEXOS E INSTRUMENTOS DE IMPLANTACIÓN", 28),
        ("7.1 Plan de Acción 5W2H de 30 Días", 28),
        ("7.2 Catálogo Fotográfico de Maquinaria y Equipos Adquiridos", 28),
        ("7.3 Plano Arquitectónico y Layout Funcional", 33),
        ("7.4 Simulación de Facturas Electrónicas y Documentos Fiscales (NF-e)", 34),
        ("7.5 Cuadro de Licencias, Habilitación y Permisos Sanitarios", 35),
        ("7.6 Instrumentos de Contratación y Recibos de Jornaleros", 36),
        ("7.7 Cuestionario Estructurado de Encuesta de Mercado", 36),
        ("7.8 Diccionario de Datos del CRM Sofia", 37),
        ("7.9 Renders y Documentación Fotográfica de los Combos Familiares", 37),
        ("CONCLUSIÓN", 41),
        ("REFERENCIAS", 42),
    ]
    for es_item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(es_item)
        if es_item.startswith(('INTRODUCCIÓN', '1 ', '2 ', '3 ', '4 ', '5 ', '6 ', '7 ', 'CONCLUSIÓN', 'REFERENCIAS')):
            r.bold = True
        p.add_run(f"\t{page}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    
    note_p = doc.add_paragraph("Paginación de referencia formateada conforme a normas académicas de la ABNT.")
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_p.paragraph_format.first_line_indent = Cm(0)
    note_p.runs[0].italic = True
    note_p.runs[0].font.size = Pt(9)
    doc.add_page_break()

    # ------------------ ELEMENTOS TEXTUALES ------------------
    # INTRODUCCIÓN
    doc.add_heading("INTRODUCCIÓN", level=1)
    add_p(doc, "El mercado contemporáneo de alimentación fuera del hogar (foodservice) en Brasil atraviesa profundas transformaciones estructurales, impulsadas por la búsqueda incesante de practicidad, la valorización del tiempo en familia durante los fines de semana y la digitalización de los canales de atención y consumo. En el contexto de las principales capitales de la Región Sur, particularmente en Curitiba - PR, el hábito cultural del consumo dominical de carnes asadas (pollo relleno, costilla braseada y asado tradicional) constituye una sólida tradición gastronómica. No obstante, el modelo convencional de los establecimientos de barrio padece frecuentemente graves ineficiencias operativas: largas filas de espera a la intemperie, falta de previsibilidad de la demanda, desperdicio severo de carnes y guarniciones no comercializadas y ausencia absoluta de relación continuada e inteligente con el cliente.")
    add_p(doc, "En este escenario, surge la oportunidad de mercado para concebir Casa de Assados Sofia, una microempresa gastronómica situada en el barrio Umbará, polo residencial en expansión en la zona sur de Curitiba. A diferencia de las rotiserías tradicionales, el emprendimiento propone un modelo operativo ajustado y estructurado, operando con cuatro combos familiares estandarizados, comercializados los sábados, domingos y en todos los días feriados nacionales, provinciales y municipales, respaldado por un sistema propietario de Gestión de Relaciones con Clientes (CRM Sofia).")

    doc.add_heading("1.1 Contextualización y Justificación", level=2)
    add_p(doc, "La justificación para la elaboración de este plan de negocios se fundamenta en la necesidad de planificar científicamente la creación de una microempresa en un sector altamente competitivo y vulnerable a la volatilidad de costos de insumos perecederos. La elección del barrio Umbará, integrante de la Administración Regional Bairro Novo de Curitiba, responde a su densidad poblacional, predominio de núcleos familiares de clase media y escasez de opciones que combinen alta calidad gastronómica artesanal con agilidad y previsibilidad de atención digital.")
    add_p(doc, "Desde el punto de vista académico y de gestión, la relevancia de este trabajo reside en la aplicación práctica de herramientas interdisciplinarias de administración e informática —finanzas, marketing de servicios, investigación operativa, gestión de inventarios y desarrollo de sistemas de información— integradas en la formación técnica del Colégio Excelência. Se demuestra que el uso riguroso de la tecnología y los datos no es privilegio exclusivo de grandes corporaciones, constituyendo, por el contrario, la palanca indispensable para la sostenibilidad, rentabilidad y mitigación de riesgos de una microempresa familiar.")

    doc.add_heading("1.2 Problema de Investigación", level=2)
    add_p(doc, "Frente a las oportunidades y desafíos inherentes al sector de alimentación de conveniencia, se formula el siguiente problema central de investigación: ¿Bajo qué condiciones de mercado, operativas, financieras y tecnológicas resulta viable y sostenible la implantación de Casa de Assados Sofia en el barrio Umbará, en Curitiba - PR?")

    doc.add_heading("1.3 Objetivos de la Investigación", level=2)
    add_p(doc, "Objetivo General:", boldlead="Objetivo General:")
    add_p(doc, "Elaborar un plan de negocios detallado, riguroso e integrado para la implantación y operación de la microempresa gastronómica Casa de Assados Sofia en el barrio Umbará, en Curitiba - PR.")
    add_p(doc, "Objetivos Específicos:", boldlead="Objetivos Específicos:")
    add_bullets(doc, [
        "Realizar el diagnóstico del macroentorno y del mercado local de alimentación de fin de semana en la región sur de Curitiba, mapeando competidores directos, indirectos y perfiles de consumidores con datos oficiales actualizados del IBGE e IPARDES.",
        "Estructurar la estrategia de marketing de servicios (4Ps), la identidad visual y slogan, el diseño de los menús comercial (impreso y digital WhatsApp) y la inserción del sistema CRM Sofia como motor de preventas y fidelización.",
        "Mapear los procesos operativos, sanitarios (RDC 216 Anvisa), la planta arquitectónica con flujo unidireccional, el catálogo de cada equipo real adquirido y la capacidad productiva de las parrillas.",
        "Definir la estrategia jurídica de contratación del equipo operativo de fin de semana (jornaleros bajo régimen de trabajo intermitente - Art. 452-A de la CLT), con minutas y recibos claros para erradicar riesgos de pasivos laborales.",
        "Desarrollar la modelización económico-financiera completa bajo el escenario confortable de R$ 38.000,00 de inversión inicial, proyectando costos fijos y variables, DRE mensual, flujo de caja a 12 meses, punto de equilibrio y payback.",
        "Presentar la proyección de ingresos incrementales generados en días feriados hábiles en el bienio 2026-2028 como mecanismo de apalancamiento financiero y reserva de seguridad.",
        "Presentar la documentación fotográfica detallada e individual de los cuatro combos gastronómicos reales y analizar la viabilidad del negocio mediante matriz FODA y mitigación de riesgos y protección de datos (LGPD)."
    ])

    doc.add_heading("1.4 Aspectos Metodológicos", level=2)
    add_p(doc, "Para el cumplimiento de los objetivos planteados, se adoptó una metodología de investigación aplicada, con enfoque mixto (cualitativo y cuantitativo) y alcance descriptivo-exploratorio. El procedimiento técnico se basó en investigación documental y bibliográfica a partir de las directrices metodológicas del SEBRAE (2013), manuales de orientación técnica del Colégio Excelência y normas académicas de la ABNT.")
    add_p(doc, "La recopilación de datos secundarios se sustentó en estadísticas oficiales recientes del Instituto Brasileño de Geografía y Estadística (IBGE Censo 2022 y estimación 2025), Instituto Paranaense de Desarrollo Económico y Social (IPARDES), registros del Municipio de Curitiba, indicadores de CEASA Paraná e informes sectoriales de la Asociación Brasileña de Bares y Restaurantes (Abrasel). La dimensión cuantitativa se desarrolló mediante ingeniería de menú y modelado de costos unitarios (CMV), construyendo simulaciones de flujo de caja y sensibilidad en modelos matemáticos integrados.")

    doc.add_heading("1.5 Fundamentación Teórica", level=2)
    add_p(doc, "La base conceptual de este trabajo se ancla en la Teoría del Desarrollo Económico de Joseph Schumpeter (1997), quien concibe al emprendedor como el agente de innovación responsable de nuevas combinaciones productivas. En la literatura de gestión contemporánea, Dornelas (2021) y Dolabela (2008) enfatizan que el plan de negocios constituye el instrumento indispensable para transformar visiones intuitivas en estrategias auditables, reduciendo sustancialmente las tasas de mortalidad prematura de las microempresas.")
    add_p(doc, "En el ámbito del marketing y la gestión de clientes, se emplean los postulados de Philip Kotler y Ronald Swift (2001). Swift define el Customer Relationship Management (CRM) como una estrategia empresarial orientada a comprender y anticipar las necesidades de clientes actuales y potenciales. En una operación gastronómica de conveniencia, el CRM actúa como herramienta de nivelación de la demanda y retención, transformando transacciones esporádicas en flujos previsibles de ingresos y reduciendo el Costo de Adquisición de Clientes (CAC) a lo largo del tiempo.")

    # CAPÍTULO 1
    doc.add_page_break()
    doc.add_heading("1 RESUMEN EJECUTIVO", level=1)
    add_p(doc, "Casa de Assados Sofia es un emprendimiento gastronómico concebido para abastecer la creciente demanda de comidas familiares de alta calidad durante los fines de semana y días feriados en Curitiba. Operando bajo el modelo de dark store/takeaway con entrega rápida en radio controlado, la empresa se especializa en cortes asados tradicionales (pollo relleno, costilla vacuna braseada, costilla de cerdo y guarniciones caseras), comercializados exclusivamente en combos listos para servir.")

    doc.add_heading("1.1 Concepto del Negocio y Propuesta de Valor", level=2)
    add_p(doc, "La propuesta de valor se sintetiza en el concepto: 'El almuerzo de domingo de su familia resuelto con excelencia, sabor artesanal y puntualidad británica'. Mientras los competidores tradicionales obligan a los consumidores a soportar largas filas e incertidumbres sobre la disponibilidad de productos, Casa de Assados Sofia opera con reservas programadas mediante el CRM, garantizando retiro sin esperas y entrega en el horario exacto con alimentos calientes y crocantes.")

    build_styled_table(doc, ["Elemento de Identificación", "Definición Estratégica"], [
        ["Razón Social Proyectada", "Casa de Assados Sofia Ltda."],
        ["Nombre Comercial", "Casa de Assados Sofia"],
        ["Forma Jurídica", "Sociedad Limitada Unipersonal (SLU)"],
        ["Encuadre Tributario", "Microempresa (ME) acogida al Simples Nacional"],
        ["Sector de Actividad", "Alimentación Fuera del Hogar / Gastronomía de Conveniencia"],
        ["Ubicación de la Unidad", "Rua Deputado Pinheiro Júnior, 1380, Umbará, Curitiba - PR, Código Postal 81930-000"],
        ["Propietario / Administrador", "Wilkin Barban Rosabal (Técnico en Administración e Informática - Colégio Excelência)"],
        ["Inversión Inicial Total", "R$ 38.000,00 (R$ 18.000,00 Capital Propio + R$ 20.000,00 Microcrédito Fomento Paraná)"],
        ["Diferencial Competitivo", "Operación reducida de 4 combos + Producción programada por CRM propio + Puntualidad"],
    ], widths=[2800, 6560])

    doc.add_heading("1.2 Perfil del Emprendedor y Competencias", level=2)
    add_p(doc, "El emprendedor Wilkin Barban Rosabal asume la administración general del negocio. Su formación técnica interdisciplinaria en Administración e Informática por el Colégio Excelência le confiere dominio en gestión financiera, planificación presupuestaria, modelado de bases de datos, arquitectura de software CRM y optimización de procesos operativos. El administrador será el responsable directo de las compras estratégicas en CEASA, el control de caja, la parametrización del CRM Sofia, la relación con proveedores y la auditoría continua de los estándares sanitarios.")

    doc.add_heading("1.3 Misión, Visión y Valores Organizacionales", level=2)
    add_bullets(doc, [
        "Misión: Brindar a las familias momentos de unión y disfrute en la mesa durante los fines de semana y días feriados, entregando asados artesanales de sabor inigualable, con puntualidad rigurosa y atención humanizada respaldada por tecnología.",
        "Visión: Consolidarse hacia 2029 como la principal referencia en comidas de fin de semana y feriados por encargo y entrega a domicilio en la región sur de Curitiba, destacándose por la consistencia gastronómica y la excelencia en relaciones con el cliente.",
        "Valores: Respeto por la tradición culinaria; Rigor sanitario y transparencia en buenas prácticas; Puntualidad innegociable; Orientación a datos con respeto a la privacidad (LGPD); Sostenibilidad y combate frontal al desperdicio de alimentos."
    ])

    doc.add_heading("1.4 Estructura Jurídica y Encuadre Tributario", level=2)
    add_p(doc, "Se optó por la constitución de una Sociedad Limitada Unipersonal (SLU), encuadrada como Microempresa (ME) bajo el régimen tributario del Simples Nacional. La elección de la SLU garantiza la separación jurídica patrimonial completa entre los bienes personales del socio y las obligaciones de la empresa. Aunque la figura de Microemprendedor Individual (MEI) ofrece facilidades burocráticas, los ingresos anuales proyectados de Casa de Assados Sofia (R$ 189.708,00 en el escenario base) y la necesidad de un equipo de cuatro colaboradores tornan al MEI legalmente inviable, evitando contingencias fiscales severas.")

    doc.add_heading("1.5 Ubicación Estratégica e Instalaciones", level=2)
    add_p(doc, "El local comercial se sitúa en la Rua Deputado Pinheiro Júnior, 1380, en el barrio Umbará. La ubicación ofrece acceso inmediato a las principales arterias colectoras de la región (Rua Nicola Pellanda y Estrada do Ganchinho), facilitando una distribución logística rápida hacia los barrios colindantes (Sítio Cercado, Pinheirinho y Ganchinho). El inmueble de 60 m² fue diseñado para albergar recepción e inspección de materias primas, congelador y refrigeración, área de preparación previa, zona de cocción aislada con sistema de campana y extracción industrial, mesa limpia de montaje, mostrador de despacho rápido y estación de lavado de utensilios.")

    doc.add_heading("1.6 Metas y Objetivos por Horizontes", level=2)
    build_styled_table(doc, ["Horizonte Temporal", "Meta Operativa / Comercial", "Evidencia / Control en CRM"], [
        ["Corto Plazo (0 a 3 meses)", "Validación del menú y piloto comercial; alcanzar la meta base de 160 combos/mes con 90% de puntualidad.", "Registro de pedidos, monitoreo de cuellos de botella, retrasos < 5% y control de mermas < 4%."],
        ["Mediano Plazo (4 a 12 meses)", "Elevación de la tasa de recompra al 45%; alcanzar entre 220 y 248 combos/mes; consolidar margen neto del 14%.", "Análisis de cohortes, historial de consumo por cliente, NPS promedio > 85 y bajas < 1%."],
        ["Largo Plazo (13 a 24 meses)", "Evaluación de compra de 2 máquinas asadoras adicionales; ampliación del radio de entrega a 8 km.", "Informes de demanda insatisfecha por saturación de capacidad, ROI incremental y LTV consolidado."]
    ], widths=[2200, 3600, 3560])

    # CAPÍTULO 2
    doc.add_page_break()
    doc.add_heading("2 ANÁLISIS DE MERCADO", level=1)
    add_p(doc, "El análisis de mercado se fundamenta en el estudio del macroentorno socioeconómico de Curitiba y en la dinámica de consumo de las familias residentes en el barrio Umbará y zonas aledañas.")

    doc.add_heading("2.1 Contexto Económico y Demográfico Actualizado de Curitiba y Umbará", level=2)
    add_p(doc, "Conforme a los datos oficiales más recientes publicados por el IBGE y el IPARDES, Curitiba registró en el Censo Demográfico 2022 una población de 1.773.718 habitantes, alcanzando una estimación actualizada de 1.830.795 habitantes. El Producto Interno Bruto (PIB) municipal totalizó R$ 120,06 mil millones, posicionando a la capital paranaense como la mayor economía municipal de toda la Región Sur de Brasil y la sexta mayor de todo el país. El PIB per cápita oficial más reciente alcanza R$ 67.691,30, reflejando un elevado nivel de ingreso y poder de compra promedio.")
    add_p(doc, "En el ámbito de la microrregión comercial del proyecto, la Administración Regional Bairro Novo —que congrega a los barrios Umbará, Ganchinho y Sítio Cercado— alberga a más de 165.000 habitantes. Se trata de un territorio predominantemente residencial, habitado por familias de clase media (estratos B2, C1 y C2), caracterizadas por jornadas laborales activas de lunes a viernes y una alta predisposición a contratar comidas elaboradas y asados los domingos y feriados.")

    doc.add_heading("2.2 Dimensionamiento del Mercado (TAM, SAM, SOM)", level=2)
    add_bullets(doc, [
        "Mercado Total Disponible (TAM): El universo global de foodservice de fin de semana y feriados en Curitiba, estimado en más de 450.000 hogares consumidores.",
        "Mercado Atendible y Servible (SAM): Hogares ubicados en el radio operativo de 5 km en torno al barrio Umbará, que abarcan aproximadamente 35.000 familias.",
        "Mercado Objetivo Obtenible (SOM): Acotado estrictamente por la capacidad de producción instalada, correspondiente a 160 combos mensuales en el escenario base (40 por fin de semana) y hasta 248 combos mensuales en la fase de madurez."
    ])

    doc.add_heading("2.3 Segmentación y Comportamiento del Consumidor Objetivo", level=2)
    add_p(doc, "El público objetivo prioritario está integrado por núcleos familiares de 3 a 6 integrantes, residentes en casas individuales o condominios horizontales de la zona sur de Curitiba. Sus principales motivaciones de compra son: (a) Ahorro del tiempo y trabajo exigidos para cocinar y limpiar platos el domingo y días festivos; (b) Deseo de consumir asados tradicionales con sabor artesanal de hogar; (c) Búsqueda de puntualidad rigurosa y facilidad de encargo anticipado por mensajería digital.")

    doc.add_heading("2.4 Mapeo y Análisis de la Competencia", level=2)
    build_styled_table(doc, ["Categoría de Competidor", "Principales Fortalezas", "Debilidades / Brechas", "Estrategia Diferenciadora de Sofia"], [
        ["Asadores Tradicionales de Barrio", "Tradición, punto de paso y cercanía vecinal.", "Colas largas, sin reservas, riesgo de quedarse sin producto, pagos difíciles y calidad irregular.", "Reserva anticipada por CRM, franja de retiro de 15 min, combos cerrados y envases sellados."],
        ["Rotiserías de Supermercados", "Precios bajos, gran afluencia de público y economía de escala.", "Carne reseca por permanencia en estufas, atención impersonal y falta de frescura.", "Cocción artesanal en lotes controlados, marinado casero exclusivo y enfoque familiar."],
        ["Plataformas y Dark Kitchens (Apps)", "Variedad ilimitada y comodidad de compra por app genérica.", "Comisiones abusivas (20% al 27%), retrasos constantes en envíos y comida fría/estropeada.", "Canal directo propio por WhatsApp/CRM, entrega con cajas térmicas rígidas en radio de 5 km."]
    ], widths=[2000, 2400, 2400, 2560])

    doc.add_heading("2.5 Proveedores Estratégicos y Matriz de Homologación", level=2)
    add_p(doc, "La política de compras de Casa de Assados Sofia aprovecha la cercanía de la Central de Abastecimiento de Paraná (CEASA Curitiba), ubicada a menos de 10 km, permitiendo la compra directa de frutas y verduras frescas los viernes por la mañana. Para carnes y envases, se estableció una matriz rigurosa de homologación con doble proveedor para asegurar continuidad de suministro.")

    build_styled_table(doc, ["Grupo de Insumo", "Proveedor Principal Homologado", "Proveedor Secundario (Plan B)", "Criterios de Auditoría y Control Sanitario"], [
        ["Aves (Pollos Enfriados)", "Frigorífico Avícola Regional (SIF)", "Distribuidor Mayorista Linha Verde", "Sello SIF/SIPPO, piezas de 1,9-2,1kg, temperatura <= 4°C en recepción."],
        ["Carnes Vacunas (Costilla)", "Frigorífico Vacuno Homologado (PR)", "Mayorista de Carnes Pinheirinho", "Inspección sanitaria, cobertura de grasa uniforme, trazabilidad de lote."],
        ["Carnes Porcinas y Chorizos", "Frigorífico Porcino Castro/PR", "Distribuidor Regional de Embutidos", "Padrón artesanal, control de sal, envasado al vacío intacto."],
        ["Hortalizas (Patatas, Mandioca)", "Productores Directos CEASA Curitiba", "Mayorista de Frutas y Verduras Bairro Novo", "Calibre homogéneo, sin defectos mecánicos y frescura visual."],
        ["Envases Térmicos Sellados", "Distribuidora de Envases PR", "Mayorista Especializado Curitiba", "Resistencia térmica a 90°C, material atóxico y cierre hermético."]
    ], widths=[1800, 2500, 2400, 2660])

    doc.add_heading("2.6 Protocolo de Validación Empírica Preliminar", level=2)
    add_p(doc, "Antes de iniciar las operaciones comerciales regulares, se implementará una prueba piloto de cuatro fines de semana con volúmenes controlados (25, 35, 45 y 55 combos semanales). Este protocolo servirá para sincronizar los tiempos de cocción de las máquinas, validar la aceptación de las recetas, capacitar al personal de empaque y calibrar los flujos de agendamiento del CRM Sofia.")

    # CAPÍTULO 3
    doc.add_page_break()
    doc.add_heading("3 PLAN DE MARKETING, IDENTIDAD VISUAL Y CRM", level=1)
    add_p(doc, "La estrategia de comercialización se centra en posicionar una marca cercana y confiable que no participa en la 'guerra destructiva de precios', sino que ofrece una propuesta superior en sabor, comodidad y puntualidad.")

    doc.add_heading("3.1 Posicionamiento Estratégico y los 4Ps de Servicios", level=2)
    build_styled_table(doc, ["Dimensión (4Ps)", "Directriz Estratégica", "Aplicación Práctica en Casa de Assados Sofia"], [
        ["Producto", "Calidad artesanal superior y estandarización estricta.", "Cuatro combos familiares balanceados, carnes marinadas por 24h, condimento propio y guarniciones frescas."],
        ["Precio", "Precio fijado por valor percibido con margen saludable.", "Precios entre R$ 69,90 y R$ 169,90, logrando un margen de contribución medio ponderado del 55,16%."],
        ["Plaza", "Distribución omnicanal enfocada en comodidad y temperatura.", "Retiro programado en mostrador (takeaway) en Umbará y entrega propia en radio de 5 km en menos de 20 min."],
        ["Promoción", "Comunicación hiperlocal segmentada y CRM de fidelización.", "Perfil de Empresa en Google, Instagram con fotos apetitosas y campañas de preventa semanal por WhatsApp."]
    ], widths=[1600, 3600, 4160])

    doc.add_heading("3.2 Ingeniería del Menú y Fichas Técnicas de los Combos", level=2)
    table_menu_es = []
    for nome, q, p, c, desc in mix_es:
        marg = p - c
        table_menu_es.append([nome, desc, money(p), money(c), money(marg), f"{(marg/p)*100:.1f}%"])
    build_styled_table(doc, ["Combo / Producto", "Composición Detallada y Porciones", "Precio Venta", "CMV Unit.", "Margen R$", "Margen %"], 
                       table_menu_es, widths=[1800, 3800, 1100, 1100, 1100, 900], font_size=8, align_right_cols=[2,3,4,5])

    add_figure_with_caption(doc, CHART_DIR / "mix_es.png", 1, "Mezcla Mensual de Ventas (Escenario Base: 160 Combos)", "Elaboración propia a partir de los supuestos del plan (2026).")

    doc.add_heading("3.3 Identidad Visual, Slogan y Comunicación de Marca", level=2)
    add_p(doc, "La identidad visual de Casa de Assados Sofia fue desarrollada para transmitir calidez de hogar, tradición culinaria rústica y excelencia de servicio. La ambientación y comunicación visual se estructuraron en los siguientes elementos:")
    add_bullets(doc, [
        "Slogan Oficial: “El auténtico sabor del domingo en la mesa de su familia.” — Destaca la importancia del almuerzo dominical en el hogar y la memoria afectiva.",
        "Sub-slogan de Conveniencia: “Tradición artesanal • Reserva sin filas • Entrega puntual” — Resume la propuesta de valor tecnológica y gastronómica.",
        "Fachada Comercial y Letrero 3D Iluminado: Fachada moderna en madera noble tratada con panel en negro carbón, letrero luminoso en acrílico 3D e iluminación escénica cálida.",
        "Paleta Cromática Oficial: Rojo Brasa (#C0392B), Dorado Asado (#D4AC0D), Azul Confianza (#1F3864) y Negro Carbón (#2C3E50).",
        "Piezas de Punto de Venta: Caballete rústico de vereda en madera (1,0m x 0,6m) con menú del día, y bolsas kraft personalizadas con sellos adhesivos de seguridad inviolables (100% caliente)."
    ])
    add_figure_with_caption(doc, IMG_BRAND, 2, "Identidad Visual Fotográfica, Fachada y Comunicación de Punto de Venta", "Mockup fotográfico realista de la fachada, señalización y envases de Casa de Assados Sofia (2026).")

    doc.add_heading("3.4 Diseño Gráfico de los Menús: Versión Impresa y Versión WhatsApp", level=2)
    add_p(doc, "Para atender con máxima eficacia tanto al cliente presencial que acude al mostrador como al usuario digital que realiza encargos remotos, se concibieron dos piezas gráficas complementarias e integradas:")
    add_p(doc, "a) Menú Comercial Impreso de Mostrador (Figura 3): Diseñado en plancha rígida con acabado negro mate y detalles dorados, ideal para exhibición en el mostrador de atención y consulta física. Presenta los cuatro combos, sus porciones y guarniciones artesanales de manera clásica y elegante.")
    add_figure_with_caption(doc, IMG_MENU_PRINT, 3, "Diseño Gráfico del Menú Comercial Impreso para Salón y Mostrador", "Fotografía de referencia del menú comercial impreso en acabado mate y dorado (2026).")

    add_p(doc, "b) Menú Digital Interactivo para WhatsApp y Mobile (Figura 4): Desarrollado con interfaz moderna optimizada para smartphones. Presenta fotografías en alta definición de los platos, tarjetas de selección rápida de combos, selector de franjas horarias de retiro de 15 minutos (ej.: 11:30, 11:45, 12:00) y botón de confirmación de pedido conectado directamente al CRM Sofia.")
    add_figure_with_caption(doc, IMG_MENU_WA, 4, "Diseño e Interfaz del Menú Digital Interactivo para WhatsApp y Mobile", "Interfaz de usuario (UI) mobile desarrollada para atención y preventa vía WhatsApp y CRM Sofia (2026).")

    doc.add_heading("3.5 Plaza y Canales de Distribución", level=2)
    add_p(doc, "El servicio se organiza en dos modalidades perfectamente coordinadas con el CRM:")
    add_bullets(doc, [
        "Retiro Programado (Takeaway): El cliente define su franja horaria (ej.: 11:45 a 12:00) durante la preventa. Al llegar al mostrador, su pedido está empaquetado y conservado en estufa térmica, completando la entrega en menos de 90 segundos.",
        "Entrega a Domicilio Propia: Operación logística circunscrita a 5 km en Umbará y barrios vecinos. Los envíos viajan en cajas térmicas rígidas en rutas cortas agrupadas, garantizando una temperatura de recepción superior a 65°C."
    ])

    doc.add_heading("3.6 Promoción, Comunicación y Presencia Digital Local", level=2)
    add_p(doc, "La captación de nuevos clientes se basa en tácticas de marketing geolocalizado de alta efectividad: (a) Optimización continua de la ficha de Google Business Profile para búsquedas como 'pollo asado Umbará' o 'parrillada de domingo Curitiba'; (b) Cuenta activa de Instagram mostrando el proceso de marinado y cocción los viernes y vísperas de feriados; (c) Degustaciones y folletos promocionales en condominios residenciales y comercios aliados de la zona.")

    doc.add_heading("3.7 Sistema CRM Sofia: Estrategia Transversal, Retención y Ventaja Competitiva", level=2)
    add_p(doc, "El CRM Sofia constituye la columna vertebral tecnológica y de marketing del negocio, diferenciando a Casa de Assados Sofia de cualquier competidor tradicional del sector. Se trata de una solución propietaria desarrollada íntegramente por el propio autor y administrador del emprendimiento (Wilkin Barban Rosabal), integrando las competencias interdisciplinarias de la Carrera Técnica en Administración e Informática del Colégio Excelência.")
    add_p(doc, "La plataforma opera bajo el dominio dinámico gratuito y seguro https://casadeasados.duckdns.org/, con certificado digital SSL/TLS Let's Encrypt para cifrado de extremo a extremo. Construido sobre una arquitectura moderna y 100% respaldada en herramientas de software libre (Linux Ubuntu Server, base de datos relacional PostgreSQL e interfaces web responsivas en HTML5/CSS3/JavaScript), el sistema posee costo de desarrollo nulo y costo de mantenimiento hiper-reducido, siendo el hospedaje del Servidor Virtual Privado (VPS) en la nube (R$ 50,00/mes) el único gasto recurrente de la infraestructura digital.")
    add_p(doc, "Mientras que las plataformas comerciales de SaaS gastronómico cobran mensualidades de entre R$ 300,00 y R$ 800,00 más comisiones sobre pedidos (1% a 3%), y los marketplaces convencionales (como iFood) retienen comisiones abusivas del 20% al 27% sobre la facturación bruta, el CRM Sofia garantiza la soberanía total de los datos de clientes, independencia de intermediarios y un ahorro directo superior a R$ 4.000,00 anuales en software, cumpliendo las siguientes funciones operativas:")
    add_bullets(doc, [
        "Campañas de Preventa Automatizadas los Viernes y Vísperas de Feriados: Mensajes personalizados por canal propio con menú interactivo y enlace directo a https://casadeasados.duckdns.org/ para reservar combos y franjas horarias.",
        "Nivelación de Capacidad en Franjas de 15 Minutos: Algoritmo de agendamiento que limita a 6 pedidos por intervalo, eliminando filas en mostrador y garantizando entregas puntuales con comida caliente.",
        "Mitigación Drástica del Desperdicio de Carnes (< 3%): El volumen de compras de insumos perecederos en CEASA los viernes se calibra con exactitud según los pedidos pre-reservados.",
        "Segmentación RFM (Recencia, Frecuencia y Monto): Clasificación automatizada de clientes VIP (compra semanal), regulares (quincenales) y en riesgo de abandono (inactivos por más de 35 días).",
        "Campañas de Reactivación y Posventa Inteligente: Envío de microencuestas de satisfacción (NPS) 3 horas tras el almuerzo e incentivos de retorno para clientes inactivos.",
        "Conformidad Integral con la LGPD (Ley nº 13.709/2018): Gestión formal de consentimiento (opt-in), canal expreso de baja (opt-out automático), base de datos con campos sensibles cifrados y estricta gobernanza de privacidad."
    ])

    doc.add_heading("3.8 Recorrido del Cliente, Embudo de Conversión e Indicadores", level=2)
    build_styled_table(doc, ["Etapa del Recorrido", "Acción del Cliente", "Punto de Contacto", "Acción del CRM Sofia", "Indicador Clave (KPI)"], [
        ["1. Descubrimiento", "Busca comida para el domingo en Google/Instagram.", "Google Maps / Redes Sociales", "Capta contacto y redirige a WhatsApp con enlace trazable.", "Costo de Adquisición (CAC) y Prospectos."],
        ["2. Preventa", "Recibe menú el viernes y selecciona combo.", "WhatsApp / Bot de Pedidos", "Confirma artículos, ofrece adicionales y reserva franja de retiro.", "Tasa de Conversión de Preventa (> 35%)."],
        ["3. Producción", "Espera el horario programado.", "Cocina del Local", "Emite comanda KDS agrupada por hora de salida de parrilla.", "Nivel de Uso de Capacidad (%)."],
        ["4. Entrega / Retiro", "Retira en mostrador o recibe en su domicilio.", "Mostrador / Repartidor", "Avisa al cliente cuando su pedido sale de la parrilla.", "Puntualidad en Entregas (> 92%) y Tiempos."],
        ["5. Posventa", "Almuerza con su familia.", "Mensaje 3h tras entrega", "Envía microencuesta de satisfacción (NPS de 1 a 5).", "Net Promoter Score (NPS > 85)."],
        ["6. Fidelización", "Recibe sugerencia de compra recurrente.", "WhatsApp el jueves siguiente", "Envía recordatorio basado en su combo preferido del historial.", "Tasa de Recompra a 30 Días (> 40%)."]
    ], widths=[1600, 2000, 1600, 2500, 1660], font_size=8)

    # CAPÍTULO 4
    doc.add_page_break()
    doc.add_heading("4 PLAN OPERATIVO Y TECNOLÓGICO", level=1)
    add_p(doc, "El plan operativo describe la distribución física del local, la capacidad instalada, el flujo del proceso productivo y la estructura tecnológica y sanitaria que aseguran la excelencia operativa de Casa de Assados Sofia.")

    doc.add_heading("4.1 Distribución Física y Flujo Sanitario Unidireccional (RDC 216)", level=2)
    add_p(doc, "En estricto cumplimiento de la Resolución RDC nº 216 de Anvisa y las normativas de Vigilancia Sanitaria de Curitiba, el local fue estructurado con un flujo lineal y unidireccional que prohíbe el cruce entre materias primas crudas y alimentos listos para servir. Las instalaciones disponen de mesadas de acero inoxidable AISI 304, paredes revestidas con azulejos lavables hasta el techo, piso antideslizante con desagües sifonados y un potente sistema profesional de campana y extracción industrial con filtro lavable y ductos hacia el exterior. La planta arquitectónica y flujo sanitario se detallan en la Figura 15.")

    doc.add_heading("4.2 Capacidad Instalada y Dimensionamiento de Equipos", level=2)
    add_p(doc, "El equipamiento productivo real adquirido fue dimensionado para abastecer cómodamente el escenario base de 160 combos mensuales (40 por fin de semana), con holgura para alcanzar 260 combos mensuales sin requerir nuevas inversiones en maquinaria. El catálogo fotográfico de cada equipo adquirido se ilustra en las Figuras 9 a 14:")
    add_bullets(doc, [
        "2 Máquinas Asadoras Giratorias a Gas para Pollo (Figura 9): Equipadas con quemadores infrarrojos a gas GLP, espadas rotativas de acero inoxidable y puertas de vidrio templado. Permiten asar hasta 40 pollos diarios en 2 tandas consecutivas.",
        "1 Parrilla Tradicional a Carbón con Tapa al Vapor y Elevador (Figura 10): Estructura reforzada con ladrillos refractarios y manivela elevadora (1,50m), proyectada para cocción lenta de costillas al vapor por 6 horas.",
        "1 Sistema de Campana Industrial con Extracción Mecánica (Figura 11): Coifa de acero inoxidable con filtros inerciales lavables de alta retención de grasa y ductos de salida, reglamentario por VISA Curitiba.",
        "1 Congelador Horizontal Comercial Doble Función de 510 Litros (Figura 12): Almacenamiento seguro de carnes con registro diario de temperatura (-18°C).",
        "1 Refrigerador Comercial Vertical de Acero Inoxidable de 4 Puertas (Figura 13): Mantenimiento higiénico de marinadas y guarniciones preparadas (+2°C a +4°C).",
        "2 Mesadas Centrales de Acero Inoxidable AISI 304 con Balanza Digital (Figura 14): Superficie aséptica con balanza computadora Inmetro, tábua de corte y cubas gastronómicas GN."
    ])

    doc.add_heading("4.3 Gestión de Cuellos de Botella y Balanceo con CRM", level=2)
    add_p(doc, "El principal cuello de botella de las asaderías dominicales ocurre por la aglomeración de retiros entre las 11:45 y las 12:45. El CRM Sofia elimina este problema dividiendo la capacidad en franjas horarias de 15 minutos (máximo de 6 pedidos por franja). Cuando una franja se completa, el sistema bloquea automáticamente ese horario y ofrece las franjas adyacentes, logrando una carga de trabajo constante de 11:00 a 14:00.")

    doc.add_heading("4.4 Mapeo del Proceso Productivo Semanal", level=2)
    build_styled_table(doc, ["Día de la Semana", "Horario", "Actividades Operativas Clave", "Control y Registro en CRM Sofia"], [
        ["Jueves", "18:00 - 20:00", "Revisión de reservas tempranas y generación de lista de compras.", "Proyección consolidada de insumos por corte y guarnición."],
        ["Viernes", "06:30 - 11:00", "Compras en CEASA y recepción de carnes homologadas con inspección.", "Registro de lotes, fechas de vencimiento y costos unitarios reales."],
        ["Viernes", "13:00 - 18:00", "Porcionado, marinado de carnes y lanzamiento de preventa semanal.", "Envío de campaña por WhatsApp y apertura de franjas horarias."],
        ["Sábado", "06:30 - 10:30", "Encendido de parrillas, inicio de cocción y elaboración de guarniciones.", "Emisión de comandas KDS ordenadas por franja horaria."],
        ["Sábado", "11:00 - 14:30", "Montaje de combos, entrega en mostrador, despacho de repartos y limpieza.", "Confirmación de entrega y seguimiento de tiempos."],
        ["Domingo / Feriado", "06:00 - 15:00", "Jornada operativa principal (concentra el pico de ventas semanales/feriados).", "Cierre de caja y envío automatizado de encuestas NPS."],
        ["Lunes", "09:00 - 11:00", "Higienización profunda, balance de inventario y análisis de métricas.", "Reporte gerencial de ingresos, márgenes y recompras."]
    ], widths=[1800, 1600, 3500, 2460], font_size=8)

    doc.add_heading("4.5 Estructura Organizacional y Estrategia de Contratación de Jornaleros", level=2)
    add_p(doc, "La operación funciona los sábados, domingos y feriados (8 a 10 días al mes). La contratación informal de jornaleros ('diaristas') conlleva un alto riesgo de pasivo laboral en la Justicia del Trabajo brasileña (TRT 9ª Región / TST) debido a la habitualidad y subordinación. Para brindar seguridad jurídica total, la empresa utilizará el Contrato de Trabajo Intermitente (Art. 452-A de la CLT, Ley nº 13.467/2017) o Contratos de Prestación de Servicios Autónomos con Recibo de Pago a Autónomo (RPA).")

    build_styled_table(doc, ["Cargo / Función", "Colaborador / Vínculo", "Jornada Mensual", "Valor Jornal", "Estrategia Jurídico-Laboral (Seguridad Total)"], [
        ["Gerente General", "Wilkin Barban (Socio)", "Integral", "Honorarios", "Socio administrador con responsabilidad civil y técnica."],
        ["Asador Principal", "Jornalero Especialista", "8-10 jornales/mes", "R$ 120,00", "Contrato Intermitente CLT (Art. 452-A) + ASO con exámenes coprológicos."],
        ["Auxiliar de Cocina 1", "Jornalero de Preparación", "8-10 jornales/mes", "R$ 120,00", "Contrato Intermitente CLT (Art. 452-A) + Capacitación RDC 216 y EPI completo."],
        ["Auxiliar de Montaje", "Jornalero de Despacho", "8-10 jornales/mes", "R$ 120,00", "Contrato Intermitente CLT (Art. 452-A) + Convocatoria vía CRM y registro de asistencia."],
        ["Repartidor de Envíos", "Repartidor Asociado", "8-10 jornales/mes", "R$ 120,00", "Contrato Autónomo con MEI propio + Licencia categoría A + Seguro de accidentes."]
    ], widths=[1800, 1800, 1400, 1300, 3060], font_size=8)

    doc.add_heading("4.6 Requisitos Sanitarios y Licenciamiento Municipal", level=2)
    add_p(doc, "La apertura formal cumplirá los trámites legales de Curitiba: (a) Consulta Previa de Viabilidad de Domicilio aprobada en la Secretaría Municipal de Urbanismo; (b) Inscripción Municipal y Licencia de Actividad; (c) Licencia Sanitaria otorgada por la Vigilancia Sanitaria Municipal; (d) Certificado de Aprobación del Cuerpo de Bomberos de Paraná (CLCB); (e) Manual de Buenas Prácticas y Procedimientos Operativos Estandarizados (POEs) disponibles en el local.")

    doc.add_heading("4.7 Arquitectura Tecnológica del CRM Sofia e Infraestructura de Bajo Costo", level=2)
    add_p(doc, "La arquitectura tecnológica del CRM Sofia fue íntegramente diseñada e implementada bajo el paradigma de microservicios ligeros y tecnologías de código abierto (Open Source), garantizando máximo rendimiento con un costo operativo marginal:")
    add_bullets(doc, [
        "Dominio Dinámico y Certificado SSL Gratuito: El sistema es accesible públicamente a través de la URL https://casadeasados.duckdns.org/, empleando el servicio de DNS dinámico gratuito (DuckDNS) con certificado TLS/SSL Let's Encrypt para tráfico 100% cifrado (HTTPS).",
        "Infraestructura en la Nube de Costo Fijo Mínimo: Hospedado en un Servidor Virtual Privado (VPS) basado en Linux Ubuntu Server LTS, cuyo costo mensual de solo R$ 50,00 representa la totalidad de los gastos de TI del negocio.",
        "Stack de Software Libre (Cero Costo de Licencias): Backend desarrollado en Python/FastAPI de alta concurrencia, base de datos relacional PostgreSQL para integridad transaccional (ACID) y frontend responsivo en HTML5, CSS3 y JavaScript puro sin dependencia de librerías propietarias.",
        "Panel KDS Operativo para Cocina: Interfaz en tiempo real que organiza las comandas de preparación por orden cronológico y franjas de retiro de 15 minutos, mostrando contadores regresivos para los asadores.",
        "Mecanismo de Respaldo y Seguridad Criptográfica: Rutina diaria automatizada de volcado de la base de datos con cifrado AES-256 y sincronización segura con almacenamiento remoto, garantizando un RPO (Recovery Point Objective) menor a 24 horas y RTO (Recovery Time Objective) de 15 minutos."
    ])

    doc.add_heading("4.8 Gestión de Inventarios (PEPS) y Sostenibilidad", level=2)
    add_p(doc, "El inventario se controla mediante el criterio PEPS (Primero en Entrar, Primero en Salir). Gracias a la preventa, el volumen de compras de carne del viernes cubre con exactitud el 90% de los pedidos reservados, reduciendo las mermas a menos del 3%. Asimismo, el aceite vegetal utilizado se almacena en bidones sellados y se entrega a una recicladora autorizada para producir biodiésel, y las cenizas de carbón se destinan a abono vegetal.")

    # CAPÍTULO 5
    doc.add_page_break()
    doc.add_heading("5 PLAN FINANCIERO", level=1)
    add_p(doc, "El plan financiero consolida las estimaciones de ingresos, egresos e inversiones expresadas en moneda corriente de 2026 bajo el escenario confortable y óptimo de R$ 38.000,00, demostrando el equilibrio contable y la rentabilidad del proyecto.")

    doc.add_heading("5.1 Inversión Inicial Total", level=2)
    build_styled_table(doc, ["Rubro de Inversión Fija (Maquinaria e Infraestructura)", "Cant", "Costo Unit. (R$)", "Subtotal (R$)"], [
        ["Máquinas asadoras giratorias a gas para pollo (nuevas c/ garantía de fábrica)", "2", "2.400,00", "4.800,00"],
        ["Parrilla braseadora a carbón con sistema elevador (1,5m reforzada)", "1", "2.200,00", "2.200,00"],
        ["Sistema de campana industrial con extracción mecánica y ducto (VISA)", "1", "4.200,00", "4.200,00"],
        ["Congelador horizontal comercial doble función 510L nuevo", "1", "3.100,00", "3.100,00"],
        ["Refrigerador comercial vertical de acero inoxidable 4 puertas nuevo", "1", "3.400,00", "3.400,00"],
        ["Mesadas centrales de trabajo en acero inoxidable AISI 304 (2,0x0,9m)", "2", "1.100,00", "2.200,00"],
        ["Computador / Terminal de mostrador touch con impresora térmica 80mm", "1", "1.500,00", "1.500,00"],
        ["Balanza digital comercial computadora con batería (homologada Inmetro)", "1", "500,00", "500,00"],
        ["Cajas térmicas reforzadas tipo baúl para motoboy (45L)", "2", "250,00", "500,00"],
        ["Hidrolavadora de alta presión profesional para limpieza pesada", "1", "800,00", "800,00"],
        ["Utensilios de cocina en acero inoxidable, cuchillos y recipientes Gastronorm", "Varios", "1.300,00", "1.300,00"],
        ["TOTAL DE INVERSIÓN FIJA", "-", "-", "24.500,00"]
    ], widths=[4800, 1000, 1800, 1760], align_right_cols=[1,2,3])

    build_styled_table(doc, ["Capital de Trabajo y Gastos Preoperativos", "Destino del Recurso", "Valor Estimado (R$)"], [
        ["Depósito de Garantía de Alquiler (3 meses) + 1er Mes de Alquiler", "Fianza locativa de 60m² en Umbará", "4.000,00"],
        ["Inventario Inicial de Carnes e Insumos para Piloto Comercial", "Materias primas y adobos de arranque", "2.500,00"],
        ["Lote Inicial de Envases Térmicos Sellados y Bolsas (1.000 un)", "Envases de seguridad para 1.000 pedidos", "1.400,00"],
        ["Tasas de Licenciamiento, Vistoria Bomberos y Apertura SLU", "Legalización comercial, sanitaria y bomberos", "800,00"],
        ["Cartel de Fachada en Madera, Letrero 3D y Pizarra de Acera", "Identidad visual del punto de venta", "1.200,00"],
        ["Marketing de Lanzamiento, Sesión Fotográfica y Degustación", "Campaña local de inauguración y captación", "1.100,00"],
        ["Fondo de Reserva de Capital de Trabajo Libre (Liquidez)", "Colchón de seguridad para primeros meses", "2.500,00"],
        ["TOTAL DE CAPITAL DE TRABAJO Y PREOPERATIVOS", "-", "13.500,00"],
        ["INVERSIÓN TOTAL REQUERIDA (FIJA + CAPITAL DE TRABAJO)", "-", "38.000,00"]
    ], widths=[4500, 3100, 1760], align_right_cols=[2])

    doc.add_heading("5.2 Estructura de Financiamiento y Fuentes de Capital", level=2)
    add_p(doc, "El capital de R$ 38.000,00 se integrará con R$ 18.000,00 (47,37%) de fondos propios del emprendedor y R$ 20.000,00 (52,63%) provenientes de la línea de microcrédito productivo orientado de Fomento Paraná (Banco do Empreendedor), financiado a 36 cuotas fijas de R$ 680,00 mensuales (tasa subsidiada para microempresas de alimentación), contempladas íntegramente en los costos fijos del proyecto.")

    doc.add_heading("5.3 Costos Variables Unitarios y CMV de los Combos", level=2)
    doc.add_heading("5.4 Costos Fijos Mensuales Detallados", level=2)
    build_styled_table(doc, ["Concepto de Costo Fijo", "Detalle y Base de Cálculo", "Valor Mensual (R$)"], [
        ["Mano de Obra Operativa (Personal Jornalero)", "4 personas x 8 jornadas de fin de semana/mes x R$ 120,00/jornada", "3.840,00"],
        ["Alquiler del Local Comercial", "Inmueble comercial de 60 m² en Rua Deputado Pinheiro Júnior", "1.000,00"],
        ["Servicios Básicos (Agua, Electricidad, Gas)", "Consumo operativo de agua potable, luz y gas GLP industrial", "350,00"],
        ["Servicios de Internet Fibra y Telefonía", "Línea comercial móvil + Conexión de fibra de 500 Mbps", "120,00"],
        ["Infraestructura Cloud y Servidor VPS CRM", "Hospedaje de servidor, base de datos y copias de seguridad", "50,00"],
        ["Honorarios Contables Mensuales", "Gestión contable, fiscal y laboral para Simples Nacional", "250,00"],
        ["Publicidad y Marketing Recurrente", "Anuncios locales en Instagram y material informativo barrial", "200,00"],
        ["Mantenimiento y Productos de Desinfección", "Detergentes profesionales, sanitizantes y revisiones técnicas", "180,00"],
        ["Cuota del Microcrédito Fomento Paraná", "Amortización de préstamo de R$ 20.000 (36 cuotas fijas)", "680,00"],
        ["Fondo de Reserva para Contingencias", "Provisión para reemplazos menores e imprevistos", "200,00"],
        ["TOTAL DE COSTOS FIJOS MENSUALES", "-", "6.870,00"]
    ], widths=[3400, 4200, 1760], align_right_cols=[2])

    doc.add_heading("5.5 Estado de Resultados del Ejercicio (DRE Proyectada)", level=2)
    build_styled_table(doc, ["Línea del Estado Financiero", "Base de Cálculo / Criterio", "Valor Mensual (R$)", "Análisis Vertical (%)"], [
        ["(=) INGRESOS BRUTOS OPERACIONALES", "160 combos en la mezcla proyectada", "15.809,00", "100,00%"],
        ["(-) Costo de Alimentos Vendidos (CMV)", "Suma ponderada de carnes, guarniciones y envases", "6.140,00", "38,84%"],
        ["(-) Impuestos del Simples Nacional (4,0%)", "Alícuota efectiva de microempresa comercial", "632,36", "4,00%"],
        ["(-) Comisiones de Medios de Pago (2,0%)", "Promedio ponderado de tarjetas de débito/crédito y PIX", "316,18", "2,00%"],
        ["(=) MARGEN DE CONTRIBUCIÓN TOTAL", "Ingresos Brutos - Costos Variables - Impuestos/Tasas", "8.720,46", "55,16%"],
        ["(-) COSTOS FIJOS OPERATIVOS TOTALES", "Estructura fija mensual detallada", "6.870,00", "43,46%"],
        ["(=) UTILIDAD OPERATIVA NETA", "Margen de Contribución - Costos Fijos", "1.850,46", "11,71%"]
    ], widths=[3400, 3100, 1600, 1260], align_right_cols=[2,3])

    add_figure_with_caption(doc, CHART_DIR / "dre_es.png", 5, "Composición del Resultado Mensual (Estado de Resultados Base)", "Elaboración propia a partir de los supuestos financieros del plan (2026).")

    doc.add_heading("5.6 Flujo de Caja Proyectado a 12 Meses", level=2)
    cash_rows_es = []
    cum = -38000.0
    for m, q in enumerate(range(160, 249, 8), 1):
        rev = q * (revenue / 160.0)
        op = rev * cm_ratio - fixed
        cum += op
        cash_rows_es.append([f"Mes {m}", str(q), money(rev), money(rev * (1 - cm_ratio)), money(fixed), money(op), money(cum)])
    build_styled_table(doc, ["Mes", "Combos", "Ingresos (R$)", "Costos Var. (R$)", "Costos Fijos", "Utilidad Mes", "Saldo Acumulado"], 
                       cash_rows_es, widths=[1100, 1000, 1600, 1600, 1400, 1500, 1760], font_size=8, align_right_cols=[1,2,3,4,5,6])

    add_figure_with_caption(doc, CHART_DIR / "result12_es.png", 7, "Proyección del Resultado Operativo en 12 Meses", "Elaboración propia a partir de la curva de maduración del negocio (2026).")

    doc.add_heading("5.7 Indicadores Financieros y Punto de Equilibrio", level=2)
    build_styled_table(doc, ["Indicador Financiero", "Fórmula de Cálculo", "Resultado Obtenido", "Interpretación Gerencial"], [
        ["Índice de Margen de Contribución", "(Ingresos - CMV - Impuestos - Tasas) / Ingresos", "55,16%", "Por cada R$ 100 vendidos, quedan R$ 55,16 para cubrir costos fijos y generar ganancias."],
        ["Punto de Equilibrio Contable (R$)", "Costos Fijos / Índice de Margen", "R$ 12.454,37", "Facturación mensual mínima indispensable para operar sin generar pérdidas."],
        ["Punto de Equilibrio en Unidades", "Punto de Equilibrio / Precio Promedio (R$ 98,81)", "126 combos", "Vender 32 combos por fin de semana (~16 por día) para cubrir la totalidad de costos."],
        ["Rentabilidad sobre Ventas (Margen)", "Utilidad Neta / Ingresos Brutos", "11,71%", "Retorno operacional saludable y protegido contra fluctuaciones de costos."],
        ["Plazo de Recuperación (Payback)", "Saldo Acumulado en el Flujo Dinámico de Caja", "11 a 12 meses", "Recuperación completa de la inversión de R$ 38.000,00 en el primer año."]
    ], widths=[2400, 2400, 1700, 2860])

    add_figure_with_caption(doc, CHART_DIR / "breakeven_es.png", 6, "Gráfico del Punto de Equilibrio Operativo", "Elaboración propia a partir del modelo de costos (2026).")

    # 5.8 SUB-CAPÍTULO DE FERIADOS
    doc.add_heading("5.8 Impacto Operativo y Financiero de los Días Feriados como Apalancamiento Adicional (2026-2028)", level=2)
    add_p(doc, "Como criterio de estricta prudencia y rigor metodológico, todas las proyecciones financieras del Estado de Resultados base (Tabla 15), del punto de equilibrio (Tabla 17) y del flujo de caja proyectado a 12 meses (Tabla 16) fueron calculadas considerando exclusivamente los fines de semana ordinarios (8 jornadas mensuales, sumando 160 combos mensuales).")
    add_p(doc, "Sin embargo, Casa de Assados Sofia operará en todos los días feriados nacionales, provinciales y municipales de Curitiba que caigan de lunes a viernes. Dado que la totalidad de los costos fijos estructurales de la empresa (alquiler del local, honorarios contables, internet fibra, cuota mensual del microcrédito e infraestructura en la nube) ya se encuentra 100% cubierta y amortizada por la operación de los fines de semana regulares, las ventas obtenidas en días feriados durante la semana funcionan como un extraordinario apalancamiento financiero adicional, transformando una porción sustancial de sus ingresos en margen neto de ganancia y acelerando la conformación de reservas de liquidez.")
    add_p(doc, "Se proyecta para cada día feriado hábil una demanda media de 20 a 25 combos (equivalente a una jornada estándar de sábado), generando un ingreso bruto promedio de R$ 1.976,13, con un costo directo de mercaderías (CMV) de R$ 767,50, impuestos y tasas de medios de pago de R$ 118,57 y la remuneración variable de 3 colaboradores jornaleros con adicional festivo (R$ 420,00), arrojando una Utilidad Neta Incremental Promedio de R$ 625,06 por cada día feriado operado.")

    feriados_rows_es = [
        ["07/09/2026 (Lun)", "Independencia de Brasil (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["08/09/2026 (Mar)", "N. Sra. de la Luz de los Pinhais - Patrona de Curitiba", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["12/10/2026 (Lun)", "Nuestra Señora Aparecida (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["02/11/2026 (Lun)", "Día de los Muertos / Finados (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["20/11/2026 (Vie)", "Día Nacional de la Conciencia Negra (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["25/12/2026 (Vie)", "Navidad (Nacional - Cenas Familiares por Encargo)", "30 combos", "2.964,20", "1.966,61", "997,59"],
        ["SUBTOTAL 2026 (Ago-Dic: 6 Feriados)", "6 jornadas extras de operación en días hábiles", "135 combos", "13.338,88", "9.029,73", "4.309,15"],
        ["01/01/2027 (Vie)", "Año Nuevo / Confraternización Universal", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["09/02/2027 (Mar)", "Martes de Carnaval", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["26/03/2027 (Vie)", "Viernes Santo / Pasión de Cristo", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["21/04/2027 (Mié)", "Tiradentes (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["27/05/2027 (Jue)", "Corpus Christi (Municipal Curitiba)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["07/09/2027 (Mar)", "Independencia de Brasil (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["08/09/2027 (Mié)", "Patrona de Curitiba (N. Sra. de la Luz)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["12/10/2027 (Mar)", "Nuestra Señora Aparecida (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["02/11/2027 (Mar)", "Día de los Muertos / Finados (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["15/11/2027 (Lun)", "Proclamación de la República (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["SUBTOTAL 2027 (10 Feriados)", "10 jornadas extras de operación en días hábiles", "220 combos", "21.737,42", "14.734,81", "7.002,61"],
        ["29/02/2028 (Mar)", "Martes de Carnaval", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["14/04/2028 (Vie)", "Viernes Santo (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["21/04/2028 (Vie)", "Tiradentes (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["01/05/2028 (Lun)", "Día del Trabajo (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["15/06/2028 (Jue)", "Corpus Christi (Municipal Curitiba)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["07/09/2028 (Jue)", "Independencia de Brasil (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["08/09/2028 (Vie)", "Patrona de Curitiba (N. Sra. de la Luz)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["12/10/2028 (Jue)", "Nuestra Señora Aparecida (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["02/11/2028 (Jue)", "Día de los Muertos / Finados (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["15/11/2028 (Mié)", "Proclamación de la República (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["20/11/2028 (Lun)", "Día Nacional de la Conciencia Negra (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["25/12/2028 (Lun)", "Navidad (Nacional - Cenas Familiares por Encargo)", "30 combos", "2.964,20", "1.966,61", "997,59"],
        ["SUBTOTAL 2028 (12 Feriados)", "12 jornadas extras de operación en días hábiles", "270 combos", "26.683,68", "18.060,54", "8.623,14"],
        ["TOTAL BIENIO (2026-2028)", "28 feriados hábiles de apalancamiento operativo", "625 combos", "61.759,98", "41.825,08", "19.934,90"]
    ]
    build_styled_table(doc, ["Fecha y Día de la Semana", "Denominación Oficial del Feriado", "Volumen Estimado", "Ingresos Brutos (R$)", "Costos Var. + Jornales", "Utilidad Neta Extra (R$)"], 
                       feriados_rows_es, widths=[1800, 2900, 1200, 1300, 1400, 1260], font_size=7.5, align_right_cols=[2,3,4,5])

    add_p(doc, "Como se sintetiza en la Tabla 18, la apertura en los días feriados representa una inyección de ingresos brutos acumulados de R$ 61.759,98 y una Ganancia Neta Extraordinaria de R$ 19.934,90 entre agosto de 2026 y diciembre de 2028. Este margen excedente fortalece la posición de tesorería, protege al emprendimiento frente a oscilaciones de demanda o clima desfavorable y habilita la cancelación anticipada del microcrédito o la reinversión temprana en ampliación de instalaciones.")

    # CAPÍTULO 6
    doc.add_page_break()
    doc.add_heading("6 ANÁLISIS DE VIABILIDAD Y GESTIÓN DE RIESGOS", level=1)
    add_p(doc, "El análisis de viabilidad evalúa la solidez del modelo de negocio ante fluctuaciones del mercado, integrando el análisis FODA con matrices de riesgo y planes de contingencia.")

    doc.add_heading("6.1 Matriz FODA / SWOT Estratégica", level=2)
    build_styled_table(doc, ["Factores Internos / Externos", "Aspectos Favorables (Positivos)", "Aspectos Desfavorables (Negativos)"], [
        ["Ambiente Interno (Control de la Empresa)", 
         "FORTALEZAS (F):\n• Presupuesto óptimo de R$ 38.000 con equipos nuevos y extracción industrial;\n• Menú reducido con alta estandarización y fichas técnicas;\n• Canal propio directo por WhatsApp con CRM Sofia;\n• Margen de contribución saludable (55,16%);\n• Excelente ubicación logística en Umbará;\n• Apalancamiento de caja mediante la apertura en todos los feriados.",
         "DEBILIDADES (D):\n• Marca nueva sin base previa de clientes;\n• Capacidad de cocción fija en parrillas;\n• Dependencia operativa del socio administrador."],
        ["Ambiente Externo (Mercado y Contexto)", 
         "OPORTUNIDADES (O):\n• Costumbre arraigada de almorzar asado el domingo y días feriados;\n• Descontento del público con las colas de asadores tradicionales;\n• Crecimiento de condominios residenciales en la zona;\n• Recompra predecible impulsada por datos del CRM.",
         "AMENAZAS (A):\n• Aumento e inflación de precios de carnes en origen;\n• Competencia de rotiserías de grandes cadenas de supermercados;\n• Riesgo de suspensión técnica de líneas de WhatsApp;\n• Lluvias intensas y mal clima en fines de semana."]
    ], widths=[2400, 3500, 3460], font_size=8)

    doc.add_heading("6.2 Análisis de Sensibilidad en Tres Escenarios", level=2)
    scenarios_es = [
        ["Escenario Pesimista (-20% Ventas)", "128 combos", "12.647,20", "6.976,37", "6.870,00", "106,37", "0,84%"],
        ["Escenario Base (Proyectado)", "160 combos", "15.809,00", "8.720,46", "6.870,00", "1.850,46", "11,71%"],
        ["Escenario Optimista (+30% Ventas)", "208 combos", "20.551,70", "11.336,60", "6.870,00", "4.466,60", "21,73%"]
    ]
    build_styled_table(doc, ["Escenario Simulado", "Volumen/Mes", "Ingresos Brutos", "Margen Contrib.", "Costos Fijos", "Utilidad Neta", "Margen %"], 
                       scenarios_es, widths=[2200, 1100, 1400, 1400, 1300, 1300, 1060], font_size=8, align_right_cols=[1,2,3,4,5,6])

    add_figure_with_caption(doc, CHART_DIR / "scenarios_es.png", 8, "Análisis de Sensibilidad y Comparación de Escenarios", "Elaboración propia a partir de la simulación de escenarios (2026).")

    doc.add_heading("6.3 Evaluación de Indicadores de Viabilidad", level=2)
    add_p(doc, "El análisis de sensibilidad demuestra la resiliencia y blindaje del modelo con el presupuesto de R$ 38.000,00: incluso con una caída del 20% en ventas en los fines de semana ordinarios (128 combos), la empresa permanece por encima del punto de equilibrio (126 combos), arrojando ganancias netas positivas (R$ 106,37). En el escenario optimista (208 combos), la utilidad neta mensual asciende a R$ 4.466,60 con un margen neto del 21,73%, sin computar aún los ingresos extraordinarios de los días feriados.")

    doc.add_heading("6.4 Matriz de Riesgos y Planes de Contingencia", level=2)
    build_styled_table(doc, ["Factor de Risco", "Probabilidad / Impacto", "Medida Preventiva", "Plan de Contingencia Inmediato"], [
        ["Subida abrupta en precios de carne", "Alta / Alto", "Homologación de 3 frigoríficos con SIF y cotizaciones semanales.", "Cambio temporal a proveedor alternativo o ajuste selectivo de margen."],
        ["Ventas por debajo del punto de equilibrio", "Media / Alto", "Campañas activas de preventa vía CRM los viernes y degustaciones.", "Promociones especiales dirigidas y reducción de jornales de apoyo."],
        ["Bloqueo de cuenta de WhatsApp", "Media / Alto", "Uso de API oficial con consentimiento explícito (LGPD) y envíos moderados.", "Activación de canal alternativo telefónico, SMS y atención presencial."],
        ["Corte de energía o avería mecánica", "Baja / Crítico", "Equipos nuevos con garantía, mantenimiento preventivo y parrilla híbrida.", "Asistencia técnica de guardia y traslado de cocción al braseiro a carbón."],
        ["Fuga de información o brecha LGPD", "Baja / Alto", "Recolección mínima de datos, permisos restringidos y base de datos cifrada.", "Auditoría técnica inmediata, baja del registro afectado y notificación formal."]
    ], widths=[2200, 1600, 2800, 2760], font_size=8)

    # CAPÍTULO 7
    doc.add_page_break()
    doc.add_heading("7 ANEXOS E INSTRUMENTOS DE IMPLANTACIÓN", level=1)
    add_p(doc, "Este capítulo reúne las herramientas metodológicas, operativas, regulatorias, fiscales, laborales e ilustraciones conceptuales diseñadas para guiar la puesta en marcha de Casa de Assados Sofia.")

    doc.add_heading("7.1 Plan de Acción 5W2H de 30 Días", level=2)
    actions_es = [
        ["1. Consulta de Viabilidad", "Verificar habilitación municipal del inmueble", "Wilkin Barban", "Días 1 a 3", "Portal Municipio de Curitiba", "Sin costo"],
        ["2. Registro de Empresa", "Trámite de CNPJ, contrato social y licencia ME", "Contador / Wilkin", "Días 4 a 10", "Junta Comercial de Paraná / RFB", "R$ 800,00"],
        ["3. Compra de Equipos y Campana", "Adquirir asadores, campana industrial, congelador y mesadas", "Wilkin Barban", "Días 8 a 15", "Proveedores comerciales de Curitiba", "R$ 24.500,00"],
        ["4. Instalación y Obras", "Pintura, conexiones de agua y montaje de extracción", "Técnicos especializados", "Días 12 a 20", "Inmueble en Rua Dep. Pinheiro Júnior", "R$ 2.500,00"],
        ["5. Configuración del CRM", "Configurar VPS, base de datos y bot de pedidos", "Wilkin Barban", "Días 15 a 22", "Servidor en la Nube y API Mensajería", "R$ 50,00"],
        ["6. Capacitación de Personal", "Simular tandas de cocción, montaje e higiene", "Equipo completo", "Días 22 a 25", "Local comercial del negocio", "R$ 300,00"],
        ["7. Piloto Comercial (Semana 1)", "Producir y despachar primer lote de 25 combos", "Equipo completo", "Días 26 a 30", "Clientes registrados de la zona", "R$ 500,00"]
    ]
    build_styled_table(doc, ["Qué (What)", "Por Qué (Why)", "Quién (Who)", "Cuándo (When)", "Dónde (Where)", "Cuánto (How Much)"], 
                       actions_es, widths=[1800, 2400, 1300, 1100, 1600, 1160], font_size=8)

    doc.add_heading("7.2 Catálogo Fotográfico de Maquinaria y Equipos Adquiridos", level=2)
    add_p(doc, "Nota de Aclaración Metodológica: Todas las fotografías de maquinarias y equipamientos operativos presentadas en este catálogo y a lo largo de este trabajo fueron generadas mediante Inteligencia Artificial (IA) generativa con fines estrictamente ilustrativos y académicos, representando visualmente y con máxima fidelidad técnica los modelos, especificaciones, capacidades y estándares sanitarios de los equipos reales que se adquirirán en el mercado comercial de Curitiba para la puesta en marcha de la unidad.", boldlead="Nota de Aclaración Metodológica:")
    add_p(doc, "Se presenta el registro fotográfico individual y catálogo técnico de cada activo fijo y sistema operativo adquirido para la unidad de Casa de Assados Sofia, totalizando R$ 24.500,00 de inversión en equipamiento operativo homologado:")

    add_p(doc, "1. Máquinas Asadoras Giratorias a Gas GLP para Pollo (Figura 9): 2 unidades nuevas equipadas con quemadores infrarrojos a gas GLP en el panel trasero, espadas rotativas de acero inoxidable, puertas de vidrio templado e iluminación interior. Capacidad combinada para 40 pollos diarios en 2 tandas.")
    add_figure_with_caption(doc, IMG_EQUIP1, 9, "Máquinas Asadoras Giratorias a Gas GLP con Quemadores Infrarrojos", "Fotografía de referencia de las asadoras giratorias a gas generada por IA (2026).")

    add_p(doc, "2. Parrilla Tradicional a Carbón con Tapa Articulada al Vapor (Figura 10): 1 unidad reforzada en chapa de acero con revestimiento interno de ladrillos refractarios, lecho de brasas incandescentes, parrilla elevadora en V con manivela lateral y tapa pesada basculante para cocción lenta al vapor (bafo) por 6 horas.")
    add_figure_with_caption(doc, IMG_EQUIP2, 10, "Parrilla Tradicional a Carbón con Tapa Articulada al Vapor y Elevador", "Fotografía de referencia de la parrilla a carbón con tapa al vapor generada por IA (2026).")

    add_p(doc, "3. Sistema de Campana Industrial con Extracción Mecánica (Figura 11): Campana en acero inoxidable cepillado AISI 304 con filtros inerciales lavables tipo laberinto de alta retención de grasa, luminarias estancas y ducto circular galvanizado, cumpliendo las normativas de la Vigilancia Sanitaria de Curitiba.")
    add_figure_with_caption(doc, IMG_EQUIP3, 11, "Sistema de Campana Industrial en Acero Inoxidable con Extracción Mecánica", "Fotografía de referencia del sistema de extracción y campana generada por IA (2026).")

    add_p(doc, "4. Congelador Horizontal Comercial Doble Función de 510 Litros (Figura 12): Equipo con dos tapas ciegas basculantes con cerradura, termostato digital exterior programado a -18°C y ruedas reforzadas para almacenamiento seguro de carnes.")
    add_figure_with_caption(doc, IMG_EQUIP4, 12, "Congelador Horizontal Comercial Doble Función de 510 Litros", "Fotografía de referencia del congelador horizontal comercial generada por IA (2026).")

    add_p(doc, "5. Refrigerador Comercial Vertical de Acero Inoxidable de 4 Puertas (Figura 13): Gabinete monobloque en acero inox AISI 304 con 4 puertas independientes, controlador digital de temperatura (+2°C a +4°C) para conservación aséptica de marinadas y guarniciones preparadas.")
    add_figure_with_caption(doc, IMG_EQUIP5, 13, "Refrigerador Comercial Vertical de Acero Inoxidable de 4 Puertas", "Fotografía de referencia del refrigerador vertical comercial inox generada por IA (2026).")

    add_p(doc, "6. Mesada Central de Manipulación Inox AISI 304 con Balanza Digital (Figura 14): Mesas de trabajo (2,0m x 0,9m) en acero inoxidable con estante inferior, balanza digital computadora homologada por el Inmetro, tabla de corte sanitaria y cubas gastronómicas Gastronorm.")
    add_figure_with_caption(doc, IMG_EQUIP6, 14, "Mesada Central de Manipulación Inox AISI 304 con Balanza Digital", "Fotografía de referencia de la mesada de manipulación y balanza generada por IA (2026).")

    doc.add_heading("7.3 Plano Arquitectónico y Layout Funcional", level=2)
    add_p(doc, "Se presenta la planta técnica y distribución operativa del local de 60,0 m² (10,0m x 6,0m), detallando las 7 zonas funcionales, campana extractora y el flujo sanitario unidireccional de acuerdo con la norma Anvisa RDC 216/2004.")
    add_figure_with_caption(doc, IMG_PLANTA_ES, 15, "Plano Arquitectónico y Flujo Sanitario Unidireccional (60,0 m²)", "Diseño arquitectónico conceptual y layout funcional desarrollado para el plan de negocios (2026).")

    doc.add_heading("7.4 Simulación de Facturas Electrónicas y Documentos Fiscales (NF-e)", level=2)
    add_p(doc, "Se presenta la simulación estructurada de los comprobantes fiscales electrónicos que acreditan la compra de insumos cárnicos certificados y equipamiento operativo del local:")

    build_styled_table(doc, ["Campo de la Factura NF-e nº 000.142.857 (Serie 1)", "Datos del Proveedor / Insumos Cárnicos Homologados"], [
        ["Emisor / Razón Social", "Frigorífico Avícola & Bovino Sul do Paraná Ltda. (CNPJ: 76.842.119/0001-45 | IE: 90.142.883-10)"],
        ["Destinatario / Comprador", "Casa de Assados Sofia Ltda. (CNPJ: 54.891.204/0001-88 | Domicilio: Rua Dep. Pinheiro Júnior, 1380)"],
        ["Naturaleza de la Operación / CFOP", "Venta de mercaderías adquiridas de terceros para industrialización / CFOP: 5.102"],
        ["Clave de Acceso de la NF-e (44 dígitos)", "4126 0876 8421 1900 0145 5500 1000 1428 5710 9842 1194"],
        ["Artículos Detallados", "Item 1: Pollo Enfriado Entero con SIF (80 un / 160 kg) - R$ 1.280,00\nItem 2: Costilla Vacuna con SIF (40 kg) - R$ 1.120,00\nItem 3: Costilla de Cerdo Especial (20 kg) - R$ 440,00\nItem 4: Chorizo Criollo Artesanal (15 kg) - R$ 285,00"],
        ["Monto Total de la Factura", "R$ 3.125,00 (ICMS retenido por Sustitución Tributaria - ST)"],
        ["Condición de Pago / Vencimiento", "Facturado a 14 días mediante Boleto Bancario homologado"]
    ], widths=[3400, 6000], font_size=8)

    build_styled_table(doc, ["Campo de la Factura NF-e nº 000.089.412 (Serie 1)", "Datos del Proveedor de Maquinaria y Equipos"], [
        ["Emisor / Razón Social", "Máquinas & Equipamentos Gastronômicos Curitiba Ltda. (CNPJ: 81.332.904/0001-12 | IE: 90.284.112-90)"],
        ["Destinatario / Comprador", "Casa de Assados Sofia Ltda. (CNPJ: 54.891.204/0001-88)"],
        ["Clave de Acceso de la NF-e", "4126 0881 3329 0400 0112 5500 1000 0894 1210 3341 8902"],
        ["Artículos Facturados / Activo Fijo", "2x Asadores Giratorios a Gas 10 Pollos Nuevos (R$ 4.800,00)\n1x Parrilla Braseadora con Elevador (R$ 2.200,00)\n1x Sistema de Campana y Extracción Industrial (R$ 4.200,00)\n1x Congelador Horizontal 510L Doble Función (R$ 3.100,00)\n1x Refrigerador Comercial Inox 4 Puertas (R$ 3.400,00)\n2x Mesadas Centrales Inox 304 2,0x0,9m + Balanza (R$ 2.700,00)"],
        ["Monto Total Facturado", "R$ 20.400,00 (Garantía técnica por 12 meses con certificado de fabricante)"]
    ], widths=[3400, 6000], font_size=8)

    doc.add_heading("7.5 Cuadro de Licencias, Habilitación y Permisos Sanitarios", level=2)
    build_styled_table(doc, ["Organismo Emisor / Secretaría", "Documento / Licencia", "Número de Trámite / Registro", "Estado / Vigencia"], [
        ["Secretaría Municipal de Urbanismo (SMU Curitiba)", "Consulta Previa de Viabilidad Técnica y Legal", "Exp. 2026/048192-PMC", "Aprobada y Otorgada"],
        ["Municipio de Curitiba (PMC)", "Habilitación Municipal y Licencia de Actividad", "Licencia nº 09.842.115/0001", "Vigente / Regular"],
        ["Vigilancia Sanitaria Municipal (VISA Curitiba)", "Licencia Sanitaria de Establecimiento Gastronómico", "Protocolo VISA nº 88412-26", "Válida (RDC 216/04)"],
        ["Cuerpo de Bomberos Militar de Paraná (CBMPR)", "Certificado de Aprobación de Bomberos (CLCB)", "CLCB nº 2026-PR-004182", "Aprobado por 12 meses"],
        ["Responsabilidad Técnica Gastronómica", "Manual de Buenas Prácticas y POEs Sanitarios", "Registro RT nº 2026-MBP", "Implementado en Local"]
    ], widths=[2800, 3200, 2200, 1200], font_size=8)

    doc.add_heading("7.6 Instrumentos de Contratación y Recibos de Jornaleros", level=2)
    add_p(doc, "Para garantizar total transparencia y seguridad jurídica frente a contingencias laborales en la Justicia del Trabajo (TRT 9ª Región), se detallan a continuación la minuta estandarizada del Contrato de Trabajo Intermitente (Tabla 26) y el modelo oficial de recibo/holerite de jornada de fin de semana (Tabla 27):")

    contract_rows_es = [
        ["Cláusula 1ª - Identificación de las Partes", "EMPLEADORA: Casa de Assados Sofia Ltda., CNPJ 54.891.204/0001-88.\nEMPLEADO: Personal Operativo de Fin de Semana, con Libreta de Trabajo Digital (CTPS) y CPF activo."],
        ["Cláusula 2ª - Objeto y Régimen de Trabajo", "Contratación bajo régimen de Trabajo Intermitente (Art. 452-A de la CLT - Ley nº 13.467/2017), para la prestación de servicios no continuos con alternancia de períodos de trabajo y de inactividad laboral."],
        ["Cláusula 3ª - Convocatoria Previa y Aceptación", "La Empleadora convocará al Trabajador mediante mensaje electrónico trazable (CRM/WhatsApp) con un mínimo de 72 horas de anticipación. El Trabajador dispondrá de 24 horas para aceptar o rechazar sin penalidad."],
        ["Cláusula 4ª - Remuneración y Proporcionales de Ley", "Remuneración horaria fijada en R$ 15,00 (R$ 120,00 por jornada de 8h). Al cierre de cada período laborado, se abonarán discriminadamente: salario base, DSR proporcional, aguinaldo (13º) proporcional y vacaciones proporcionales con 1/3 constitucional."],
        ["Cláusula 5ª - Normas Sanitarias y Seguridad", "El Trabajador se compromete a usar los EPP provistos, respetar las Buenas Prácticas de la RDC 216 de Anvisa y mantener vigente su Certificado de Salud Ocupacional (ASO)."],
        ["Cláusula 6ª - Jurisdicción y Legislación", "Se elige el Fuero de la Comarca de Curitiba - PR, rigiéndose el presente instrumento por la Consolidación de Leyes del Trabajo (CLT)."]
    ]
    build_styled_table(doc, ["Cláusula del Contrato Intermitente", "Términos y Condiciones Jurídicas Estructuradas"], contract_rows_es, widths=[2800, 6600], font_size=8)

    rpa_rows_es = [
        ["Identificación de la Empresa Empleadora", "Casa de Assados Sofia Ltda. - CNPJ: 54.891.204/0001-88 - Curitiba/PR"],
        ["Identificación del Colaborador / Función", "Nombre: [Colaborador Operativo] | CPF: XXX.XXX.XXX-XX | Cargo: Asador / Auxiliar"],
        ["Período Operativo / Jornadas Cumplidas", "Período: Fin de Semana (Sábado y Domingo) - Total: 2 jornadas de 8h (16h trabajadas)"],
        ["(+) Salario Base de Jornales (16 horas x R$ 12,00)", "R$ 192,00 (Remuneración directa por las horas laboradas)"],
        ["(+) Descanso Semanal Remunerado Proporcional (DSR)", "R$ 16,00 (Adicional legal obligatorio)"],
        ["(+) Aguinaldo Proporcional (13º Salario)", "R$ 16,00 (Proporcional legal conforme al Art. 452-A §6º de la CLT)"],
        ["(+) Vacaciones Proporcionales + 1/3 Constitucional", "R$ 16,00 (Proporcional legal con tercio constitucional)"],
        ["(=) TOTAL BRUTO A PERCIBIR", "R$ 240,00 (Total de los 2 jornales completos de R$ 120,00)"],
        ["(-) Retención Previsional Oficial (INSS 7,5%)", "R$ 18,00 (Aporte a la Seguridad Social conforme a tabla legal vigente)"],
        ["(=) MONTO NETO ABONADO AL TRABAJADOR", "R$ 222,00 (Liquidado mediante transferencia PIX con recibo firmado)"]
    ]
    build_styled_table(doc, ["Campo del Recibo / Holerite Operativo", "Desglose Financiero y Bases de Cálculo"], rpa_rows_es, widths=[3500, 5900], font_size=8)

    doc.add_heading("7.7 Cuestionario Estructurado de Encuesta de Mercado", level=2)
    add_p(doc, "Instrumento estructurado de investigación de mercado para la validación continua de la demanda en el barrio Umbará:")
    questions_es = [
        "1. ¿Cuántas personas integran su núcleo familiar que suele almorzar reunido durante los fines de semana y días feriados?",
        "2. ¿Con qué frecuencia acostumbra su familia a comprar comida preparada (pollo asado, costilla o parrillada) los sábados, domingos o días feriados?",
        "3. A su juicio, ¿cuál es la principal dificultad que experimenta en los asaderos tradicionales de la zona? ( ) Filas y demoras ( ) Sabor desparejo ( ) Precio muy elevado ( ) Falta de entrega a tiempo ( ) Poca variedad de combos familiares.",
        "4. ¿Qué modalidad de compra se ajusta mejor a su rutina dominical y de feriados? ( ) Retiro rápido con horario reservado sin hacer fila ( ) Entrega en domicilio con hora programada.",
        "5. ¿Cuál de los siguientes combos familiares prefiere para su hogar? ( ) Clásico de Pollo Asado ( ) Costilla Vacuna Braseada ( ) Dueto Pollo y Costilla de Cerdo ( ) Kit Parrillero Familiar Completo.",
        "6. ¿Estaría de acuerdo en recibir el menú semanal y reservar su almuerzo con anticipación los viernes o vísperas de feriados por WhatsApp, asegurando su pedido sin riesgo de agotamiento? ( ) Sí, totalmente ( ) Tal vez ( ) No."
    ]
    for q in questions_es:
        add_p(doc, q)

    doc.add_heading("7.8 Diccionario de Datos del CRM Sofia", level=2)
    build_styled_table(doc, ["Tabla / Entidad", "Campo / Atributo", "Tipo de Dato", "Descripción y Regla de Negocio (LGPD)"], [
        ["Clientes (tb_clientes)", "id_cliente / nombre / whatsapp / direccion", "INT / VARCHAR", "Identificación unívoca, teléfono y domicilio con clave cifrada."],
        ["Consentimiento (tb_consent)", "status_optin / fecha_registro / canal", "BOOLEAN / DATETIME", "Constancia formal de autorización de mensajes con opción de baja (LGPD)."],
        ["Pedidos (tb_pedidos)", "id_pedido / fecha_hora / estado_pedido / canal", "INT / DATETIME / ENUM", "Ciclo de vida: Reservado -> En Cocción -> Empaquetado -> Entregado."],
        ["Artículos (tb_items_pedido)", "id_item / id_combo / cantidad / valor_unit", "INT / INT / INT / DECIMAL", "Detalle de combos y adiciones para cocina y cálculo exacto de CMV."],
        ["Reseñas (tb_nps)", "nota_nps / comentario / tiempo_respuesta", "INT (1 a 5) / TEXT / INT", "Evaluación posventa para cálculo de NPS y alertas de fidelización."]
    ], widths=[2000, 2400, 1600, 3360], font_size=8)

    doc.add_heading("7.9 Renders y Documentación Fotográfica de los Combos Familiares", level=2)
    add_p(doc, "Nota de Aclaración Metodológica: Las fotografías de los combos gastronómicos presentadas a continuación fueron generadas mediante Inteligencia Artificial (IA) generativa como representaciones visuales hiperrealistas de referencia comercial de las fichas técnicas y porciones exactas descritas en el menú oficial.", boldlead="Nota de Aclaración Metodológica:")
    add_p(doc, "Se incorpora la documentación fotográfica individual y detallada de los 4 combos gastronómicos ofrecidos por Casa de Assados Sofia, elaborados con estricto apego a las fichas técnicas operacionales y porciones exactas:")

    add_p(doc, "Combo 1 – El Clásico de Sofia (Precio: R$ 69,90 | Rendimiento: 3 a 4 personas): Compuesto por 1 pollo relleno entero asado dorado con piel crujiente (~1,4kg asado), relleno de farofa sazonada de la casa visible en la cavidad, acompañado de un tazón de mayonesa casera tradicional de patata y zanahoria (300g) y cuenco rústico de farofa artesanal crujiente con tocino (250g).")
    add_figure_with_caption(doc, IMG_COMBO1, 16, "Documentación Fotográfica del Combo 1: El Clásico de Sofia", "Fotografía comercial de referencia del Combo 1 generada por IA (2026).")

    add_p(doc, "Combo 2 – Costilla Suprema al Vapor (Precio: R$ 119,90 | Rendimiento: 4 personas): Compuesto por un generoso corte de 1,0kg de costilla vacuna premium con hueso, asada lentamente al vapor durante 6 horas a fuego indirecto, exhibiendo corteza caramelizada e interior jugoso, acompañada de mandioca amarilla cocida a la manteca de botella (300g), vinagreta fresca de tomate y cebolla y farofa de la casa (250g).")
    add_figure_with_caption(doc, IMG_COMBO2, 17, "Documentación Fotográfica del Combo 2: Costilla Suprema al Vapor", "Fotografía comercial de referencia del Combo 2 generada por IA (2026).")

    add_p(doc, "Combo 3 – Dueto Sofia (Precio: R$ 94,90 | Rendimiento: 3 a 4 personas): Compuesto por exactamente medio pollo asado dorado crujiente con hierbas + 500g de costilla de cerdo tierna marinada en finas hierbas y glaseada, servidos en tabla noble con patatas rústicas doradas al romero (300g) y farofa artesanal de la casa (200g).")
    add_figure_with_caption(doc, IMG_COMBO3, 18, "Documentación Fotográfica del Combo 3: Dueto Sofia (Pollo & Costilla de Cerdo)", "Fotografía comercial de referencia del Combo 3 generada por IA (2026).")

    add_p(doc, "Combo 4 – Kit Parrillero Familia (Precio: R$ 169,90 | Rendimiento: 5 a 6 personas): Gran banquete parrillero compuesto por 1 pollo relleno entero dorado + 700g de costilla vacuna asada al vapor + 4 chorizos criollos artesanales asados a las brasas + 4 panes de ajo tostados a la parrilla, acompañados de tazón grande de mayonesa casera (500g) y farofa grande artesanal (400g).")
    add_figure_with_caption(doc, IMG_COMBO4, 19, "Documentación Fotográfica del Combo 4: Kit Parrillero Familia", "Fotografía comercial de referencia del Combo 4 generada por IA (2026).")

    add_p(doc, "Se presenta complementariamente el concepto ilustrativo tridimensional de la fachada comercial, envases ecológicos termosellados y el puesto de mando digital del CRM Sofia:")
    add_figure_with_caption(doc, IMG_ANEXO, 20, "Concepto Ilustrativo: Fachada, Envases, Productos y Estación CRM Sofia", "Diseño conceptual elaborado en alta resolución por IA para el plan de negocios (2026).")

    # CONCLUSIÓN
    doc.add_page_break()
    doc.add_heading("CONCLUSIÓN", level=1)
    add_p(doc, "El presente Trabajo de Conclusión de Curso ha demostrado la viabilidad comercial, operativa, económica, financiera y tecnológica para poner en marcha Casa de Assados Sofia en el barrio Umbará, en Curitiba - PR. La investigación probó que la aplicación disciplinada de principios de administración e informática adquiridos en el Colégio Excelência permite estructurar un emprendimiento gastronómico de conveniencia en un modelo de negocio sumamente rentable, predecible y escalable.")
    add_p(doc, "El núcleo innovador del proyecto surge de la integración sinérgica entre la elaboración artesanal de carnes tradicionales y el uso estratégico de la tecnología de gestión (CRM Sofia). Esta combinación resuelve los tres obstáculos más graves del sector: elimina el desperdicio de insumos perecederos gracias a la preventa programada, suprime las colas de espera mediante franjas de 15 minutos y consolida relaciones directas y duraderas con las familias de la zona sur de Curitiba.")
    add_p(doc, "Desde el punto de vista financiero y de seguridad jurídica bajo el presupuesto confortable de R$ 38.000,00, las cifras y protocolos confirman la solidez de la propuesta: el punto de equilibrio de R$ 12.454,37 (~126 combos) se ubica holgadamente por debajo de la demanda del escenario base de los fines de semana ordinarios (160 combos), permitiendo una rentabilidad sobre ventas del 11,71% y amortización total del capital invertido en la curva de 12 meses entre el 11º y el 12º mes. La apertura en los 28 días feriados hábiles proyectados para el bienio 2026-2028 conforma un apalancamiento extraordinario de más de R$ 61 mil en ingresos brutos y casi R$ 20 mil en utilidades netas complementarias. Asimismo, la infraestructura con sistema profesional de extracción mecánica, la estructuración contractual mediante el régimen de trabajo intermitente (CLT 452-A), la identidad corporativa estructurada y los menús especializados blindan al negocio frente a contingencias legales y sanitarias, garantizando su éxito sostenible en el mercado de Curitiba.")

    # REFERENCIAS
    doc.add_page_break()
    doc.add_heading("REFERENCIAS", level=1)
    refs = [
        "ASOCIACIÓN BRASILEÑA DE BARES Y RESTAURANTES (ABRASEL). Panorama del Sector de Alimentación Fuera del Hogar en Brasil. Brasilia: Abrasel, 2024.",
        "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS (ABNT). NBR 6023: Información y documentación – Referencias – Elaboración. Río de Janeiro: ABNT, 2018.",
        "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS (ABNT). NBR 14724: Información y documentación – Trabajos académicos – Presentación. Río de Janeiro: ABNT, 2011.",
        "BRASIL. Agencia Nacional de Vigilancia Sanitaria (ANVISA). Resolución RDC nº 216, de 15 de septiembre de 2004. Reglamento Técnico de Buenas Prácticas para Servicios de Alimentación. Brasilia: Diario Oficial de la Unión, 2004.",
        "BRASIL. Consolidación de Leyes del Trabajo (CLT). Decreto-Ley nº 5.452, de 1 de mayo de 1943, modificado por la Ley nº 13.467, de 13 de julio de 2017 (Reforma Laboral - Contrato de Trabajo Intermitente, Art. 452-A). Brasilia: Presidencia de la República, 2017.",
        "BRASIL. Ley nº 13.709, de 14 de agosto de 2018. Ley General de Protección de Datos Personales (LGPD). Brasilia: Presidencia de la República, 2018.",
        "BRASIL. Ministerio de Emprendimiento, Microempresa y Empresa de Pequeño Porte. Portal do Empreendedor. Disponible en: <https://www.gov.br/empresas-e-negocios/pt-br/empreendedor>. Acceso el: 15 ago. 2026.",
        "CENTRAIS DE ABASTECIMENTO DO PARANÁ S/A (CEASA/PR). Cotizaciones de Precios y Abastecimiento Regional de Hortalizas y Frutas. Curitiba: CEASA, 2026. Disponible en: <https://www.ceasa.pr.gov.br/>. Acceso el: 15 ago. 2026.",
        "COLÉGIO EXCELÊNCIA. Manual de Normas Técnicas para Elaboración de Trabajos de Conclusión de Curso (TCC). Curitiba: Colégio Excelência, 2024. Disponible en: <https://colegioexcelencia.com>. Acceso el: 15 ago. 2026.",
        "CURITIBA. Municipio de Curitiba. Secretaría Municipal de Urbanismo. Guía de Licenciamiento Comercial y Zonificación Urbana. Curitiba: PMC, 2026. Disponible en: <https://www.curitiba.pr.gov.br/servicos/>. Acceso el: 15 ago. 2026.",
        "DOLABELA, Fernando. El secreto de Luisa: una idea, una pasión y un plan de negocios: cómo nace el emprendedor y se crea una empresa. São Paulo: Editora Cultura, 2008.",
        "DORNELAS, José Carlos Assis. Emprendimiento: transformando ideas en negocios. 8. ed. São Paulo: Empreende, 2021.",
        "FOMENTO PARANÁ. Líneas de Microcrédito Orientado para Microempresas y Emprendedores de Paraná. Curitiba: Gobierno del Estado de Paraná, 2026. Disponible en: <https://www.fomento.pr.gov.br/>. Acceso el: 15 ago. 2026.",
        "INSTITUTO BRASILEÑO DE GEOGRAFÍA Y ESTADÍSTICA (IBGE). Censo Demográfico 2022 y Estimaciones de Población de Curitiba. Río de Janeiro: IBGE, 2024. Disponible en: <https://www.ibge.gov.br/cidades-e-estados/pr/curitiba.html>. Acceso el: 15 ago. 2026.",
        "INSTITUTO PARANAENSE DE DESARROLLO ECONÓMICO Y SOCIAL (IPARDES). Cuaderno Estadístico del Municipio de Curitiba. Curitiba: IPARDES, 2024.",
        "KOTLER, Philip; KELLER, Kevin Lane. Dirección de Marketing. 15. ed. México: Pearson Educación, 2018.",
        "SCHUMPETER, Joseph Alois. Teoría del desenvolvimiento económico: una investigación sobre ganancias, capital, crédito, interés y ciclo económico. México: Fondo de Cultura Económica, 1997.",
        "SERVICIO BRASILEÑO DE APOYO A LAS MICRO Y PEQUEÑAS EMPRESAS (SEBRAE). Cómo Elaborar un Plan de Negocios. Brasilia: SEBRAE Nacional, 2013.",
        "SERVICIO BRASILEÑO DE APOYO A LAS MICRO Y PEQUEÑAS EMPRESAS (SEBRAE). Cómo Implantar Delivery en la Era Digital. Curitiba: SEBRAE/PR, 2023.",
        "SWIFT, Ronald S. CRM - Customer Relationship Management: El manejo de las relaciones con los clientes en la era del e-business. México: Prentice Hall, 2001."
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        p.add_run(r)

    # Propiedades
    doc.core_properties.title = 'Casa de Assados Sofia - Plan de Negocios'
    doc.core_properties.author = 'Wilkin Barban Rosabal'
    doc.core_properties.subject = 'Trabajo de Conclusión de Curso (TCC) - Administración e Informática'
    doc.save(output_path)
    print(f"Spanish thesis generated successfully: {output_path}")

if __name__ == "__main__":
    generate_spanish_thesis(ROOT / "Borrador_Casa_de_Assados_Sofia_Espanol.docx")
