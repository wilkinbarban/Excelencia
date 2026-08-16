import os
import shutil
import math
from pathlib import Path
import docx
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT = Path(r"c:\Users\wilki\OneDrive\Documentos\Trabajo de Curso")
CHART_DIR = ROOT / "_work" / "charts"
IMG_ANEXO = ROOT / "_work" / "anexo_casa_assados_sofia.png"

mix_pt = [
    ("O Clássico da Sofia", 70, 69.90, 26.50, "1 Frango recheado inteiro (~1,4kg assado), farofa artesanal crocante (250g), maionese caseira de batata (300g). Serve 3-4 pessoas."),
    ("Costela Suprema", 35, 119.90, 48.00, "1kg de Costela bovina premium assada lentamente no bafo por 6 horas, mandioca na manteiga (300g), vinagrete e farofa (250g). Serve 4 pessoas."),
    ("Dueto Sofia", 35, 94.90, 36.00, "Meio frango assado + 500g de Costelinha de porco marinada em ervas, batatas rústicas (300g) e farofa da casa (200g). Serve 3-4 pessoas."),
    ("Kit Churrasco Família", 20, 169.90, 68.00, "1 Frango recheado + 700g de Costela no bafo + 4 Linguiças toscanas, maionese grande (500g), farofa grande (400g) e 4 pães de alho. Serve 5-6 pessoas."),
]

mix_es = [
    ("El Clásico de Sofia", 70, 69.90, 26.50, "1 Pollo relleno entero (~1,4kg asado), farofa artesanal crocante (250g), mayonesa casera de patata (300g). Rinde 3-4 personas."),
    ("Costilla Suprema", 35, 119.90, 48.00, "1kg de Costilla vacuna premium asada lentamente al vapor por 6 horas, mandioca a la manteca (300g), vinagreta y farofa (250g). Rinde 4 personas."),
    ("Dueto Sofia", 35, 94.90, 36.00, "Medio pollo asado + 500g de Costilla de cerdo marinada en hierbas, patatas rústicas (300g) y farofa de la casa (200g). Rinde 3-4 personas."),
    ("Kit Parrillero Familia", 20, 169.90, 68.00, "1 Pollo relleno + 700g de Costilla vacuna braseada + 4 Chorizos criollos, mayonesa grande (500g), farofa grande (400g) y 4 panes de ajo. Rinde 5-6 personas."),
]

revenue = sum(q * p for _, q, p, c, _ in mix_pt)
cmv = sum(q * c for _, q, p, c, _ in mix_pt)
tax = revenue * 0.04
fees = revenue * 0.02
fixed = 6690.00
profit = revenue - cmv - tax - fees - fixed
cm_ratio = (revenue - cmv - tax - fees) / revenue
breakeven = fixed / cm_ratio
payback = 20000.0 / profit

def money(x):
    return f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="B0B0B0", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxa}" w:type="dxa"/>')
    tcPr.append(tcW)

def build_styled_table(doc, headers, rows, widths=None, font_size=8.5, align_right_cols=None):
    if align_right_cols is None:
        align_right_cols = []
    
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t, color="B0B0B0", sz="5")
    
    header_tr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
    header_tr.append(tblHeader)
    
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = str(h)
        set_cell_shading(c, "E2ECF6")
        set_cell_margins(c, top=100, bottom=100, left=130, right=130)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else (WD_ALIGN_PARAGRAPH.CENTER if len(str(h)) < 18 else WD_ALIGN_PARAGRAPH.LEFT)
        for r in p.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            
    for row_idx, row in enumerate(rows):
        row_cells = t.add_row().cells
        bg_color = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row):
            c = row_cells[i]
            c.text = str(val)
            if bg_color != "FFFFFF":
                set_cell_shading(c, bg_color)
            set_cell_margins(c, top=80, bottom=80, left=130, right=130)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(font_size)
                if row[0] in ("TOTAL", "Total", "TOTAL MENSAL", "TOTAL MENSUAL") or (isinstance(val, str) and "TOTAL" in val):
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
                    
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                set_cell_width(row.cells[i], w)
        tblPr = t._tbl.tblPr
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{sum(widths)}" w:type="dxa"/>')
        tblPr.append(tblW)
        
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)
    p_spacer.paragraph_format.line_spacing = 1.0
    return t

def add_figure_with_caption(doc, img_path, fig_num, title, source, is_pt, width_cm=15.5):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(3)
    
    run = p_img.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    
    label_prefix = f"Figura {fig_num} – "
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Cm(0)
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(2)
    r_lbl = p_cap.add_run(label_prefix)
    r_lbl.bold = True
    r_lbl.font.size = Pt(10)
    r_lbl.font.name = "Arial"
    r_title = p_cap.add_run(title)
    r_title.font.size = Pt(10)
    r_title.font.name = "Arial"
    
    src_prefix = "Fonte: " if is_pt else "Fuente: "
    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_src.paragraph_format.first_line_indent = Cm(0)
    p_src.paragraph_format.space_before = Pt(0)
    p_src.paragraph_format.space_after = Pt(8)
    r_src = p_src.add_run(src_prefix + source)
    r_src.italic = True
    r_src.font.size = Pt(9)
    r_src.font.name = "Arial"

def configure_doc_styles(doc):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    
    for name, size, space_b, space_a in [
        ('Heading 1', 14, 18, 8),
        ('Heading 2', 12, 14, 6),
        ('Heading 3', 12, 10, 4)
    ]:
        s = styles[name]
        s.font.name = 'Arial'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        s.paragraph_format.first_line_indent = Cm(0)
        s.paragraph_format.space_before = Pt(space_b)
        s.paragraph_format.space_after = Pt(space_a)
        s.paragraph_format.keep_with_next = True
        
    styles['Heading 1'].font.all_caps = True
    
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.first_line_indent = Cm(0)
    r_foot = footer.add_run("Casa de Assados Sofia | ")
    r_foot.font.size = Pt(9)
    r_foot.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    
    fld = parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>')
    footer._p.append(fld)

def add_p(doc, text, boldlead=None):
    p = doc.add_paragraph()
    if boldlead and text.startswith(boldlead):
        r_lead = p.add_run(boldlead)
        r_lead.bold = True
        p.add_run(text[len(boldlead):])
    else:
        p.add_run(text)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item)

print('Base generator setup completed successfully.')
