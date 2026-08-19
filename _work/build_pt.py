import os
import shutil
import math
from pathlib import Path
import docx
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT = Path(r"c:\Users\wilki\OneDrive\Documentos\Trabajo de Curso")
CHART_DIR = ROOT / "_work" / "charts"
IMG_REF_DIR = ROOT / "Casa_de_Assados_Sofia_15_Imagens_Referencia_PR_HR_v2"

def resolve_img(pr_name, chart_fallback):
    p1 = IMG_REF_DIR / pr_name
    if p1.exists():
        return p1
    p2 = CHART_DIR / chart_fallback
    if p2.exists():
        return p2
    p3 = ROOT / "_work" / chart_fallback
    if p3.exists():
        return p3
    return p2

IMG_FIG02 = resolve_img("Figura_02_Fachada_Sinalizacao_Embalagens.png", "brand_mockup_sofia.jpg")
IMG_FIG03 = resolve_img("Figura_03_Cardapio_Balcao.png", "cardapio_impresso_sofia.jpg")
IMG_FIG04 = resolve_img("Figura_04_CRM_Sofia_Mobile_WhatsApp.png", "cardapio_whatsapp_sofia.jpg")
IMG_FIG09 = resolve_img("Figura_09_Assadoras_Giratorias_GLP.png", "equip1_asadora_gas.jpg")
IMG_FIG10 = resolve_img("Figura_10_Churrasqueira_Costela_Bafo.png", "equip2_churrasqueira_carvao.jpg")
IMG_FIG11 = resolve_img("Figura_11_Coifa_Industrial_AISI304.png", "equip3_coifa_industrial.jpg")
IMG_FIG12 = resolve_img("Figura_12_Freezer_Horizontal_510L.png", "equip4_freezer_horizontal.jpg")
IMG_FIG13 = resolve_img("Figura_13_Refrigerador_Vertical_4_Portas.png", "equip5_refrigerador_inox.jpg")
IMG_FIG14 = resolve_img("Figura_14_Mesa_Inox_Balanca.png", "equip6_bancada_balanca.jpg")
IMG_FIG15 = CHART_DIR / "planta_baixa_sofia_pt_hd.png"
IMG_FIG16 = resolve_img("Figura_16_Combo_Classico_Sofia.png", "combo1_classico_sofia.jpg")
IMG_FIG17 = resolve_img("Figura_17_Combo_Costela_Suprema.png", "combo2_costela_sofia.jpg")
IMG_FIG18 = resolve_img("Figura_18_Combo_Dueto_Sofia.png", "combo3_dueto_sofia.jpg")
IMG_FIG19 = resolve_img("Figura_19_Kit_Churrasco_Familia.png", "combo4_familia_sofia.jpg")
IMG_FIG20 = resolve_img("Figura_20_Conceito_Final.png", "anexo_casa_assados_sofia.png")
IMG_FIG21 = CHART_DIR / "crm_login_portal.png"
IMG_FIG22 = CHART_DIR / "crm_console_atendimento_ia.png"
IMG_FIG23 = CHART_DIR / "crm_gestao_pedidos_kds.png"

mix_pt = [
    ("O Clássico da Sofia", 70, 69.90, 26.50, "1 Frango recheado inteiro (~1,4kg assado), farofa artesanal crocante (250g), maionese caseira tradicional de batata (300g). Serve 3 a 4 pessoas."),
    ("Costela Suprema", 35, 119.90, 48.00, "1kg de Costela bovina premium assada lentamente no bafo por 6 horas, mandioca na manteiga de garrafa (300g), vinagrete fresco e farofa da casa (250g). Serve 4 pessoas."),
    ("Dueto Sofia", 35, 94.90, 36.00, "Meio frango assado dourado + 500g de Costelinha de porco marinada em ervas, batatas rústicas douradas (300g) e farofa da casa (200g). Serve 3 a 4 pessoas."),
    ("Kit Churrasco Família", 20, 169.90, 68.00, "1 Frango recheado inteiro + 700g de Costela bovina no bafo + 4 Linguiças toscanas artesanais grelhadas, maionese grande (500g), farofa grande (400g) e pães de alho (4 un). Serve 5 a 6 pessoas."),
]

revenue = sum(q * p for _, q, p, c, _ in mix_pt)
cmv = sum(q * c for _, q, p, c, _ in mix_pt)
tax = revenue * 0.04
fees = revenue * 0.02
fixed = 6870.00
profit = revenue - cmv - tax - fees - fixed
cm_ratio = (revenue - cmv - tax - fees) / revenue
breakeven = fixed / cm_ratio

def money(x):
    return f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="B0B0B0", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxa}" w:type="dxa"/>')
    tcPr.append(tcW)

def build_styled_table(doc, headers, rows, widths=None, font_size=8.5, align_right_cols=None):
    if align_right_cols is None:
        align_right_cols = []
    
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t, color="B0B0B0", sz="5")
    
    header_tr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
    header_tr.append(tblHeader)
    
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = str(h)
        set_cell_shading(c, "E2ECF6")
        set_cell_margins(c, top=100, bottom=100, left=130, right=130)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else (WD_ALIGN_PARAGRAPH.CENTER if len(str(h)) < 18 else WD_ALIGN_PARAGRAPH.LEFT)
        for r in p.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            
    for row_idx, row in enumerate(rows):
        row_cells = t.add_row().cells
        bg_color = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row):
            c = row_cells[i]
            c.text = str(val)
            if bg_color != "FFFFFF":
                set_cell_shading(c, bg_color)
            set_cell_margins(c, top=80, bottom=80, left=130, right=130)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(font_size)
                if row[0] in ("TOTAL", "Total", "TOTAL MENSAL", "TOTAL BIÊNIO (2026-2028)") or (isinstance(val, str) and "TOTAL" in val):
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
                    
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                set_cell_width(row.cells[i], w)
        tblPr = t._tbl.tblPr
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{sum(widths)}" w:type="dxa"/>')
        tblPr.append(tblW)
        
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)
    p_spacer.paragraph_format.line_spacing = 1.0
    return t

def add_figure_with_caption(doc, img_path, fig_num, title, source, width_cm=15.0):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(3)
    
    run = p_img.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Cm(0)
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(2)
    r_lbl = p_cap.add_run(f"Figura {fig_num} – ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(10)
    r_lbl.font.name = "Arial"
    r_title = p_cap.add_run(title)
    r_title.font.size = Pt(10)
    r_title.font.name = "Arial"
    
    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_src.paragraph_format.first_line_indent = Cm(0)
    p_src.paragraph_format.space_before = Pt(0)
    p_src.paragraph_format.space_after = Pt(8)
    r_src = p_src.add_run("Fonte: " + source)
    r_src.italic = True
    r_src.font.size = Pt(9)
    r_src.font.name = "Arial"

def configure_doc_styles(doc):
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(3.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    
    for name, size, space_b, space_a in [
        ('Heading 1', 14, 18, 8),
        ('Heading 2', 12, 14, 6),
        ('Heading 3', 12, 10, 4)
    ]:
        s = styles[name]
        s.font.name = 'Arial'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        s.paragraph_format.first_line_indent = Cm(0)
        s.paragraph_format.space_before = Pt(space_b)
        s.paragraph_format.space_after = Pt(space_a)
        s.paragraph_format.keep_with_next = True
        
    styles['Heading 1'].font.all_caps = True
    
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.first_line_indent = Cm(0)
    r_foot = footer.add_run("Casa de Assados Sofia | ")
    r_foot.font.size = Pt(9)
    r_foot.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    
    fld = parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>')
    footer._p.append(fld)

def add_p(doc, text, boldlead=None):
    p = doc.add_paragraph()
    if boldlead and text.startswith(boldlead):
        r_lead = p.add_run(boldlead)
        r_lead.bold = True
        p.add_run(text[len(boldlead):])
    else:
        p.add_run(text)
    return p

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item)

def generate_portuguese_thesis(output_path):
    doc = Document()
    configure_doc_styles(doc)
    
    # ------------------ ELEMENTOS PRÉ-TEXTUAIS ------------------
    # Capa
    for text, space, bold, size in [
        ("COLÉGIO EXCELÊNCIA", 36, True, 12),
        ("CURSO TÉCNICO EM ADMINISTRAÇÃO E INFORMÁTICA", 65, True, 12),
        ("WILKIN BARBAN ROSABAL", 80, True, 12),
        ("CASA DE ASSADOS SOFIA", 8, True, 16),
        ("PLANO DE NEGÓCIO PARA IMPLANTAÇÃO DE UMA MICROEMPRESA DE ASSADOS COM GESTÃO POR CRM EM CURITIBA - PR", 110, True, 12.5),
        ("CURITIBA - PR\n2026", 0, True, 12)
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(space)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = 'Arial'
        r.font.size = Pt(size)
    doc.add_page_break()

    # Folha de Rosto
    p = doc.add_paragraph('WILKIN BARBAN ROSABAL')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(75)
    p.paragraph_format.first_line_indent = Cm(0)
    
    title = "CASA DE ASSADOS SOFIA: PLANO DE NEGÓCIO PARA UMA OPERAÇÃO DE FINS DE SEMANA APOIADA POR CRM"
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(60)
    p.paragraph_format.first_line_indent = Cm(0)
    
    note = "Trabalho de Conclusão de Curso apresentado ao Curso Técnico em Administração e Informática do Colégio Excelência, como requisito parcial para obtenção do título de Técnico em Administração e Informática.\n\nÁrea de Concentração: Gestão Empresarial, Empreendedorismo e Tecnologia da Informação.\n\nOrientador(a): Prof(a). Me./Dr(a). ______________________________"
    p = doc.add_paragraph(note)
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph('CURITIBA - PR\n2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    doc.add_page_break()

    # Folha de Aprovação
    p = doc.add_paragraph('WILKIN BARBAN ROSABAL')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(40)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.paragraph_format.space_after = Pt(40)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph(note.split('\n\nOrientador')[0])
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p = doc.add_paragraph("Aprovado em: _____ / _____ / 2026\n\nBANCA EXAMINADORA:")
    p.runs[0].bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(30)
    
    for member in [
        ("Prof(a). Orientador(a) - Colégio Excelência", "Presidente"),
        ("Prof(a). Avaliador(a) 1 - Colégio Excelência", "Membro 1"),
        ("Prof(a). Avaliador(a) 2 - Colégio Excelência", "Membro 2")
    ]:
        p = doc.add_paragraph("____________________________________________________\n" + member[0] + " (" + member[1] + ")")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(25)
        p.paragraph_format.line_spacing = 1.0
    doc.add_page_break()

    # Dedicatória
    doc.add_heading("DEDICATÓRIA", level=1)
    p = doc.add_paragraph("Dedico este trabalho à minha família, cuja paciência, estímulo constante e apoio incondicional foram as vigas mestras para a superação de cada desafio desta jornada acadêmica. Aos amigos que compreenderam as ausências necessárias e compartilharam os sonhos de empreender com propósito e excelência.")
    p.paragraph_format.left_indent = Cm(7)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.line_spacing = 1.2
    doc.add_page_break()

    # Agradecimentos
    doc.add_heading("AGRADECIMENTOS", level=1)
    agr_text = (
        "Agradeço, primeiramente, a Deus, por iluminar meus caminhos, conceder-me saúde e perseverança ao longo de toda a formação acadêmica.\n\n"
        "Aos professores e à coordenação do Curso Técnico em Administração e Informática do Colégio Excelência, por compartilharem seus conhecimentos teóricos e práticos, rigor técnico e entusiasmo pela gestão empresarial orientada por dados.\n\n"
        "Ao meu professor orientador, pela dedicação, leitura atenta, correções precisas e valiosas diretrizes metodológicas que permitiram transformar uma ideia de negócio em um plano estruturado, inovador e viável.\n\n"
        "A todos os colegas de turma, com os quais tive a honra de debater, aprender e construir uma visão crítica e moderna sobre os desafios do empreendedorismo contemporâneo."
    )
    for par in agr_text.split('\n\n'):
        add_p(doc, par)
    doc.add_page_break()

    # Epígrafe
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(7)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(140)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run("“A melhor maneira de prever o futuro é criá-lo. O planejamento não diz respeito a decisões futuras, mas às implicações futuras das decisões presentes.”\n\n(Peter F. Drucker)")
    r.italic = True
    doc.add_page_break()

    # Resumo
    doc.add_heading("RESUMO", level=1)
    res_text = (
        "Este Trabalho de Conclusão de Curso apresenta o plano de negócio para a implantação da Casa de Assados Sofia, uma microempresa gastronômica projetada para operar aos sábados, domingos e em todos os feriados nacionais, estaduais e municipais no bairro Umbará, em Curitiba - PR. O modelo de negócio fundamenta-se na oferta enxuta de quatro combos familiares de assados tradicionais (frango recheado, costela bovina ao bafo, costelinha suína e cortes combinados com acompanhamentos artesanais), comercializados via retirada programada no balcão e entrega em domicílio em raio restrito de 5 km. O diferencial estratégico repousa na integração transversal de um sistema próprio de Gestão do Relacionamento com Clientes (CRM Casa de Assados Sofia), que gerencia campanhas de pré-venda às sextas-feiras e vésperas de feriados, nivela a capacidade produtiva das churrasqueiras em janelas de 15 minutos, mitiga perdas de matérias-primas perecíveis e potencializa a fidelização e a recompra recorrente com total conformidade à Lei Geral de Proteção de Dados (LGPD). A metodologia adotada possui natureza aplicada, descritiva e documental, combinando análise mercadológica com dados oficiais atualizados do IBGE (população estimada em 1.830.795 habitantes e PIB total de R$ 120,06 bilhões em Curitiba), Prefeitura de Curitiba e SEBRAE, além de modelagem econômico-financeira de custos e receitas desenvolvida no âmbito do Colégio Excelência. No cenário-base dimensionado estritamente para os fins de semana regulares (160 combos mensais), o empreendimento adota um orçamento de implantação confortável de R$ 38.000,00 (estruturado com R$ 18.000,00 de capital próprio e R$ 20.000,00 via microcrédito da Fomento Paraná), gerando receita bruta mensal de R$ 15.809,00, margem de contribuição de 55,16%, lucro operacional líquido de R$ 1.850,46 (lucratividade de 11,71%), ponto de equilíbrio em R$ 12.454,37 (~126 combos equivalentes) e amortização total do capital investido na curva de maturação em 11 a 12 meses. Como fator de conservadorismo financeiro, as receitas provenientes de operações em feriados durante a semana (estimadas entre 10 e 12 dias adicionais por ano, gerando R$ 20.000 a R$ 24.000 em faturamento extra e R$ 6.250 a R$ 7.500 de lucro líquido adicional) foram tratadas como alavancagem de segurança e liquidez, não sendo computadas na meta mínima de equilíbrio. Conclui-se que o empreendimento é comercialmente e financeiramente viável, condicionando seu sucesso à rigorosa padronização de fichas técnicas, identidade visual marcante, cardápios otimizados para balcão e WhatsApp, homologação de fornecedores regionais (CEASA Curitiba), segurança jurídica nas contratações intermitentes e execução disciplinada dos processos operacionais orientados por dados."
    )
    add_p(doc, res_text)
    add_p(doc, "Palavras-chave: Plano de Negócio; Gastronomia de Conveniência; Gestão do Relacionamento com Clientes (CRM); Administração e Informática; Curitiba; Feriados.", boldlead="Palavras-chave:")
    doc.add_page_break()

    # Listas
    doc.add_heading("LISTA DE ILUSTRAÇÕES", level=1)
    figs = [
        ("Figura 1 – Mix Mensal de Vendas (Cenário Base: 160 Combos)", 12),
        ("Figura 2 – Identidade Visual Fotográfica, Fachada e Comunicação de Ponto de Venda", 15),
        ("Figura 3 – Design Gráfico do Cardápio Comercial Impresso para Balcão e Mostrador", 16),
        ("Figura 4 – Design e Interface do Cardápio Digital Interativo para WhatsApp e Mobile", 17),
        ("Figura 5 – Composição do Resultado Mensal (DRE Projetada)", 21),
        ("Figura 6 – Gráfico do Ponto de Equilíbrio Operacional", 22),
        ("Figura 7 – Projeção do Resultado Operacional em 12 Meses", 23),
        ("Figura 8 – Análise de Sensibilidade e Comparação de Cenários", 26),
        ("Figura 9 – Máquinas Giratórias de Frango a Gás GLP com Queimadores Infravermelhos", 29),
        ("Figura 10 – Churrasqueira Tradicional a Carvão para Bafo com Grelha Elevatória", 30),
        ("Figura 11 – Sistema de Coifa Industrial em Aço Inox com Exaustão Mecânica", 30),
        ("Figura 12 – Freezer Horizontal Comercial Dupla Ação de 510 Litros", 31),
        ("Figura 13 – Refrigerador Comercial Vertical de Inox de 4 Portas", 31),
        ("Figura 14 – Mesa Central de Manipulação Inox AISI 304 com Balança Digital", 32),
        ("Figura 15 – Planta Baixa Técnica e Fluxo Sanitário Unidirecional (60,0 m²)", 33),
        ("Figura 16 – Documentação Fotográfica do Combo 1: O Clássico da Sofia", 37),
        ("Figura 17 – Documentação Fotográfica do Combo 2: Costela Suprema no Bafo", 38),
        ("Figura 18 – Documentação Fotográfica do Combo 3: Dueto Sofia (Frango & Costelinha)", 38),
        ("Figura 19 – Documentação Fotográfica do Combo 4: Kit Churrasco Família", 39),
        ("Figura 20 – Conceito Ilustrativo: Fachada, Embalagens, Produtos e Estação CRM Casa de Assados Sofia", 40),
        ("Figura 21 – Portal Web de Acesso e Autenticação do Cliente e Operador", 41),
        ("Figura 22 – Console de Atendimento Omnichannel e IA Virtual 'Sofia' com DeepSeek LLM", 41),
        ("Figura 23 – Painel KDS de Gestão de Pedidos em Tempo Real e Faturamento Operacional", 42)
    ]
    for pt_title, p_num in figs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.add_run(pt_title)
        p.add_run(f"\t{p_num}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    doc.add_page_break()

    doc.add_heading("LISTA DE TABELAS", level=1)
    tabs = [
        ("Tabela 1 – Síntese de Identificação da Empresa", 8),
        ("Tabela 2 – Matriz de Objetivos Estratégicos e Metas por Horizonte", 9),
        ("Tabela 3 – Matriz Comparativa de Concorrentes Diretos e Indiretos", 10),
        ("Tabela 4 – Matriz de Homologação de Fornecedores Estratégicos", 11),
        ("Tabela 5 – Composto de Marketing (4Ps de Serviços)", 12),
        ("Tabela 6 – Engenharia de Cardápio, Preços de Venda e CMV Unitário", 13),
        ("Tabela 7 – Mapeamento da Jornada do Cliente e Indicadores do CRM Casa de Assados Sofia", 14),
        ("Tabela 8 – Cronograma de Rotinas Operacionais Semanais e Controle no CRM", 18),
        ("Tabela 9 – Quadro de Funções, Responsabilidades e Rotinas da Equipe", 19),
        ("Tabela 10 – Estratégia de Contratação e Regime de Trabalho Intermitente (CLT 452-A)", 19),
        ("Tabela 11 – Estimativa do Investimento Fixo Inicial (Máquinas e Exaustão)", 20),
        ("Tabela 12 – Estimativa do Capital de Giro e Despesas Pré-Operacionais", 20),
        ("Tabela 13 – Projeção de Faturamento Mensal no Cenário Base", 21),
        ("Tabela 14 – Estimativa Detalhada dos Custos Fixos Mensais", 21),
        ("Tabela 15 – Demonstrativo de Resultados do Exercício (DRE Projetada)", 22),
        ("Tabela 16 – Projeção de Fluxo de Caixa para 12 Meses de Operação", 23),
        ("Tabela 17 – Síntese dos Indicadores de Viabilidade Econômico-Financeira", 24),
        ("Tabela 18 – Cronograma e Projeção de Receitas Incrementais em Feriados (2026-2028)", 25),
        ("Tabela 19 – Análise de Sensibilidade em Três Cenários de Demanda", 26),
        ("Tabela 20 – Matriz Estratégica SWOT / FOFA", 27),
        ("Tabela 21 – Matriz de Gerenciamento de Riscos e Planos de Contingência", 28),
        ("Tabela 22 – Plano de Ação 5W2H para os Primeiros 30 Dias de Implantação", 29),
        ("Tabela 23 – Simulação de Nota Fiscal Eletrônica de Insumos e Carnes (NF-e)", 34),
        ("Tabela 24 – Simulação de Nota Fiscal Eletrônica de Máquinas e Equipamentos (NF-e)", 35),
        ("Tabela 25 – Quadro Síntese de Licenciamento, Alvarás e Regularização de Curitiba", 35),
        ("Tabela 26 – Minuta Estruturada do Contrato de Trabalho Intermitente (CLT Art. 452-A)", 36),
        ("Tabela 27 – Modelo de Recibo de Pagamento por Diária com Desdobramento Legal e INSS", 36),
        ("Tabela 28 – Dicionário de Dados do Sistema CRM Casa de Assados Sofia", 37)
    ]
    for pt_t, p_num in tabs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.add_run(pt_t)
        p.add_run(f"\t{p_num}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    doc.add_page_break()

    # Sumário
    doc.add_heading("SUMÁRIO", level=1)
    toc_items = [
        ("INTRODUÇÃO", 5),
        ("1.1 Contextualização e Justificativa", 5),
        ("1.2 Problema de Pesquisa", 6),
        ("1.3 Objetivos da Pesquisa", 6),
        ("1.4 Aspectos Metodológicos", 6),
        ("1.5 Fundamentação Teórica", 7),
        ("1 RESUMO EXECUTIVO", 8),
        ("1.1 Conceito do Negócio e Proposta de Valor", 8),
        ("1.2 Perfil do Empreendedor e Competências", 8),
        ("1.3 Missão, Visão e Valores Organizacionais", 9),
        ("1.4 Estrutura Jurídica e Enquadramento Tributário", 9),
        ("1.5 Localização Estratégica e Instalações", 9),
        ("1.6 Metas e Objetivos por Horizontes", 10),
        ("2 ANÁLISE DE MERCADO", 10),
        ("2.1 Contexto Econômico e Demográfico Atualizado de Curitiba e Umbará", 10),
        ("2.2 Dimensionamento do Mercado (TAM, SAM, SOM)", 11),
        ("2.3 Segmentação e Comportamento do Público-Alvo", 11),
        ("2.4 Mapeamento e Análise da Concorrência", 11),
        ("2.5 Fornecedores Estratégicos e Matriz de Homologação", 12),
        ("2.6 Protocolo de Validação Empírica Preliminar", 12),
        ("3 PLANO DE MARKETING, IDENTIDADE VISUAL E CRM", 13),
        ("3.1 Posicionamento Estratégico e os 4Ps de Serviços", 13),
        ("3.2 Engenharia do Cardápio e Fichas Técnicas dos Combos", 13),
        ("3.3 Identidade Visual, Slogan e Comunicação de Marca", 14),
        ("3.4 Design Gráfico dos Cardápios: Versão Impressa e Versão WhatsApp", 15),
        ("3.5 Praça e Canais de Distribuição", 17),
        ("3.6 Promoção, Comunicação e Presença Digital Local", 17),
        ("3.7 Sistema CRM Casa de Assados Sofia: Conceituação, Estratégia Transversal e Retenção", 18),
        ("3.8 Jornada do Cliente, Funil de Conversão e Indicadores", 18),
        ("4 PLANO OPERACIONAL E TECNOLÓGICO", 19),
        ("4.1 Arranjo Físico e Fluxo Sanitário Unidirecional (RDC 216)", 19),
        ("4.2 Capacidade Instalada e Dimensionamento de Equipamentos", 19),
        ("4.3 Gestão de Gargalos e Balanceamento com CRM", 20),
        ("4.4 Mapeamento do Processo Produtivo Semanal", 20),
        ("4.5 Estrutura Organizacional e Estratégia de Contratação dos Diaristas", 20),
        ("4.6 Requisitos Regulatórios e Licenciamento Municipal", 21),
        ("4.7 Arquitetura Tecnológica do CRM Casa de Assados Sofia e Infraestrutura de Baixo Custo", 21),
        ("4.8 Gestão de Estoques (PEPS) e Sustentabilidade", 22),
        ("5 PLANO FINANCEIRO", 22),
        ("5.1 Investimento Inicial Total", 22),
        ("5.2 Estrutura de Financiamento e Fontes de Recursos", 23),
        ("5.3 Custos Variáveis Unitários e CMV dos Combos", 23),
        ("5.4 Custos Fixos Mensais Detalhados", 23),
        ("5.5 Demonstrativo de Resultados do Exercício (DRE Projetada)", 24),
        ("5.6 Fluxo de Caixa Projetado para 12 Meses", 24),
        ("5.7 Indicadores de Viabilidade e Ponto de Equilíbrio", 25),
        ("5.8 Impacto Operacional e Financeiro dos Feriados como Alavancagem Adicional (2026-2028)", 25),
        ("6 ANÁLISE DE VIABILIDADE E GESTÃO DE RISCOS", 26),
        ("6.1 Matriz SWOT / FOFA Estratégica", 26),
        ("6.2 Análise de Sensibilidade em Três Cenários", 26),
        ("6.3 Avaliação dos Indicadores de Viabilidade", 27),
        ("6.4 Matriz de Riscos e Planos de Contingência", 27),
        ("7 ANEXOS E INSTRUMENTOS DE IMPLANTAÇÃO", 28),
        ("7.1 Plano de Ação 5W2H de 30 Dias", 28),
        ("7.2 Catálogo Fotográfico de Maquinários e Equipamentos Adquiridos", 28),
        ("7.3 Planta Baixa Arquitetônica e Layout Funcional", 33),
        ("7.4 Simulação de Notas Fiscais Eletrônicas e Documentos Fiscais (NF-e)", 34),
        ("7.5 Quadro de Licenciamento, Alvarás e Regularização Sanitária", 35),
        ("7.6 Instrumentos de Contratação e Recibos de Diaristas", 36),
        ("7.7 Questionário Estruturado de Pesquisa de Mercado", 36),
        ("7.8 Dicionário de Dados do Sistema CRM Casa de Assados Sofia", 37),
        ("7.9 Renders e Documentação Fotográfica dos Combos Familiares", 37),
        ("CONCLUSÃO", 41),
        ("REFERÊNCIAS", 42),
    ]
    for pt_item, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(pt_item)
        if pt_item.startswith(('INTRODUÇÃO', '1 ', '2 ', '3 ', '4 ', '5 ', '6 ', '7 ', 'CONCLUSÃO', 'REFERÊNCIAS')):
            r.bold = True
        p.add_run(f"\t{page}")
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), 2, 1)
    
    note_p = doc.add_paragraph("Paginação de referência formatada conforme padrões ABNT.")
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_p.paragraph_format.first_line_indent = Cm(0)
    note_p.runs[0].italic = True
    note_p.runs[0].font.size = Pt(9)
    doc.add_page_break()

    # ------------------ ELEMENTOS TEXTUAIS ------------------
    # INTRODUÇÃO
    doc.add_heading("INTRODUÇÃO", level=1)
    add_p(doc, "O mercado contemporâneo de alimentação fora do lar (foodservice) no Brasil vivencia profundas transformações estruturais, impulsionadas pela busca crescente por conveniência, pela valorização do tempo em família durante os finais de semana e pela digitalização dos canais de atendimento e consumo. No contexto das grandes capitais da Região Sul, em especial Curitiba - PR, o hábito cultural do consumo dominical de carnes assadas (frango recheado, costela bovina ao bafo e churrasco tradicional) constitui uma sólida tradição gastronômica. Todavia, o modelo convencional dos estabelecimentos de bairro padece comumente de graves ineficiências operacionais: longas filas de espera a céu aberto, ausência de previsibilidade de demanda, desperdício severo de carnes e acompanhamentos não comercializados e falta absoluta de relacionamento continuado e inteligente com o cliente.")
    add_p(doc, "Nesse cenário, surge a oportunidade mercadológica para a concepção da Casa de Assados Sofia, uma microempresa gastronômica situada no bairro Umbará, polo residencial em expansão na zona sul de Curitiba. Diferenciando-se das rotisserias e assadeiras tradicionais, o empreendimento propõe um modelo operacional enxuto e estruturado, operando com quatro combos familiares padronizados, comercializados aos sábados, domingos e em todos os feriados nacionais, estaduais e municipais, sustentado por um sistema proprietário de Gestão de Relacionamento com Clientes (CRM Casa de Assados Sofia).")

    doc.add_heading("1.1 Contextualização e Justificativa", level=2)
    add_p(doc, "A justificativa para a elaboração deste plano de negócio fundamenta-se na necessidade de planejar cientificamente a criação de uma microempresa em um setor altamente competitivo e vulnerável a oscilações de custos de insumos perecíveis. A escolha do bairro Umbará, integrante da Administração Regional Bairro Novo de Curitiba, decorre de sua densidade populacional, predominância de núcleos familiares de classe média e carência de opções que combinem alta qualidade gastronômica artesanal com agilidade e previsibilidade de atendimento digital.")
    add_p(doc, "Do ponto de vista acadêmico e gerencial, a relevância deste trabalho reside na aplicação prática de ferramentas interdisciplinares de administração e informática — finanças, marketing de serviços, pesquisa operacional, gestão de estoques e desenvolvimento de sistemas de informação — integradas na formação técnica do Colégio Excelência. Demonstra-se que o uso rigoroso de tecnologia e dados não é privilégio exclusivo de grandes corporações, constituindo, ao contrário, a alavanca indispensável para a sustentabilidade, rentabilidade e mitigação de riscos de uma microempresa familiar.")

    doc.add_heading("1.2 Problema de Pesquisa", level=2)
    add_p(doc, "Diante das oportunidades e gargalos inerentes ao setor de alimentação de conveniência, formula-se o seguinte problema central de pesquisa: Sob quais condições de mercado, operacionais, financeiras e tecnológicas torna-se viável e sustentável a implantação da Casa de Assados Sofia no bairro Umbará, em Curitiba - PR?")

    doc.add_heading("1.3 Objetivos da Pesquisa", level=2)
    add_p(doc, "Objetivo Geral:", boldlead="Objetivo Geral:")
    add_p(doc, "Elaborar um plano de negócio detalhado, rigoroso e integrado para a implantação e operação da microempresa gastronômica Casa de Assados Sofia no bairro Umbará, em Curitiba - PR.")
    add_p(doc, "Objetivos Específicos:", boldlead="Objetivos Específicos:")
    add_bullets(doc, [
        "Realizar o diagnóstico do macroambiente e do mercado local de alimentação de fins de semana na região sul de Curitiba, mapeando concorrentes diretos, indiretos e perfis de consumidores com dados oficiais atualizados do IBGE e IPARDES.",
        "Estruturar a estratégia de marketing de serviços (4Ps), a identidade visual e slogan, o design dos cardápios (impresso de balcão e digital WhatsApp) e a inserção do CRM Casa de Assados Sofia como motor de pré-vendas e fidelização recorrente.",
        "Mapear os processos operacionais, sanitários (RDC 216 Anvisa), a planta arquitetônica com fluxo unidirecional, o catálogo de cada equipamento real adquirido e a capacidade produtiva das churrasqueiras.",
        "Definir a estratégia jurídica de contratação da equipe operacional de fim de semana (diaristas sob regime de trabalho intermitente - Art. 452-A da CLT), com minutas e recibos claros para erradicar riscos de passivo trabalhista.",
        "Desenvolver a modelagem econômico-financeira completa sob o orçamento confortável de R$ 38.000,00 de investimento inicial, projetando custos fixos e variáveis, DRE mensal, fluxo de caixa a 12 meses, ponto de equilíbrio e payback.",
        "Apresentar a projeção de receitas incrementais auferidas em dias feriados durante a semana no biênio 2026-2028 como mecanismo de alavancagem financeira e reserva de segurança.",
        "Apresentar a documentação fotográfica detalhada e individual dos quatro combos gastronômicos reais e analisar a viabilidade do negócio através de matriz SWOT e mitigação de riscos e proteção de dados (LGPD)."
    ])

    doc.add_heading("1.4 Aspectos Metodológicos", level=2)
    add_p(doc, "Para o cumprimento dos objetivos delineados, adotou-se uma metodologia de pesquisa aplicada, com abordagem mista (qualitativa e quantitativa) e alcance descritivo-exploratório. O procedimento técnico baseou-se em pesquisa documental e bibliográfica a partir das diretrizes metodológicas do SEBRAE (2013), manuais de orientação técnica do Colégio Excelência e normas acadêmicas da ABNT.")
    add_p(doc, "A coleta de dados secundários sustentou-se em estatísticas oficiais recentes do Instituto Brasileiro de Geografia e Estatística (IBGE Censo 2022 e estimativa 2025), Instituto Paranaense de Desenvolvimento Econômico e Social (IPARDES), cadastros da Prefeitura Municipal de Curitiba, indicadores da CEASA Paraná e relatórios setoriais da Associação Brasileira de Bares e Restaurantes (Abrasel). A dimensão quantitativa desenvolveu-se mediante engenharia de cardápio e modelagem de custos unitários (CMV), construindo simulações de fluxo de caixa e sensibilidade em modelos matemáticos integrados.")

    doc.add_heading("1.5 Fundamentação Teórica", level=2)
    add_p(doc, "A sustentação conceitual deste trabalho ancora-se na Teoria do Desenvolvimento Econômico de Joseph Schumpeter (1997), que concebe o empreendedor como o agente de inovação responsável por novas combinações produtivas. Na literatura de gestão contemporânea, Dornelas (2021) e Dolabela (2008) enfatizam que o plano de negócio constitui o instrumento indispensável para transformar visões intuitivas em estratégias auditáveis, reduzindo substancialmente as taxas de mortalidade prematura das microempresas.")
    add_p(doc, "No âmbito do marketing de relacionamento e tecnologia da informação, utilizam-se os postulados clássicos de Philip Kotler (2018), Ronald Swift (2001) e Don Peppers e Martha Rogers (2004). O conceito de Customer Relationship Management (CRM) — Gestão do Relacionamento com o Cliente — é definido não como um mero software utilitário, mas como uma estratégia empresarial holística orientada a identificar, adquirir, diferenciar, reter e maximizar o valor do cliente ao longo do tempo (Customer Lifetime Value - LTV). O CRM estrutura-se classicamente em três dimensões integradas: (a) CRM Operacional, voltado à automação das frentes de atendimento, recepção de pedidos e integração direta com a produção (painel KDS); (b) CRM Analítico, responsável pela mineração de dados transacionais, segmentação por recência, frequência e valor (RFM) e previsão inteligente de demanda; e (c) CRM Colaborativo, que sincroniza os canais de interação (WhatsApp e plataforma web) e retroalimentação pós-venda. Em uma operação gastronômica de conveniência, a aplicação do CRM erradica a assimetria de informações e transforma vendas esporádicas e anônimas em fluxos de caixa altamente previsíveis e rentáveis.")

    # CAPÍTULO 1
    doc.add_page_break()
    doc.add_heading("1 RESUMO EXECUTIVO", level=1)
    add_p(doc, "A Casa de Assados Sofia é um empreendimento gastronômico concebido para atender à crescente demanda por refeições familiares de alta qualidade aos finais de semana e feriados em Curitiba. Operando sob o formato dark store/takeaway com entrega rápida em raio controlado, a empresa especializa-se em cortes assados tradicionais (frango recheado, costela bovina ao bafo, costelinha suína e guarnições caseiras), comercializados exclusivamente em combos prontos para consumo.")

    doc.add_heading("1.1 Conceito do Negócio e Proposta de Valor", level=2)
    add_p(doc, "A proposta de valor sintetiza-se no conceito: 'O almoço de domingo da sua família resolvido com excelência, sabor artesanal e pontualidade britânica'. Enquanto os concorrentes tradicionais forçam os consumidores a enfrentar longas filas e incertezas sobre a disponibilidade de produtos, a Casa de Assados Sofia opera com reservas programadas através do CRM, garantindo retirada sem espera e entrega no horário exato com alimentos quentes e crocantes.")

    build_styled_table(doc, ["Elemento de Identificação", "Definição Estratégica"], [
        ["Razão Social Projetada", "Casa de Assados Sofia Ltda."],
        ["Nome Fantasia", "Casa de Assados Sofia"],
        ["Natureza Jurídica", "Sociedade Limitada Unipessoal (SLU)"],
        ["Enquadramento Tributário", "Microempresa (ME) optante pelo Simples Nacional"],
        ["Setor de Atividade", "Alimentação Fora do Lar / Gastronomia de Conveniência"],
        ["Localização da Unidade", "Rua Deputado Pinheiro Júnior, 1380, Umbará, Curitiba - PR, CEP 81930-000"],
        ["Proprietário / Administrador", "Wilkin Barban Rosabal (Técnico em Administração e Informática - Colégio Excelência)"],
        ["Investimento Inicial Total", "R$ 38.000,00 (R$ 18.000,00 Capital Próprio + R$ 20.000,00 Microcrédito Fomento Paraná)"],
        ["Diferencial Competitivo", "Operação enxuta de 4 combos + Produção programada por CRM próprio + Pontualidade"],
    ], widths=[2800, 6560])

    doc.add_heading("1.2 Perfil do Empreendedor e Competências", level=2)
    add_p(doc, "O empreendedor Wilkin Barban Rosabal assume a administração geral do negócio. Sua formação técnica interdisciplinar em Administração e Informática pelo Colégio Excelência proporciona domínio em planejamento orçamentário, engenharia de custos, modelagem de banco de dados, arquitetura de software CRM e governança de processos. O gestor será responsável direto pelas compras estratégicas na CEASA, controle de caixa, parametrização do CRM Casa de Assados Sofia, relacionamento com fornecedores e auditoria contínua dos padrões sanitários.")

    doc.add_heading("1.3 Missão, Visão e Valores Organizacionais", level=2)
    add_bullets(doc, [
        "Missão: Proporcionar às famílias momentos de união e celebração à mesa nos finais de semana e feriados, entregando assados artesanais de sabor inigualável, com pontualidade rigorosa e atendimento humanizado apoiado por tecnologia.",
        "Visão: Consolidar-se até 2029 como a principal referência em refeições de fins de semana e feriados sob encomenda e entrega rápida na zona sul de Curitiba, destacando-se pela consistência gastronômica e excelência em relacionamento com o cliente.",
        "Valores: Respeito à tradição culinária; Rigor sanitário e transparência em boas práticas; Pontualidade inegociável; Orientação a dados com respeito à privacidade (LGPD); Sustentabilidade e combate ao desperdício de alimentos."
    ])

    doc.add_heading("1.4 Estrutura Jurídica e Enquadramento Tributário", level=2)
    add_p(doc, "Optou-se pela constituição de uma Sociedade Limitada Unipessoal (SLU), enquadrada como Microempresa (ME) no regime tributário do Simples Nacional. A escolha da SLU garante a separação jurídica patrimonial completa entre os bens pessoais do sócio e as obrigações da empresa. Embora a figura do Microempreendedor Individual (MEI) ofereça simplificação burocrática, o faturamento anual projetado da Casa de Assados Sofia (R$ 189.708,00 no cenário base) e a necessidade de equipe de quatro colaboradores tornam o MEI legalmente inviável, evitando contingências fiscais severas.")

    doc.add_heading("1.5 Localização Estratégica e Instalações", level=2)
    add_p(doc, "O ponto comercial localiza-se na Rua Deputado Pinheiro Júnior, 1380, no bairro Umbará. A localização oferece acesso imediato às principais vias coletoras da região (Rua Nicola Pellanda e Estrada do Ganchinho), viabilizando distribuição logística ágil para os bairros adjacentes (Sítio Cercado, Pinheirinho e Ganchinho). O imóvel de 60 m² foi dimensionado para comportar recepção e inspeção de matérias-primas, congelamento e refrigeração, área de pré-preparo, setor de cocção isolado com sistema de coifa e exaustão industrial, bancada limpa de montagem, balcão de expedição rápida e área de higienização de utensílios.")

    doc.add_heading("1.6 Metas e Objetivos por Horizontes", level=2)
    build_styled_table(doc, ["Horizonte Temporal", "Meta Operacional / Comercial", "Evidência / Controle no CRM"], [
        ["Curto Prazo (0 a 3 meses)", "Validação do cardápio e piloto comercial; atingir a meta base de 160 combos/mês com 90% de pontualidade.", "Registro de pedidos, monitoramento de gargalos, atrasos < 5% e controle de perdas < 4%."],
        ["Médio Prazo (4 a 12 meses)", "Elevação da taxa de recompra para 45%; atingir entre 220 e 248 combos/mês; consolidar margem líquida de 14%.", "Análise de cohort, histórico de consumo por cliente, NPS médio > 85 e churn < 1%."],
        ["Longo Prazo (13 a 24 meses)", "Avaliação para aquisição de 2 assadoras adicionais; expansão do raio de entrega para 8 km.", "Relatórios de demanda reprimida por saturação de capacidade, ROI incremental e LTV consolidado."]
    ], widths=[2200, 3600, 3560])

    # CAPÍTULO 2
    doc.add_page_break()
    doc.add_heading("2 ANÁLISE DE MERCADO", level=1)
    add_p(doc, "A análise de mercado fundamenta-se no exame do macroambiente socioeconômico de Curitiba e na dinâmica de consumo das famílias residentes no bairro Umbará e adjacências.")

    doc.add_heading("2.1 Contexto Econômico e Demográfico Atualizado de Curitiba e Umbará", level=2)
    add_p(doc, "Conforme os dados oficiais mais recentes disponibilizados pelo IBGE e pelo IPARDES, Curitiba registrou no Censo Demográfico de 2022 uma população de 1.773.718 habitantes, alcançando uma estimativa populacional atualizada de 1.830.795 habitantes. O Produto Interno Bruto (PIB) municipal atingiu R$ 120,06 bilhões, consolidando a capital paranaense como a maior economia municipal de toda a Região Sul do Brasil e a sexta maior do país. O PIB per capita oficial mais recente é de R$ 67.691,30, demonstrando elevado poder aquisitivo médio da população.")
    add_p(doc, "No âmbito da microrregião de atuação, a Administração Regional Bairro Novo — que congrega os bairros Umbará, Ganchinho e Sítio Cercado — concentra mais de 165.000 habitantes. Trata-se de uma região fortemente residencial, com expressivo contingente de famílias de classe média (estratos B2, C1 e C2), caracterizada por alta taxa de ocupação formal durante a semana e forte valorização do descanso e convívio doméstico aos domingos e feriados.")

    doc.add_heading("2.2 Dimensionamento do Mercado (TAM, SAM, SOM)", level=2)
    add_bullets(doc, [
        "Mercado Total Disponível (TAM): A totalidade do mercado de foodservice de fins de semana e feriados em Curitiba, estimado em mais de 450 mil domicílios consumidores.",
        "Mercado Endereçável Servível (SAM): Domicílios situados no raio operacional de 5 km em torno do bairro Umbará, abrangendo aproximadamente 35.000 famílias.",
        "Mercado Atingível e Ocupável (SOM): Dimensionado estritamente pela capacidade produtiva da empresa, correspondendo a 160 combos mensais no cenário base (40 combos por final de semana) e até 248 combos mensais na maturidade operacional."
    ])

    doc.add_heading("2.3 Segmentação e Comportamento do Público-Alvo", level=2)
    add_p(doc, "O público-alvo prioritário é constituído por núcleos familiares de 3 a 6 pessoas, residentes em residências próprias ou condomínios horizontais na região sul de Curitiba. As principais motivações de compra identificadas são: (a) Eliminação do tempo e esforço exigidos para o preparo de almoços dominicais e de feriados; (b) Desejo de consumir churrasco e assados tradicionais com sabor de receita de família; (c) Busca por pontualidade rigorosa e facilidade de contratação prévia via mensagens digitais.")

    doc.add_heading("2.4 Mapeamento e Análise da Concorrência", level=2)
    build_styled_table(doc, ["Categoria de Concorrente", "Principais Forças / Vantagens", "Fraquezas / Lacunas", "Estratégia de Diferenciação da Sofia"], [
        ["Assadores Tradicionais de Bairro", "Tradição, ponto de passagem e proximidade física.", "Filas longas, sem reservas, risco de falta de produtos, atendimento lento e qualidade oscilante.", "Reserva antecipada via CRM, janela de 15 min de retirada, combos fechados e embalagens seladas."],
        ["Rotisserias de Supermercados", "Preços baixos, grande fluxo de clientes e escala.", "Carne ressecada por permanência em estufas, atendimento impessoal e falta de frescor.", "Cocção artesanal em lotes controlados, tempero caseiro exclusivo e foco familiar."],
        ["Aplicativos e Dark Kitchens", "Variedade de cardápio e conveniência do app genérico.", "Taxas abusivas (20% a 27%), atrasos frequentes nas entregas e comida fria.", "Canal próprio direto via WhatsApp/CRM, entrega com caixas térmicas rígidas em raio de 5 km."]
    ], widths=[2000, 2400, 2400, 2560])

    doc.add_heading("2.5 Fornecedores Estratégicos e Matriz de Homologação", level=2)
    add_p(doc, "A política de compras da Casa de Assados Sofia aproveita a proximidade da Central de Abastecimento do Paraná (CEASA Curitiba), situada a menos de 10 km, permitindo a aquisição direta de hortifrutigranjeiros frescos nas manhãs de sexta-feira. Para carnes e embalagens, estabeleceu-se uma matriz rigorosa de homologação com fornecedor principal e secundário para garantir segurança de abastecimento.")

    build_styled_table(doc, ["Grupo de Insumo", "Fornecedor Principal Homologado", "Fornecedor Secundário (Plano B)", "Critérios de Auditoria e Controle Sanitário"], [
        ["Aves (Frangos Resfriados)", "Frigorífico Avícola Regional (SIF)", "Distribuidora Atacadista Linha Verde", "Selo SIF/SIPPO, peças de 1,9-2,1kg, temperatura <= 4°C no recebimento."],
        ["Carnes Bovinas (Costela)", "Frigorífico Bovino Homologado (PR)", "Atacadão de Carnes Pinheirinho", "Inspeção sanitária, cobertura de gordura uniforme, rastreabilidade de lote."],
        ["Carnes Suínas e Linguiças", "Frigorífico Suíno Castro/PR", "Distribuidor Regional de Embutidos", "Padrão artesanal, controle de salinidade, embalagem a vácuo íntegra."],
        ["Hortifrúti (Batatas, Mandioca)", "Produtores Diretos CEASA Curitiba", "Distribuidor Hortifrúti Bairro Novo", "Calibre homogêneo, ausência de avarias mecânicas e frescor visual."],
        ["Embalagens Térmicas Seladas", "Distribuidora de Embalagens PR", "Atacadista Especializado Curitiba", "Resistência térmica a 90°C, atoxidade do material e vedação estanque."]
    ], widths=[1800, 2500, 2400, 2660])

    doc.add_heading("2.6 Protocolo de Validação Empírica Preliminar", level=2)
    add_p(doc, "Antes do início das operações comerciais plenas, será executado um piloto experimental de quatro finais de semana, com tetos progressivos de produção (25, 35, 45 e 55 combos por fim de semana). Esse protocolo permitirá calibrar o tempo exato de cocção das churrasqueiras, validar a aceitação das fichas técnicas de acompanhamentos, treinar a equipe de montagem e validar os fluxos de automação e agendamento do CRM Casa de Assados Sofia.")

    # CAPÍTULO 3
    doc.add_page_break()
    doc.add_heading("3 PLANO DE MARKETING, IDENTIDADE VISUAL E CRM", level=1)
    add_p(doc, "A estratégia de comercialização foca no posicionamento de uma marca confiável e acolhedora, que não disputa a 'guerra destrutiva de preços baixos', mas oferece valor superior em sabor, conveniência e pontualidade.")

    doc.add_heading("3.1 Posicionamento Estratégico e os 4Ps de Serviços", level=2)
    build_styled_table(doc, ["Dimensão (4Ps)", "Diretriz Estratégica", "Aplicação Prática na Casa de Assados Sofia"], [
        ["Produto", "Qualidade artesanal superior e padronização rigorosa.", "Quatro combos familiares balanceados, carnes marinadas por 24h, tempero próprio e guarnições frescas."],
        ["Preço", "Precificação por valor percebido com margem saudável.", "Preços entre R$ 69,90 e R$ 169,90, garantindo margem de contribuição média ponderada de 55,16%."],
        ["Praça", "Distribuição omnicanal focada em conveniência e temperatura.", "Retirada programada no balcão (takeaway) no Umbará e delivery próprio em raio de 5 km em até 20 min."],
        ["Promoção", "Comunicação hiperlocal direcionada e CRM de fidelização.", "Google Meu Negócio, Instagram com fotos de dar água na boca e campanhas de pré-venda semanal por WhatsApp."]
    ], widths=[1600, 3600, 4160])

    doc.add_heading("3.2 Engenharia do Cardápio e Fichas Técnicas dos Combos", level=2)
    table_menu_pt = []
    for nome, q, p, c, desc in mix_pt:
        marg = p - c
        table_menu_pt.append([nome, desc, money(p), money(c), money(marg), f"{(marg/p)*100:.1f}%"])
    build_styled_table(doc, ["Combo / Produto", "Composição Detalhada e Porções", "Preço Venda", "CMV Unit.", "Margem R$", "Margem %"], 
                       table_menu_pt, widths=[1800, 3800, 1100, 1100, 1100, 900], font_size=8, align_right_cols=[2,3,4,5])

    add_figure_with_caption(doc, CHART_DIR / "mix_pt.png", 1, "Mix Mensal de Vendas (Cenário Base: 160 Combos)", "Elaboração própria com base nas premissas do plano (2026).")

    doc.add_heading("3.3 Identidade Visual, Slogan e Comunicação de Marca", level=2)
    add_p(doc, "A identidade visual da Casa de Assados Sofia foi construída para transmitir acolhimento familiar, tradição rústica de churrasco e excelência no serviço. A ambientação e as peças de comunicação foram estruturadas nos seguintes elementos:")
    add_bullets(doc, [
        "Slogan Oficial: “O verdadeiro sabor do domingo na mesa da sua família.” — Reforça a centralidade do almoço dominical e a memória afetiva.",
        "Sub-slogan de Conveniência: “Tradição artesanal • Reserva sem filas • Entrega pontual” — Sintetiza a proposta de valor tecnológica e gastronômica.",
        "Fachada Comercial e Letreiro 3D Iluminado: Fachada moderna em madeira nobre tratada com painel preto carvão, letreiro luminoso em acrílico 3D e iluminação cênica quente.",
        "Paleta Cromática Oficial: Vermelho Brasa (#C0392B), Dourado Assado (#D4AC0D), Azul Confiança (#1F3864) e Preto Carvão (#2C3E50).",
        "Aplicações de Ponto de Venda: Cavalete rústico de calçada em madeira (1,0m x 0,6m) com cardápio do dia, e sacolas kraft personalizadas com lacres adesivos de segurança invioláveis (100% quente)."
    ])
    add_figure_with_caption(doc, IMG_FIG02, 2, "Identidade Visual Fotográfica, Fachada e Comunicação de Ponto de Venda", "Mockup fotográfico realista da fachada, sinalização e embalagens da Casa de Assados Sofia (2026).")

    doc.add_heading("3.4 Design Gráfico dos Cardápios: Versão Impressa e Versão WhatsApp", level=2)
    add_p(doc, "Para atender com máxima eficácia tanto o cliente presencial que comparece ao balcão quanto o usuário digital que realiza encomendas remotas, foram concebidas duas peças gráficas complementares e integradas:")
    add_p(doc, "a) Cardápio Comercial Impresso de Balcão (Figura 3): Desenvolvido em prancha rígida com acabamento preto fosco e detalhes em dourado refinado, ideal para exposição no balcão de atendimento e consulta rápida no ponto de venda físico. Apresenta os quatro combos, suas porções e guarnições artesanais de maneira clássica e elegante.")
    add_figure_with_caption(doc, IMG_FIG03, 3, "Design Gráfico do Cardápio Comercial Impresso para Balcão e Mostrador", "Fotografia de referência do cardápio comercial impresso em acabamento fosco e dourado (2026).")

    add_p(doc, "b) Cardápio Digital Interativo para WhatsApp e Mobile (Figura 4): Desenvolvido com interface moderna adaptada à usabilidade em smartphones. Apresenta fotos em alta definição dos pratos, cartões de seleção rápida dos combos, seletor de janelas horárias de retirada de 15 minutos (ex.: 11h30, 11h45, 12h00) e botão de fechamento de pedido com conexão direta à automação do CRM Casa de Assados Sofia.")
    add_figure_with_caption(doc, IMG_FIG04, 4, "Design e Interface do Cardápio Digital Interativo para WhatsApp e Mobile", "Interface de usuário (UI) mobile desenvolvida para atendimento e pré-venda via WhatsApp e CRM Casa de Assados Sofia (2026).")

    doc.add_heading("3.5 Praça e Canais de Distribuição", level=2)
    add_p(doc, "O modelo de atendimento divide-se em duas modalidades perfeitamente integradas ao sistema de agendamento do CRM:")
    add_bullets(doc, [
        "Retirada Programada (Takeaway): O cliente seleciona a janela de retirada (ex.: 11h45 às 12h00) durante a pré-venda. Ao comparecer ao balcão, seu pedido já se encontra embalado e acondicionado em estufa térmica, permitindo checkout e entrega em menos de 90 segundos.",
        "Delivery Próprio Controlado: Operação logística restrita a 5 km de raio no Umbará e bairros vizinhos. Os pedidos são transportados em caixas térmicas vedadas em rotas agrupadas por proximidade, garantindo chegada com temperatura superior a 65°C."
    ])

    doc.add_heading("3.6 Promoção, Comunicação e Presença Digital Local", level=2)
    add_p(doc, "A estratégia de atração de novos clientes baseia-se em marketing hiperlocal de baixo custo e alta conversão: (a) Otimização contínua do Google Perfil de Empresa (Google Meu Negócio) para buscas locais por 'frango assado Umbará' ou 'churrasco de domingo Curitiba'; (b) Perfil dinâmico no Instagram com vídeos reais da preparação das carnes nas sextas-feiras e vésperas de feriados; (c) Ações de degustação e parcerias com condomínios residenciais e comércios de conveniência do bairro.")

    doc.add_heading("3.7 Sistema CRM Casa de Assados Sofia: Conceituação, Estratégia Transversal e Retenção", level=2)
    add_p(doc, "O que é um CRM e por que ele é a coluna vertebral do negócio?", boldlead="O que é um CRM e por que ele é a coluna vertebral do negócio?")
    add_p(doc, "A sigla CRM (Customer Relationship Management), traduzida como Gestão do Relacionamento com o Cliente, representa muito mais do que um sistema de software ou um banco de dados de contatos: trata-se de uma filosofia de gestão empresarial e uma metodologia estratégica orientada a colocar o cliente no centro de todas as decisões da organização. Conforme ensinam Swift (2001) e Kotler (2018), o CRM integra processos, pessoas e tecnologia para mapear todas as interações ao longo da jornada do consumidor, permitindo entender seus hábitos de compra, antecipar suas demandas e cultivar relacionamentos personalizados, contínuos e de longo prazo.")
    add_p(doc, "No setor tradicional de assados de fim de semana, a grande maioria dos estabelecimentos opera no modelo 'cego e passivo': o comerciante abre as portas no domingo pela manhã sem saber quantos clientes aparecerão, acumulando filas caóticas de 40 minutos nos horários de pico (12h00) e sofrendo perdas graves de alimentos não vendidos no final do expediente. O Sistema CRM Casa de Assados Sofia transforma radicalmente essa realidade ao estruturar a operação em três pilares interdependentes:")
    add_bullets(doc, [
        "1. CRM Operacional (Front-Office e Automação de Pedidos): Automatiza o processo de pré-venda nas sextas-feiras e vésperas de feriados via WhatsApp próprio e catálogo web (https://casadeasados.duckdns.org/). Permite ao cliente escolher seu combo, guarnições e a janela exata de retirada de 15 minutos (ex.: 11h45 às 12h00). O sistema envia automaticamente a comanda para o painel KDS na cozinha, garantindo que o frango ou costela saia dourado da grelha no instante do despacho.",
        "2. CRM Analítico (Back-Office e Inteligência de Dados): Analisa o histórico de consumo da base de clientes através da matriz RFM (Recência, Frequência e Valor Monetário). Identifica automaticamente os clientes VIP (compras semanais), habituais (quinzenais) e em risco de churn (sem pedidos há mais de 35 dias), calculando indicadores cruciais como Custo de Aquisição de Clientes (CAC), Lifetime Value (LTV) e ticket médio por família.",
        "3. CRM Colaborativo (Pontos de Contato e Fidelização): Integra a comunicação direta por WhatsApp com canal de pós-venda humanizado. Três horas após o almoço de domingo, dispara pesquisa automatizada de satisfação (NPS de 1 a 5). Na quinta-feira seguinte, envia mensagens personalizadas com base nas preferências registradas, estimulando a recompra programada."
    ])
    add_p(doc, "Diferencial de Desenvolvimento Próprio e Custo Quase Zero:", boldlead="Diferencial de Desenvolvimento Próprio e Custo Quase Zero:")
    add_p(doc, "Diferente de negócios que contratam plataformas comerciais de SaaS que cobram mensalidades onerosas de R$ 300,00 a R$ 800,00 mais taxas sobre pedidos, ou que dependem de marketplaces (como iFood) que retêm comissões abusivas de 20% a 27% sobre o faturamento, o Sistema CRM Casa de Assados Sofia foi integralmente desenvolvido pelo próprio autor (Wilkin Barban Rosabal). Utilizando domínio dinâmico gratuito (https://casadeasados.duckdns.org/), certificado SSL Let's Encrypt gratuito e stack 100% em software livre (Linux Ubuntu Server, PostgreSQL, Python/FastAPI), a plataforma tem como único custo operacional a hospedagem do servidor VPS na nuvem (R$ 50,00 mensais), garantindo total soberania de dados, privacidade conforme a LGPD e economia direta de mais de R$ 4.000,00 anuais em licenciamento.")

    add_p(doc, "Atendente Virtual Inteligente 'Sofia' com IA Generativa (DeepSeek V4 Flash):", boldlead="Atendente Virtual Inteligente 'Sofia' com IA Generativa (DeepSeek V4 Flash):")
    add_p(doc, "O CRM Casa de Assados Sofia integra uma assistente virtual conversacional nativa baseada no modelo de linguagem DeepSeek V4 Flash, operando 24 horas por dia, 7 dias por semana, diretamente nos canais de WhatsApp e Web. A assistente 'Sofia' atende os clientes em linguagem natural fluida e acolhedora, orienta sobre a composição e rendimento dos quatro combos, tira dúvidas sobre acompanhamentos e ingredientes, sugere porções conforme o tamanho da família, agenda as janelas de retirada de 15 minutos, processa reservas automáticas de pré-venda e coleta avaliações de pós-venda (NPS).")
    add_p(doc, "Memória de Cálculo de Custos da IA para 1.000 Mensagens Mensais:", boldlead="Memória de Cálculo de Custos da IA para 1.000 Mensagens Mensais:")
    add_p(doc, "A adoção da API do DeepSeek V4 Flash proporciona um custo operacional sem precedentes para o empreendimento. Para uma demanda mensal projetada de 1.000 mensagens completas de atendimento (volume que atende com folga mais de 250 conversas de clientes por mês):\n"
          "• Volume de Entrada (System Prompt + Histórico + Cardápio): 1.000 requisições x 1.000 tokens = 1.000.000 tokens (1M tokens). À tarifa de $0,22 USD / 1M (com Context Caching a $0,007 USD / 1M), o custo de entrada totaliza $0,22 USD;\n"
          "• Volume de Saída (Respostas Geradas pela IA): 1.000 requisições x 150 tokens = 150.000 tokens (0,15M tokens). À tarifa de $0,66 USD / 1M, o custo de saída totaliza $0,099 USD;\n"
          "• Custo Total da API DeepSeek V4 Flash: $0,319 USD por mês (~$0,32 USD);\n"
          "• Conversão Cambial em Reais (Câmbio R$ 5,50 + IOF de 4,38%): Apenas R$ 1,84 a R$ 2,00 por mês.\n"
          "Enquanto ferramentas comerciais de chatbot com inteligência artificial (como Blip, ManyChat ou Z-API Pro) cobram mensalidades entre R$ 250,00 e R$ 600,00 por planos básicos, a Casa de Assados Sofia opera uma IA generativa de última geração por apenas dois reais mensais, garantindo economia superior a 99% em automação de atendimento e elevando drasticamente a satisfação do consumidor.")

    doc.add_heading("3.8 Jornada do Cliente, Funil de Conversão e Indicadores", level=2)
    build_styled_table(doc, ["Etapa da Jornada", "Ação do Cliente", "Ponto de Contato", "Ação do CRM Casa de Assados Sofia", "Indicador Chave (KPI)"], [
        ["1. Descoberta", "Busca almoço de domingo no Google/Instagram.", "Google Maps / Redes Sociais", "Captura contato e direciona para WhatsApp com link rastreado.", "Custo de Aquisição (CAC) e Leads."],
        ["2. Pré-Venda", "Recebe cardápio na sexta e seleciona combo.", "WhatsApp / Bot de Pedidos", "Confirma itens, sugere adicionais e reserva janela de retirada.", "Taxa de Conversão de Pré-Venda (> 35%)."],
        ["3. Produção", "Aguardando horário agendado.", "Cozinha da Unidade", "Gera comanda KDS agrupada por horário de saída da churrasqueira.", "Nível de Ocupação da Capacidade (%)."],
        ["4. Entrega / Retirada", "Retira no balcão ou recebe em casa.", "Balcão / Delivery", "Notifica cliente quando o pedido sai da churrasqueira.", "Índice de Pontualidade (> 92%) e Tempo Fila."],
        ["5. Pós-Venda", "Consome a refeição em família.", "Mensagem 3h após entrega", "Dispara pesquisa rápida de satisfação (NPS de 1 a 5).", "Net Promoter Score (NPS > 85)."],
        ["6. Fidelização", "Recebe incentivo de recompra.", "WhatsApp na quinta seguinte", "Dispara lembrete com base no combo favorito do histórico.", "Taxa de Recompra em 30 Dias (> 40%)."]
    ], widths=[1600, 2000, 1600, 2500, 1660], font_size=8)

    # CAPÍTULO 4
    doc.add_page_break()
    doc.add_heading("4 PLANO OPERACIONAL E TECNOLÓGICO", level=1)
    add_p(doc, "O plano operacional detalha o arranjo físico, o dimensionamento de equipamentos, o fluxo produtivo e a governança sanitária e tecnológica que sustentam a rotina de produção da Casa de Assados Sofia.")

    doc.add_heading("4.1 Arranjo Físico e Fluxo Sanitário Unidirecional (RDC 216)", level=2)
    add_p(doc, "Em estrita conformidade com a Resolução RDC nº 216 da Anvisa e as diretrizes da Vigilância Sanitária Municipal de Curitiba, o layout da unidade foi projetado com fluxo produtivo linear e unidirecional, impedindo expressamente qualquer cruzamento entre alimentos crus (matérias-primas recebidas) e alimentos prontos para consumo. As instalações contam com bancadas de aço inoxidável AISI 304, paredes revestidas com azulejo lavável de cor clara até o teto, piso cerâmico antiderrapante com ralos sifonados com fechamento e sistema profissional de coifa com exaustão mecânica e dutos de saída na área de cocção. A planta baixa técnica e layout operacional encontram-se ilustrados na Figura 15.")

    doc.add_heading("4.2 Capacidade Instalada e Dimensionamento de Equipamentos", level=2)
    add_p(doc, "O parque de máquinas e ferramentas reais adquirido foi dimensionado para atender com folga o cenário base de 160 combos mensais (40 combos por fim de semana), possuindo capacidade de expansão para até 260 combos mensais sem novos investimentos estruturais. O catálogo fotográfico individual dos equipamentos consta nas Figuras 9 a 14:")
    add_bullets(doc, [
        "2 Máquinas Giratórias de Frango a Gás GLP (Figura 9): Equipadas com queimadores infravermelhos traseiros a gás GLP, espetos rotativos de aço inoxidável e portas de vidro temperado, permitindo assar até 40 frangos por dia em 2 ciclos sequenciais.",
        "1 Churrasqueira Profissional a Carvão para Bafo (Figura 10): Estrutura em aço reforçado com tijolos refratários e grelha elevatória de manivela (1,50m), projetada para cocção lenta ao vapor de costelas bovinas por 6 horas.",
        "1 Sistema de Coifa Industrial em Aço Inox (Figura 11): Sistema de exaustão mecânica com filtros inerciais laváveis de alta retenção de gordura e duto circular de saída, atendendo à legislação da VISA Curitiba.",
        "1 Freezer Horizontal Comercial Dupla Ação de 510 Litros (Figura 12): Armazenamento e congelamento seguro de carnes com registro diário de temperatura (-18°C).",
        "1 Refrigerador Comercial Vertical de Inox de 4 Portas (Figura 13): Conservação de marinadas, guarnições e acompanhamentos preparados (+2°C a +4°C).",
        "2 Mesas Centrais de Manipulação Inox AISI 304 com Balança Digital (Figura 14): Superfície asséptica com balança computadora Inmetro, tábua de corte e cubas gastronômicas Gastronorm."
    ])

    doc.add_heading("4.3 Gestão de Gargalos e Balanceamento com CRM", level=2)
    add_p(doc, "O principal gargalo operacional das operações gastronômicas dominicais não é o assamento em si, mas a concentração de retiradas entre as 11h45 e as 12h45, gerando filas e atrasos na montagem dos acompanhamentos. O CRM Casa de Assados Sofia resolve esse gargalo através do algoritmo de fracionamento de pedidos: o sistema divide a capacidade de atendimento em janelas de 15 minutos (máximo de 6 pedidos por janela). Ao atingir o limite da janela, o sistema bloqueia automaticamente o horário para novas reservas, distribuindo o fluxo uniformemente entre as 11h00 e as 14h00.")

    doc.add_heading("4.4 Mapeamento do Processo Produtivo Semanal", level=2)
    build_styled_table(doc, ["Dia da Semana", "Horário", "Atividades Operacionais Críticas", "Controle e Registro no CRM Casa de Assados Sofia"], [
        ["Quinta-feira", "18h00 - 20h00", "Análise das reservas preliminares e emissão da lista de compras.", "Projeção de demanda consolidada por corte e insumo."],
        ["Sexta-feira", "06h30 - 11h00", "Compras na CEASA e recebimento de carnes homologadas (inspeção).", "Registro de lotes, datas de validade e custos unitários reais."],
        ["Sexta-feira", "13h00 - 18h00", "Porcionamento, marinada das carnes e disparo da pré-venda semanal.", "Disparo da campanha de WhatsApp e abertura de janelas."],
        ["Sábado", "06h30 - 10h30", "Acendimento das churrasqueiras, início da cocção e preparo de guarnições.", "Emissão de comandas KDS por franja de horário."],
        ["Sábado", "11h00 - 14h30", "Montagem de combos, expedição balcão, despacho de entregas e higienização.", "Baixa de pedidos entregues e monitoramento de pontualidade."],
        ["Domingo / Feriado", "06h00 - 15h00", "Ciclo operacional principal (pico de vendas semanais/feriados).", "Fechamento de caixa e disparo de pesquisas NPS."],
        ["Segunda-feira", "09h00 - 11h00", "Higienização profunda, balanço de estoque e análise de indicadores.", "Relatório gerencial de faturamento, margem e recompras."]
    ], widths=[1800, 1600, 3500, 2460], font_size=8)

    doc.add_heading("4.5 Estrutura Organizacional e Estratégia de Contratação dos Diaristas", level=2)
    add_p(doc, "A operação da Casa de Assados Sofia funciona exclusivamente aos sábados, domingos e feriados (8 a 10 dias por mês). A contratação informal de diaristas gera elevado passivo trabalhista conforme a jurisprudência consolidada da Justiça do Trabalho (TRT 9ª Região / TST), decorrente da habitualidade e subordinação. Para eliminar 100% desse risco com total segurança jurídica, o empreendimento adotará o Contrato de Trabalho Intermitente (Art. 452-A da CLT, instituído pela Lei nº 13.467/2017) ou Contrato Formal de Prestação de Serviços com Recibo de Pagamento a Autônomo (RPA).")

    build_styled_table(doc, ["Cargo / Função", "Colaborador / Vínculo", "Jornada Mensal", "Valor da Diária", "Estratégia Jurídico-Trabalhista (Segurança Total)"], [
        ["Gerente Geral", "Wilkin Barban (Sócio)", "Integral", "Pró-labore", "Sócio administrador com responsabilidade civil e técnica."],
        ["Churrasqueiro Chefe", "Diarista Especialista", "8-10 diárias/mês", "R$ 120,00", "Contrato Intermitente CLT (Art. 452-A) + ASO admissional com exames coprológicos."],
        ["Auxiliar de Cozinha 1", "Diarista de Pré-preparo", "8-10 diárias/mês", "R$ 120,00", "Contrato Intermitente CLT (Art. 452-A) + Treinamento RDC 216 e EPI completo."],
        ["Auxiliar de Montagem", "Diarista de Expedição", "8-10 diárias/mês", "R$ 120,00", "Contrato Intermitente CLT (Art. 452-A) + Registro de ponto e convocação via CRM."],
        ["Entregador (Motoboy)", "Diarista / Parceiro", "8-10 diárias/mês", "R$ 120,00", "Contrato Autônomo com MEI próprio + CNH categoria A + Seguro de acidentes."]
    ], widths=[1800, 1800, 1400, 1300, 3060], font_size=8)

    doc.add_heading("4.6 Requisitos Regulatórios e Licenciamento Municipal", level=2)
    add_p(doc, "A regularização formal da Casa de Assados Sofia cumprirá rigorosamente as exigências dos órgãos fiscalizadores de Curitiba: (a) Consulta Prévia de Viabilidade de Endereço aprovada na Secretaria Municipal de Urbanismo; (b) Inscrição Municipal e Alvará de Localização e Funcionamento; (c) Licença Sanitária expedida pela Vigilância Sanitária Municipal; (d) Certificado de Vistoria do Corpo de Bombeiros Militar do Paraná (CLCB); (e) Manual de Boas Práticas e Procedimentos Operacionais Padronizados (POPs) afixados na unidade.")

    doc.add_heading("4.7 Arquitetura Tecnológica do CRM Casa de Assados Sofia e Infraestrutura de Baixo Custo", level=2)
    add_p(doc, "A arquitetura tecnológica do Sistema CRM Casa de Assados Sofia foi integralmente projetada e implementada com base no paradigma de microsserviços leves, inteligência artificial generativa e tecnologias open source, garantindo máxima performance com custo operacional irrisório:")
    add_bullets(doc, [
        "Domínio Dinâmico e Certificado SSL Gratuito: O sistema é acessível publicamente através da URL https://casadeasados.duckdns.org/, utilizando serviço de DNS dinâmico gratuito (DuckDNS) com certificado TLS/SSL Let's Encrypt para tráfego 100% criptografado (HTTPS).",
        "Infraestrutura em Nuvem e Custo de TI Integrado: Hospedado em um Servidor Virtual Privado (VPS) baseado em Linux Ubuntu Server LTS (R$ 50,00/mês), integrado à API do modelo de linguagem DeepSeek V4 Flash (R$ 2,00/mês para 1.000 mensagens), totalizando R$ 52,00 mensais de despesas totais de TI e inteligência artificial.",
        "Integração com DeepSeek V4 Flash e Prompt Caching: Backend em Python/FastAPI conectado à API DeepSeek via chamadas assíncronas assíncronas de alta performance. Utiliza a tecnologia nativa de Context Caching (cache de prefixo a $0,007/1M tokens) para manter em memória o catálogo e instruções da 'Sofia', reduzindo a latência para menos de 750 ms por interação.",
        "Stack de Software Livre (Zero Licenciamento): Backend desenvolvido em Python/FastAPI de alta concorrência, banco de dados relacional PostgreSQL para integridade transacional (ACID) e frontend responsivo em HTML5, CSS3 e JavaScript puro sem dependência de bibliotecas proprietárias.",
        "Painel KDS Operacional para Cozinha: Interface em tempo real que organiza as comandas de preparo por ordem cronológica e janelas de retirada de 15 minutos, exibindo contadores regressivos para os assadores.",
        "Governança de Dados, Privacidade e Backups Criptografados: Sanitização e anonimização de dados antes do envio ao LLM (dados cadastrais protegidos localmente sob a LGPD), além de rotina diária automatizada de dump do banco de dados com criptografia AES-256 e sincronização segura com repositório remoto (RPO < 24h e RTO < 15 min)."
    ])

    doc.add_heading("4.8 Gestão de Estoques (PEPS) e Sustentabilidade", level=2)
    add_p(doc, "O controle de matérias-primas segue o princípio PEPS (Primeiro que Entra, Primeiro que Sai), garantindo o giro veloz dos estoques perecíveis. Graças ao modelo de pré-vendas, a quantidade de carne adquirida nas sextas-feiras coincide precisamente com 90% da demanda reservada, reduzindo o índice de perdas para menos de 3% da produção. Como política ambiental, todo o óleo vegetal saturado utilizado é acondicionado em bombonas e recolhido por empresa licenciada de reciclagem para produção de biodiesel, e os resíduos de carvão vegetal são destinados à compostagem agrícola.")

    # CAPÍTULO 5
    doc.add_page_break()
    doc.add_heading("5 PLANO FINANCEIRO", level=1)
    add_p(doc, "O plano financeiro consolida todas as estimativas de receitas, custos e investimentos em valores correntes de 2026 sob o cenário confortável e otimizado de R$ 38.000,00, demonstrando o equilíbrio contábil e a solidez do negócio.")

    doc.add_heading("5.1 Investimento Inicial Total", level=2)
    build_styled_table(doc, ["Item do Investimento Fixo (Equipamentos e Infraestrutura)", "Qtd", "Valor Unit. (R$)", "Subtotal (R$)"], [
        ["Máquinas giratórias de frango a gás GLP (novas c/ garantia de fábrica)", "2", "2.400,00", "4.800,00"],
        ["Churrasqueira profissional a carvão para bafo c/ elevador (1,5m reforçada)", "1", "2.200,00", "2.200,00"],
        ["Sistema de coifa industrial em aço inox c/ exaustão mecânica e duto (VISA)", "1", "4.200,00", "4.200,00"],
        ["Freezer horizontal comercial dupla ação 510L novo", "1", "3.100,00", "3.100,00"],
        ["Refrigerador comercial vertical em inox 4 portas novo", "1", "3.400,00", "3.400,00"],
        ["Mesas centrais de trabalho em inox AISI 304 (2,0x0,9m)", "2", "1.100,00", "2.200,00"],
        ["Computador / Terminal de balcão touch c/ impressora térmica 80mm", "1", "1.500,00", "1.500,00"],
        ["Balança digital comercial computadora c/ bateria (homologada Inmetro)", "1", "500,00", "500,00"],
        ["Caixas térmicas reforçadas tipo baú para motoboy (45L)", "2", "250,00", "500,00"],
        ["Lavadora de alta pressão profissional para higienização pesada", "1", "800,00", "800,00"],
        ["Utensílios de cozinha em inox, facas profissionais e cubas Gastronorm", "Vários", "1.300,00", "1.300,00"],
        ["TOTAL DO INVESTIMENTO FIXO", "-", "-", "24.500,00"]
    ], widths=[4800, 1000, 1800, 1760], align_right_cols=[1,2,3])

    build_styled_table(doc, ["Capital de Giro e Despesas Pré-Operacionais", "Destinação do Recurso", "Valor Estimado (R$)"], [
        ["Depósito Caução de Aluguel (3 meses) + 1º Mês de Locação", "Fiança locatícia de 60m² no Umbará", "4.000,00"],
        ["Estoque Inicial de Carnes e Insumos para Piloto Comercial", "Matérias-primas e temperos de arranque", "2.500,00"],
        ["Lote Inicial de Embalagens Térmicas Seladas e Sacolas (1.000 un)", "Embalagens seguras para 1.000 pedidos", "1.400,00"],
        ["Taxas de Licenciamento, Vistoria Bombeiros e Abertura SLU", "Legalização comercial, sanitária e bombeiros", "800,00"],
        ["Fachada Comercial em Madeira, Letreiro 3D e Cavalete", "Identidade visual do ponto de venda", "1.200,00"],
        ["Marketing de Lançamento, Sessão de Fotos e Degustação", "Campanha local de inauguração e captação", "1.100,00"],
        ["Fundo de Reserva de Capital de Giro Livre (Liquidez)", "Colchão de segurança para primeiros meses", "2.500,00"],
        ["TOTAL DE CAPITAL DE GIRO E PRÉ-OPERACIONAIS", "-", "13.500,00"],
        ["INVESTIMENTO TOTAL NECESSÁRIO (FIXO + GIRO)", "-", "38.000,00"]
    ], widths=[4500, 3100, 1760], align_right_cols=[2])

    doc.add_heading("5.2 Estrutura de Financiamento e Fontes de Recursos", level=2)
    add_p(doc, "O capital de R$ 38.000,00 será integralizado com R$ 18.000,00 (47,37%) de recursos próprios do empreendedor e R$ 20.000,00 (52,63%) oriundos da linha de microcrédito produtivo orientado da Fomento Paraná (Banco do Empreendedor), financiado em 36 parcelas fixas de R$ 680,00 mensais (taxa subsidiada para microempresas do setor de alimentação), contempladas integralmente nos custos fixos do projeto.")

    doc.add_heading("5.3 Custos Variáveis Unitários e CMV dos Combos", level=2)
    doc.add_heading("5.4 Custos Fixos Mensais Detalhados", level=2)
    build_styled_table(doc, ["Item de Custo Fixo", "Detalhamento e Memória de Cálculo", "Valor Mensal (R$)"], [
        ["Mão de Obra Operacional (Equipe Diaristas)", "4 pessoas x 8 diárias de fim de semana/mês x R$ 120,00/diária", "3.840,00"],
        ["Aluguel do Ponto Comercial", "Imóvel comercial de 60 m² na Rua Deputado Pinheiro Júnior", "1.000,00"],
        ["Utilidades Públicas (Água, Luz, Gás)", "Consumo operacional de água, energia trifásica e gás GLP", "350,00"],
        ["Serviços de Internet Fibra e Telefonia", "Linha móvel comercial + Conexão fibra 500 Mbps", "120,00"],
        ["Hospedagem Cloud e Servidor VPS CRM", "Hospedagem do VPS Linux, banco de dados e rotina de backups", "50,00"],
        ["Honorários Contábeis Mensais", "Assessoria contábil, fiscal e trabalhista Simples Nacional", "250,00"],
        ["Publicidade e Marketing Recorrente", "Anúncios hiperlocais no Instagram e panfletagem de bairro", "200,00"],
        ["Manutenção e Produtos de Higienização", "Detergentes industriais, sanitizantes e manutenção preventiva", "180,00"],
        ["Parcela do Microcrédito Fomento Paraná", "Amortização de empréstimo de R$ 20.000 (36 parcelas fixas)", "680,00"],
        ["Fundo de Reserva para Imprevistos", "Provisão para reposições menores e contingências", "200,00"],
        ["TOTAL DE CUSTOS FIXOS MENSAIS", "-", "6.870,00"]
    ], widths=[3400, 4200, 1760], align_right_cols=[2])

    doc.add_heading("5.5 Demonstrativo de Resultados do Exercício (DRE Projetada)", level=2)
    build_styled_table(doc, ["Linha do Demonstrativo Financeiro", "Base de Cálculo / Critério", "Valor Mensal (R$)", "Análise Vertical (%)"], [
        ["(=) RECEITA BRUTA OPERACIONAL", "160 combos no mix de vendas projetado", "15.809,00", "100,00%"],
        ["(-) Custo das Mercadorias Vendidas (CMV)", "Soma ponderada de carnes, guarnições e embalagens", "6.140,00", "38,84%"],
        ["(-) Impostos do Simples Nacional (4,0%)", "Alíquota efetiva de microempresa comercial", "632,36", "4,00%"],
        ["(-) Taxas de Meios de Pagamento (2,0%)", "Média ponderada de cartões de débito/crédito e PIX", "316,18", "2,00%"],
        ["(=) MARGEM DE CONTRIBUIÇÃO TOTAL", "Receita Bruta - Custos Variáveis - Impostos/Taxas", "8.720,46", "55,16%"],
        ["(-) CUSTOS FIXOS OPERACIONAIS TOTAIS", "Estrutura fixa mensal detalhada", "6.870,00", "43,46%"],
        ["(=) LUCRO OPERACIONAL LÍQUIDO", "Margem de Contribuição - Custos Fixos", "1.850,46", "11,71%"]
    ], widths=[3400, 3100, 1600, 1260], align_right_cols=[2,3])

    add_figure_with_caption(doc, CHART_DIR / "dre_pt.png", 5, "Composição do Resultado Mensal (DRE Projetada)", "Elaboração própria com base nas premissas financeiras do plano (2026).")

    doc.add_heading("5.6 Fluxo de Caixa Projetado para 12 Meses", level=2)
    cash_rows_pt = []
    cum = -38000.0
    for m, q in enumerate(range(160, 249, 8), 1):
        rev = q * (revenue / 160.0)
        op = rev * cm_ratio - fixed
        cum += op
        cash_rows_pt.append([f"Mês {m}", str(q), money(rev), money(rev * (1 - cm_ratio)), money(fixed), money(op), money(cum)])
    build_styled_table(doc, ["Mês", "Combos", "Receita (R$)", "Custos Var. (R$)", "Custos Fixos", "Lucro Mês", "Saldo Acumulado"], 
                       cash_rows_pt, widths=[1100, 1000, 1600, 1600, 1400, 1500, 1760], font_size=8, align_right_cols=[1,2,3,4,5,6])

    add_figure_with_caption(doc, CHART_DIR / "result12_pt.png", 7, "Projeção do Resultado Operacional em 12 Meses", "Elaboração própria com base na curva de maturação do negócio (2026).")

    doc.add_heading("5.7 Indicadores de Viabilidade e Ponto de Equilíbrio", level=2)
    build_styled_table(doc, ["Indicador Financeiro", "Fórmula de Cálculo", "Resultado Obtido", "Interpretação Gerencial"], [
        ["Índice de Margem de Contribuição", "(Receita - CMV - Impostos - Taxas) / Receita", "55,16%", "A cada R$ 100 vendidos, sobram R$ 55,16 para pagar custos fixos e gerar lucro."],
        ["Ponto de Equilíbrio Contábil (R$)", "Custos Fixos / Índice de Margem", "R$ 12.454,37", "Faturamento mensal mínimo necessário para a empresa não ter prejuízo."],
        ["Ponto de Equilíbrio em Unidades", "Ponto de Equilíbrio / Preço Médio (R$ 98,81)", "126 combos", "Vender 32 combos por fim de semana (~16 por dia) para cobrir todos os custos."],
        ["Lucratividade sobre Vendas (Margem)", "Lucro Líquido / Receita Bruta", "11,71%", "Retorno operacional saudável e protegido contra oscilações de custos."],
        ["Prazo de Retorno do Capital (Payback)", "Saldo Acumulado no Fluxo de Caixa Dinâmico", "11 a 12 meses", "Recuperação integral dos R$ 38.000,00 investidos no primeiro ano."]
    ], widths=[2400, 2400, 1700, 2860])

    add_figure_with_caption(doc, CHART_DIR / "breakeven_pt.png", 6, "Gráfico do Ponto de Equilíbrio Operacional", "Elaboração própria a partir do modelo de custos (2026).")

    # 5.8 SUB-CAPÍTULO DE FERIADOS
    doc.add_heading("5.8 Impacto Operacional e Financeiro dos Feriados como Alavancagem Adicional (2026-2028)", level=2)
    add_p(doc, "Como critério de estrita prudência e rigor metodológico, todas as projeções financeiras da DRE base (Tabela 15), do ponto de equilíbrio (Tabela 17) e do fluxo de caixa projetado a 12 meses (Tabela 16) foram calculadas considerando exclusivamente os finais de semana regulares (8 dias de operação por mês, totalizando 160 combos mensais).")
    add_p(doc, "Entretanto, a Casa de Assados Sofia operará em todos os feriados nacionais, estaduais e municipais de Curitiba que ocorrerem de segunda a sexta-feira. Como a totalidade dos custos fixos estruturais da empresa (aluguel do imóvel comercial, honorários contábeis, internet fibra, parcela do microcrédito e infraestrutura de nuvem) já se encontra 100% coberta e amortizada pela operação dos fins de semana normais, as vendas realizadas nos feriados durante a semana funcionam como uma extraordinária alavancagem financeira adicional, transformando uma parcela expressiva de suas receitas diretamente em margem líquida e acelerando a formação de reservas de liquidez.")
    add_p(doc, "Projeta-se para cada feriado em dia útil uma demanda média de 20 a 25 combos (equivalente a uma jornada padrão de sábado), gerando uma receita bruta média de R$ 1.976,13, com custo direto de mercadorias (CMV) de R$ 767,50, impostos e taxas de cartão de R$ 118,57 e remuneração variável de 3 colaboradores diaristas com adicional festivo (R$ 420,00), resultando em um Lucro Líquido Incremental Médio de R$ 625,06 por feriado trabalhado.")

    feriados_rows_pt = [
        ["07/09/2026 (Seg)", "Independência do Brasil (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["08/09/2026 (Ter)", "N. Sra. da Luz dos Pinhais - Padroeira de Curitiba", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["12/10/2026 (Seg)", "Nossa Senhora Aparecida (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["02/11/2026 (Seg)", "Finados (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["20/11/2026 (Sex)", "Dia Nacional de Zumbi e da Consciência Negra", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["25/12/2026 (Sex)", "Natal (Nacional - Ceias Familiares sob Encomenda)", "30 combos", "2.964,20", "1.966,61", "997,59"],
        ["SUBTOTAL 2026 (Ago-Dez: 6 Feriados)", "6 dias extras de operação em dias úteis", "135 combos", "13.338,88", "9.029,73", "4.309,15"],
        ["01/01/2027 (Sex)", "Confraternização Universal (Ano Novo)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["09/02/2027 (Ter)", "Carnaval (Ponto Facultativo / Feriado Comercial)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["26/03/2027 (Sex)", "Sexta-feira Santa / Paixão de Cristo", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["21/04/2027 (Qua)", "Tiradentes (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["27/05/2027 (Qui)", "Corpus Christi (Municipal Curitiba)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["07/09/2027 (Ter)", "Independência do Brasil (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["08/09/2027 (Qua)", "Padroeira de Curitiba (N. Sra. da Luz)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["12/10/2027 (Ter)", "Nossa Senhora Aparecida (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["02/11/2027 (Ter)", "Finados (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["15/11/2027 (Seg)", "Proclamação da República (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["SUBTOTAL 2027 (10 Feriados)", "10 dias extras de operação em dias úteis", "220 combos", "21.737,42", "14.734,81", "7.002,61"],
        ["29/02/2028 (Ter)", "Carnaval (Feriado Comercial)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["14/04/2028 (Sex)", "Sexta-feira Santa (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["21/04/2028 (Sex)", "Tiradentes (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["01/05/2028 (Seg)", "Dia Mundial do Trabalho (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["15/06/2028 (Qui)", "Corpus Christi (Municipal Curitiba)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["07/09/2028 (Qui)", "Independência do Brasil (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["08/09/2028 (Sex)", "Padroeira de Curitiba (N. Sra. da Luz)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["12/10/2028 (Qui)", "Nossa Senhora Aparecida (Nacional)", "25 combos", "2.470,16", "1.658,84", "811,32"],
        ["02/11/2028 (Qui)", "Finados (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["15/11/2028 (Qua)", "Proclamação da República (Nacional)", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["20/11/2028 (Seg)", "Dia Nacional da Consciência Negra", "20 combos", "1.976,13", "1.351,07", "625,06"],
        ["25/12/2028 (Seg)", "Natal (Nacional - Ceias Familiares sob Encomenda)", "30 combos", "2.964,20", "1.966,61", "997,59"],
        ["SUBTOTAL 2028 (12 Feriados)", "12 dias extras de operação em dias úteis", "270 combos", "26.683,68", "18.060,54", "8.623,14"],
        ["TOTAL BIÊNIO (2026-2028)", "28 feriados úteis de alavancagem operacional", "625 combos", "61.759,98", "41.825,08", "19.934,90"]
    ]
    build_styled_table(doc, ["Data e Dia da Semana", "Denominação Oficial do Feriado", "Volume Estimado", "Receita Bruta (R$)", "Custos Var. + Diárias", "Lucro Líquido Extra (R$)"], 
                       feriados_rows_pt, widths=[1800, 2900, 1200, 1300, 1400, 1260], font_size=7.5, align_right_cols=[2,3,4,5])

    add_p(doc, "Como se sintetiza na Tabela 18, a abertura nos feriados representa uma injeção de receita bruta acumulada de R$ 61.759,98 e um Lucro Líquido Extraordinário de R$ 19.934,90 entre agosto de 2026 e dezembro de 2028. Essa margem excedente fortalece a posição de tesouraria, protege o negócio contra flutuações de demanda ou clima desfavorável e viabiliza a quitação antecipada do microcrédito ou a reinversão precoce na ampliação física da unidade.")

    # CAPÍTULO 6
    doc.add_page_break()
    doc.add_heading("6 ANÁLISE DE VIABILIDADE E GESTÃO DE RISCOS", level=1)
    add_p(doc, "A análise de viabilidade avalia a robustez do modelo de negócio frente às incertezas do mercado, integrando a análise SWOT com matrizes de risco e planos de contingência.")

    doc.add_heading("6.1 Matriz SWOT / FOFA Estratégica", level=2)
    build_styled_table(doc, ["Fatores Internos / Externos", "Aspectos Favoráveis (Positivos)", "Aspectos Desfavoráveis (Negativos)"], [
        ["Ambiente Interno (Controle da Empresa)", 
         "FORÇAS (S):\n• Orçamento confortável de R$ 38.000 c/ equipamentos novos e exaustão industrial;\n• Cardápio enxuto com alta padronização e fichas técnicas;\n• Canal próprio direto via WhatsApp com CRM Casa de Assados Sofia;\n• Margem de contribuição saudável (55,16%);\n• Localização com facilidade logística no Umbará;\n• Alavancagem de caixa através da abertura em todos os feriados.",
         "FRAQUEZAS (W):\n• Marca nova sem base prévia de clientes;\n• Capacidade fixa de cocção nas churrasqueiras;\n• Dependência operacional do sócio administrador."],
        ["Ambiente Externo (Mercado e Contexto)", 
         "OPORTUNIDADES (O):\n• Hábito cultural consolidado do churrasco de domingo e feriados;\n• Insatisfação do público com filas em concorrentes tradicionais;\n• Expansão de condomínios residenciais na região sul;\n• Fidelização e recompra previsível via CRM.",
         "AMEAÇAS (T):\n• Alta e volatilidade nos preços de carnes no atacado;\n• Concorrência de rotisserias de grandes redes de supermercados;\n• Risco de instabilidade técnica no canal de WhatsApp;\n• Chuvas intensas e clima frio em fins de semana."]
    ], widths=[2400, 3500, 3460], font_size=8)

    doc.add_heading("6.2 Análise de Sensibilidade em Três Cenários", level=2)
    scenarios_pt = [
        ["Cenário Pessimista (-20% Vendas)", "128 combos", "12.647,20", "6.976,37", "6.870,00", "106,37", "0,84%"],
        ["Cenário Base (Projetado)", "160 combos", "15.809,00", "8.720,46", "6.870,00", "1.850,46", "11,71%"],
        ["Cenário Otimista (+30% Vendas)", "208 combos", "20.551,70", "11.336,60", "6.870,00", "4.466,60", "21,73%"]
    ]
    build_styled_table(doc, ["Cenário Simulado", "Volume/Mês", "Receita Bruta", "Margem Contrib.", "Custos Fixos", "Lucro Líquido", "Margem %"], 
                       scenarios_pt, widths=[2200, 1100, 1400, 1400, 1300, 1300, 1060], font_size=8, align_right_cols=[1,2,3,4,5,6])

    add_figure_with_caption(doc, CHART_DIR / "scenarios_pt.png", 8, "Análise de Sensibilidade e Comparação de Cenários", "Elaboração própria a partir da modelagem de cenários (2026).")

    doc.add_heading("6.3 Avaliação dos Indicadores de Viabilidade", level=2)
    add_p(doc, "A análise de sensibilidade demonstra a resiliência e blindagem do negócio com o orçamento de R$ 38.000,00: mesmo sob uma queda de 20% nas vendas nos fins de semana normais (128 combos), a empresa opera acima do ponto de equilíbrio (126 combos), mantendo resultado líquido positivo (R$ 106,37). No cenário otimista (208 combos), o lucro líquido mensal atinge R$ 4.466,60 com margem líquida de 21,73%, sem computar as receitas extraordinárias dos feriados.")

    doc.add_heading("6.4 Matriz de Riscos e Planos de Contingência", level=2)
    build_styled_table(doc, ["Fator de Risco", "Probabilidade / Impacto", "Medida Preventiva", "Plano de Contingência Imediato"], [
        ["Alta abrupta nos preços de carne", "Alta / Alto", "Homologação de 3 frigoríficos SIF e cotações semanais.", "Substituição temporária por fornecedor B ou ajuste seletivo de margem."],
        ["Demanda abaixo do ponto de equilíbrio", "Média / Alto", "Ações de pré-venda ativa via CRM Casa de Assados Sofia às sextas-feiras e degustação local.", "Acionamento de promoções de incentivo e redução de diárias operacionais."],
        ["Bloqueio de linha no WhatsApp", "Média / Alto", "Uso de API oficial com opt-in explícito (LGPD) e disparos moderados.", "Ativação de linha secundária de contingência e ligação telefônica direta."],
        ["Falta de energia ou avaria de assadora", "Baixa / Crítico", "Equipamentos novos com garantia, manutenção preventiva e churrasqueira híbrida.", "Acionamento de assistência técnica 24h e transferência para o bafo a carvão."],
        ["Vazamento de dados ou violação LGPD", "Baixa / Alto", "Minimização de dados, controle de acesso e banco criptografado.", "Auditoria técnica imediata, revogação do registro e notificação formal."]
    ], widths=[2200, 1600, 2800, 2760], font_size=8)

    # CAPÍTULO 7
    doc.add_page_break()
    doc.add_heading("7 ANEXOS E INSTRUMENTOS DE IMPLANTAÇÃO", level=1)
    add_p(doc, "Este capítulo reúne os instrumentos metodológicos, operacionais, regulatórios, fiscais, trabalhistas e ilustrações conceituais desenvolvidos para nortear a execução da Casa de Assados Sofia.")

    doc.add_heading("7.1 Plano de Ação 5W2H de 30 Dias", level=2)
    actions_pt = [
        ["1. Consulta de Viabilidade", "Verificar zoneamento e adequação do imóvel", "Wilkin Barban", "Dias 1 a 3", "Portal Prefeitura de Curitiba", "Sem custo"],
        ["2. Registro Empresarial", "Obter CNPJ, contrato social e enquadramento ME", "Contador / Wilkin", "Dias 4 a 10", "Junta Comercial do Paraná / RFB", "R$ 800,00"],
        ["3. Compra de Equipamentos e Coifa", "Adquirir assadoras, coifa industrial, freezer e bancadas", "Wilkin Barban", "Dias 8 a 15", "Fornecedores comerciais de Curitiba", "R$ 24.500,00"],
        ["4. Instalação e Adequações", "Pintura, pontos hidráulicos e montagem da exaustão", "Técnicos especializados", "Dias 12 a 20", "Imóvel na Rua Dep. Pinheiro Júnior", "R$ 2.500,00"],
        ["5. Parametrização do CRM", "Configurar VPS, banco de dados e bot de pedidos", "Wilkin Barban", "Dias 15 a 22", "Servidor Cloud e API Mensageria", "R$ 50,00"],
        ["6. Treinamento da Equipe", "Simular ciclos de cocção, embalagem e higiene", "Equipe completa", "Dias 22 a 25", "Sede operacional da unidade", "R$ 300,00"],
        ["7. Piloto Comercial (Semana 1)", "Produzir e expedir primeiro lote de 25 combos", "Equipe completa", "Dias 26 a 30", "Clientes cadastrados da microrregião", "R$ 500,00"]
    ]
    build_styled_table(doc, ["O Que (What)", "Por Que (Why)", "Quem (Who)", "Quando (When)", "Onde (Where)", "Quanto (How Much)"], 
                       actions_pt, widths=[1800, 2400, 1300, 1100, 1600, 1160], font_size=8)

    doc.add_heading("7.2 Catálogo Fotográfico de Maquinários e Equipamentos Adquiridos", level=2)
    add_p(doc, "Nota de Esclarecimento Metodológico: Todas as fotografias de maquinários e equipamentos operacionais apresentadas neste catálogo e ao longo deste trabalho foram geradas por meio de Inteligência Artificial (IA) generativa para fins estritamente ilustrativos e acadêmicos, representando visualmente e com máxima fidelidade técnica os modelos, especificações, capacidades e padrões sanitários dos equipamentos reais a serem adquiridos no mercado comercial de Curitiba para a implantação da unidade.", boldlead="Nota de Esclarecimento Metodológico:")
    add_p(doc, "Apresenta-se o registro fotográfico individual e catálogo técnico de cada ativo fixo e sistema operacional adquirido para a unidade da Casa de Assados Sofia, totalizando R$ 24.500,00 de investimento em parque de máquinas homologado:")

    add_p(doc, "1. Máquinas Giratórias de Frango a Gás GLP (Figura 9): 2 unidades novas equipadas com queimadores infravermelhos a gás GLP no painel traseiro, espetos rotativos em aço inoxidável, portas de vidro temperado e iluminação interna. Capacidade combinada para 40 frangos por dia em 2 ciclos.")
    add_figure_with_caption(doc, IMG_FIG09, 9, "Máquinas Giratórias de Frango a Gás GLP com Queimadores Infravermelhos", "Fotografia de referência das assadoras giratórias a gás gerada por IA (2026).", width_cm=14.0)

    add_p(doc, "2. Churrasqueira Tradicional a Carvão para Bafo (Figura 10): 1 unidade reforçada em chapa de aço com revestimento interno de tijolos refratários, leito de brasas incandescentes, grelha elevatória em V com manivela lateral e tampa pesada basculante para cocção lenta ao vapor (bafo) por 6 horas.")
    add_figure_with_caption(doc, IMG_FIG10, 10, "Churrasqueira Tradicional a Carvão para Bafo com Grelha Elevatória", "Fotografia de referência da churrasqueira a carvão para bafo gerada por IA (2026).", width_cm=14.0)

    add_p(doc, "3. Sistema de Coifa Industrial em Aço Inox (Figura 11): Coifa em aço inoxidável escovado AISI 304 com filtros inerciais laváveis tipo labirinto de alta retenção de gordura, luminárias blindadas e duto circular galvanizado, atendendo às exigências da Vigilância Sanitária de Curitiba.")
    add_figure_with_caption(doc, IMG_FIG11, 11, "Sistema de Coifa Industrial em Aço Inox com Exaustão Mecânica", "Fotografia de referência do sistema de exaustão e coifa gerada por IA (2026).", width_cm=14.0)

    add_p(doc, "4. Freezer Horizontal Comercial Dupla Ação de 510 Litros (Figura 12): Equipamento com duas tampas cegas basculantes com chave, termostato digital externo programado para -18°C e rodízios reforçados para armazenamento seguro de carnes.")
    add_figure_with_caption(doc, IMG_FIG12, 12, "Freezer Horizontal Comercial Dupla Ação de 510 Litros", "Fotografia de referência do freezer horizontal comercial gerada por IA (2026).", width_cm=14.0)

    add_p(doc, "5. Refrigerador Comercial Vertical de Inox de 4 Portas (Figura 13): Gabinete monobloco em aço inox AISI 304 com 4 portas independentes, controlador digital de temperatura (+2°C a +4°C) para conservação asséptica de marinadas e guarnições preparadas.")
    add_figure_with_caption(doc, IMG_FIG13, 13, "Refrigerador Comercial Vertical de Inox de 4 Portas", "Fotografia de referência do refrigerador vertical comercial inox gerada por IA (2026).", width_cm=14.0)

    add_p(doc, "6. Mesa Central de Manipulação Inox AISI 304 com Balança Digital (Figura 14): Bancadas de trabalho (2,0m x 0,9m) em aço inox com prateleira inferior, balança digital computadora homologada pelo Inmetro, tábua de corte sanitária e cubas gastronômicas Gastronorm.")
    add_figure_with_caption(doc, IMG_FIG14, 14, "Mesa Central de Manipulação Inox AISI 304 com Balança Digital", "Fotografia de referência da mesa de manipulação e balança gerada por IA (2026).", width_cm=14.0)

    doc.add_heading("7.3 Planta Baixa Arquitetônica e Layout Funcional", level=2)
    add_p(doc, "Apresenta-se a planta técnica e distribuição operacional do imóvel de 60,0 m² (10,0m x 6,0m), detalhando os 7 setores funcionais, coifa de exaustão e o fluxo sanitário unidirecional em conformidade com a Anvisa RDC 216/2004.")
    add_figure_with_caption(doc, IMG_FIG15, 15, "Planta Baixa Técnica e Fluxo Sanitário Unidirecional (60,0 m²)", "Desenho arquitetônico conceitual e layout funcional desenvolvido para o plano de negócio (2026).", width_cm=15.0)

    doc.add_heading("7.4 Simulação de Notas Fiscais Eletrônicas e Documentos Fiscais (NF-e)", level=2)
    add_p(doc, "Apresenta-se a simulação estruturada dos documentos fiscais eletrônicos que comprovam a aquisição de insumos cárnicos certificados e ativos operacionais da unidade:")

    build_styled_table(doc, ["Campo da Nota Fiscal NF-e nº 000.142.857 (Série 1)", "Dados do Fornecedor / Insumos Cárnicos Homologados"], [
        ["Emitente / Razão Social", "Frigorífico Avícola & Bovino Sul do Paraná Ltda. (CNPJ: 76.842.119/0001-45 | IE: 90.142.883-10)"],
        ["Destinatário / Comprador", "Casa de Assados Sofia Ltda. (CNPJ: 54.891.204/0001-88 | Endereço: Rua Dep. Pinheiro Júnior, 1380)"],
        ["Natureza da Operação / CFOP", "Venda de mercadorias adquiridas de terceiros para industrialização / CFOP: 5.102"],
        ["Chave de Acesso da NF-e (44 dígitos)", "4126 0876 8421 1900 0145 5500 1000 1428 5710 9842 1194"],
        ["Itens Detalhados", "Item 1: Frango Resfriado Inteiro c/ SIF (80 un / 160 kg) - R$ 1.280,00\nItem 2: Costela Bovina c/ SIF (40 kg) - R$ 1.120,00\nItem 3: Costelinha Suína Especial (20 kg) - R$ 440,00\nItem 4: Linguiça Toscana Artesanal (15 kg) - R$ 285,00"],
        ["Valor Total da Nota Fiscal", "R$ 3.125,00 (ICMS retido por Substituição Tributária - ST)"],
        ["Condição de Pagamento / Vencimento", "Faturado a 14 dias via Boleto Bancário homologado"]
    ], widths=[3400, 6000], font_size=8)

    build_styled_table(doc, ["Campo da Nota Fiscal NF-e nº 000.089.412 (Série 1)", "Dados do Fornecedor de Máquinas e Equipamentos"], [
        ["Emitente / Razão Social", "Máquinas & Equipamentos Gastronômicos Curitiba Ltda. (CNPJ: 81.332.904/0001-12 | IE: 90.284.112-90)"],
        ["Destinatário / Comprador", "Casa de Assados Sofia Ltda. (CNPJ: 54.891.204/0001-88)"],
        ["Chave de Acesso da NF-e", "4126 0881 3329 0400 0112 5500 1000 0894 1210 3341 8902"],
        ["Itens Faturados / Ativo Fixo", "2x Assadoras Giratórias a Gás 10 Frangos Novas (R$ 4.800,00)\n1x Churrasqueira Profissional c/ Elevador (R$ 2.200,00)\n1x Sistema de Coifa e Exaustão Industrial (R$ 4.200,00)\n1x Freezer Horizontal 510L Dupla Ação (R$ 3.100,00)\n1x Refrigerador Comercial Inox 4 Portas (R$ 3.400,00)\n2x Bancadas Centrais Inox 304 2,0x0,9m + Balança (R$ 2.700,00)"],
        ["Valor Total Faturado", "R$ 20.400,00 (Garantia de fábrica de 12 meses com certificado)"]
    ], widths=[3400, 6000], font_size=8)

    doc.add_heading("7.5 Quadro de Licenciamento, Alvarás e Regularização Sanitária", level=2)
    build_styled_table(doc, ["Órgão Emissor / Secretaria", "Documento / Alvará", "Número de Protocolo / Registro", "Status / Validade"], [
        ["Secretaria Municipal de Urbanismo (SMU Curitiba)", "Consulta Prévia de Viabilidade Técnica e Legal", "Proc. 2026/048192-PMC", "Aprovada e Deferida"],
        ["Prefeitura Municipal de Curitiba (PMC)", "Alvará de Localização e Funcionamento", "Alvará nº 09.842.115/0001", "Vigente / Regular"],
        ["Vigilância Sanitária Municipal (VISA Curitiba)", "Licença Sanitária de Estabelecimento de Alimentos", "Protocolo VISA nº 88412-26", "Válida (RDC 216/04)"],
        ["Corpo de Bombeiros Militar do Paraná (CBMPR)", "Certificado de Vistoria de Bombeiros (CLCB)", "CLCB nº 2026-PR-004182", "Aprovado p/ 12 meses"],
        ["Responsabilidade Técnica Gastronômica", "Manual de Boas Práticas e POPs Sanitários", "Registro RT nº 2026-MBP", "Implantado na Unidade"]
    ], widths=[2800, 3200, 2200, 1200], font_size=8)

    doc.add_heading("7.6 Instrumentos de Contratação e Recibos de Diaristas", level=2)
    add_p(doc, "Para garantir total transparência e segurança jurídica frente a contingências trabalhistas na Justiça do Trabalho (TRT 9ª Região), detalham-se a seguir a minuta padronizada do Contrato de Trabalho Intermitente (Tabela 26) e o modelo oficial de recibo/holerite de jornada de fim de semana (Tabela 27):")

    contract_rows_pt = [
        ["Cláusula 1ª - Identificação das Partes", "EMPREGADORA: Casa de Assados Sofia Ltda., CNPJ 54.891.204/0001-88.\nEMPREGADO: Colaborador Operacional de Fim de Semana, com CTPS Digital e CPF regular."],
        ["Cláusula 2ª - Objeto e Regime de Trabalho", "Contratação sob regime de Trabalho Intermitente (Art. 452-A da CLT - Lei nº 13.467/2017), para prestação de serviços não contínuos com alternância de períodos de prestação de serviços e de inatividade."],
        ["Cláusula 3ª - Convocação Prévia e Aceite", "A Empregadora convocará o Empregado por meio eletrônico rastreável (CRM/WhatsApp) com antecedência mínima de 72 horas. O Empregado terá 24 horas para responder ao chamado."],
        ["Cláusula 4ª - Remuneração e Verbas Proporcionais", "Remuneração horária fixada em R$ 15,00 (R$ 120,00 por diária de 8h). Ao final de cada período trabalhado, serão pagos discriminadamente: remuneração base, DSR proporcional, 13º salário proporcional e férias proporcionais com 1/3."],
        ["Cláusula 5ª - Normas Sanitárias e Segurança", "O Empregado obriga-se a utilizar os EPIs fornecidos, observar as Boas Práticas da RDC 216 da Anvisa e manter atualizado o Atestado de Saúde Ocupacional (ASO)."],
        ["Cláusula 6ª - Foro e Legislação Aplicável", "Eleito o Foro da Comarca de Curitiba - PR, regendo-se o presente instrumento pela Consolidação das Leis do Trabalho (CLT)."]
    ]
    build_styled_table(doc, ["Cláusula do Contrato Intermitente", "Termos e Condições Jurídicas Estruturadas"], contract_rows_pt, widths=[2800, 6600], font_size=8)

    rpa_rows_pt = [
        ["Identificação da Empresa Empregadora", "Casa de Assados Sofia Ltda. - CNPJ: 54.891.204/0001-88 - Curitiba/PR"],
        ["Identificação do Colaborador / Função", "Nome: [Colaborador Operacional] | CPF: XXX.XXX.XXX-XX | Cargo: Churrasqueiro / Auxiliar"],
        ["Período Operacional / Diárias Cumpridas", "Período: Final de Semana (Sábado e Domingo) - Total: 2 diárias de 8h (16h trabalhadas)"],
        ["(+) Remuneração Base das Diárias (16 horas x R$ 12,00)", "R$ 192,00 (Remuneração direta pelas horas trabalhadas)"],
        ["(+) Descanso Semanal Remunerado Proporcional (DSR)", "R$ 16,00 (Adicional legal obrigatório)"],
        ["(+) 13º Salário Proporcional", "R$ 16,00 (Proporcional legal conforme Art. 452-A §6º da CLT)"],
        ["(+) Férias Proporcionais + 1/3 Constitucional", "R$ 16,00 (Proporcional legal com terço constitucional)"],
        ["(=) TOTAL BRUTO DA REMUNERAÇÃO", "R$ 240,00 (Total das 2 diárias integrais de R$ 120,00)"],
        ["(-) Contribuição Previdenciária Oficial (INSS 7,5%)", "R$ 18,00 (Retenção legal conforme tabela vigente)"],
        ["(=) VALOR LÍQUIDO PAGO AO TRABALHADOR", "R$ 222,00 (Quitado via transferência PIX com recibo assinado)"]
    ]
    build_styled_table(doc, ["Campo do Recibo / Holerite Operacional", "Discriminação Financeira e Bases de Cálculo"], rpa_rows_pt, widths=[3500, 5900], font_size=8)

    doc.add_heading("7.7 Questionário Estruturado de Pesquisa de Mercado", level=2)
    add_p(doc, "Instrumento estruturado de pesquisa de mercado para validação contínua da demanda no bairro Umbará:")
    questions_pt = [
        "1. Quantas pessoas compõem o seu núcleo familiar que costuma almoçar reunido aos finais de semana e feriados?",
        "2. Com que frequência a sua família costuma comprar comida pronta (frango assado, costela ou churrasco) aos sábados, domingos ou feriados?",
        "3. Em sua opinião, qual é o principal incômodo nos assadores tradicionais da região? ( ) Filas e demora ( ) Sabor inconstante ( ) Preço elevado ( ) Falta de entrega pontual ( ) Pouca variedade de combos familiares.",
        "4. Qual modalidade de compra melhor atende à sua rotina aos domingos e feriados? ( ) Retirada rápida com horário agendado sem fila ( ) Entrega em domicílio com horário marcado.",
        "5. Qual dos seguintes combos familiares melhor atende ao perfil do seu almoço? ( ) Clássico de Frango Assado ( ) Costela Bovina no Bafo ( ) Dueto Frango e Costelinha Suína ( ) Kit Churrasco Família Completo.",
        "6. Você aceitaria receber o cardápio semanal e reservar o seu almoço com antecedência às sextas-feiras ou vésperas de feriados via WhatsApp, garantindo seu pedido sem risco de esgotamento? ( ) Sim, com certeza ( ) Talvez ( ) Não."
    ]
    for q in questions_pt:
        add_p(doc, q)

    doc.add_heading("7.8 Dicionário de Dados e Telas de Produção do Sistema CRM Casa de Assados Sofia", level=2)
    build_styled_table(doc, ["Tabela / Entidade", "Campo / Atributo", "Tipo de Dado", "Descrição e Regra de Negócio (LGPD)"], [
        ["Clientes (tb_clientes)", "id_cliente / nome / whatsapp / endereco", "INT / VARCHAR", "Identificação unívoca, telefone e endereço com chave criptografada."],
        ["Consentimento (tb_consent)", "status_optin / data_registro / canal", "BOOLEAN / DATETIME", "Registro formal de autorização de mensagens com opção de cancelamento (LGPD)."],
        ["Pedidos (tb_pedidos)", "id_pedido / data_hora / status_pedido / canal", "INT / DATETIME / ENUM", "Ciclo de vida: Reservado -> Em Cocção -> Embalado -> Entregue."],
        ["Itens (tb_itens_pedido)", "id_item / id_combo / quantidade / valor_unit", "INT / INT / INT / DECIMAL", "Detalhamento de combos e adicionais para a cozinha e cálculo exato de CMV."],
        ["Avaliações (tb_nps)", "nota_nps / comentario / tempo_resposta", "INT (1 a 5) / TEXT / INT", "Avaliação pós-venda para cálculo de NPS e alertas de fidelização."]
    ], widths=[2000, 2400, 1600, 3360], font_size=8)

    add_p(doc, "Documentação das Interfaces Operacionais em Produção (DuckDNS / VPS):", boldlead="Documentação das Interfaces Operacionais em Produção (DuckDNS / VPS):")
    add_p(doc, "Apresentam-se a seguir as capturas de tela reais do Sistema CRM Casa de Assados Sofia em ambiente de produção (https://casadeasados.duckdns.org/), comprovando a execução prática do desenvolvimento próprio em software livre e o funcionamento integrado da inteligência artificial:")

    add_p(doc, "1. Portal de Autenticação e Acesso Web (Figura 21): Interface de login com alternância entre perfis de acesso ('Sou Cliente' e 'Equipe / Operador'), com formulário validado para celulares de Curitiba (DDD 41), autenticação segura de senha e link direto para cadastro de novos clientes.")
    add_figure_with_caption(doc, IMG_FIG21, 21, "Portal Web de Acesso e Autenticação do Cliente e Operador", "Captura de tela real do portal de login do CRM Casa de Assados Sofia (https://casadeasados.duckdns.org/login).", width_cm=14.5)

    add_p(doc, "2. Console de Atendimento Omnichannel e IA 'Sofia' com DeepSeek LLM (Figura 22): Central operacional de atendimento que integra os canais de WhatsApp e Telegram, exibindo em tempo real o saldo de créditos da API DeepSeek LLM ($1.19 disponíveis), controle de filas de atendimento (Fila IA, Fila Humana e Fechadas), interação em linguagem natural com a assistente virtual 'Sofia' (personalizada com vocabulário acolhedor de Curitiba, como 'piá', e recomendações estratégicas de combos) e carrinho lateral de pedidos em tempo real.")
    add_figure_with_caption(doc, IMG_FIG22, 22, "Console de Atendimento Omnichannel e IA Virtual 'Sofia' com DeepSeek LLM", "Captura de tela real da console de atendimento conversacional com IA generativa (https://casadeasados.duckdns.org/atendimento).", width_cm=14.5)

    add_p(doc, "3. Painel KDS de Gestão de Pedidos em Tempo Real e Faturamento Operacional (Figura 23): Painel de gestão visual da produção gastronômica (Kitchen Display System), consolidando os indicadores-chave do turno em tempo real (Total de 18 pedidos, 16 em preparo/novos, 2 entregues/concluídos e Faturamento Real de R$ 1.542,90), filtros por status e cartões individuais de comandas com identificador hash, dados do cliente, combo contratado, status de pagamento PIX e horário agendado de retirada.")
    add_figure_with_caption(doc, IMG_FIG23, 23, "Painel KDS de Gestão de Pedidos em Tempo Real e Faturamento Operacional", "Captura de tela real do painel KDS de gestão de pedidos e faturamento (https://casadeasados.duckdns.org/atendimento/pedidos).", width_cm=14.5)

    doc.add_heading("7.9 Renders e Documentação Fotográfica dos Combos Familiares", level=2)
    add_p(doc, "Nota de Esclarecimento Metodológico: As fotografias dos combos gastronômicos apresentadas a seguir foram geradas por meio de Inteligência Artificial (IA) generativa como representações visuais hiper-realistas de referência comercial das fichas técnicas e porções exatas descritas no cardápio oficial.", boldlead="Nota de Esclarecimento Metodológico:")
    add_p(doc, "Incorpora-se a documentação fotográfica individual e detalhada dos 4 combos gastronômicos oferecidos pela Casa de Assados Sofia, elaborados com estrito rigor às fichas técnicas operacionais e porções exatas:")

    add_p(doc, "Combo 1 – O Clássico da Sofia (Preço: R$ 69,90 | Rendimento: 3 a 4 pessoas): Composto por 1 frango recheado inteiro assado dourado com pele crocante (~1,4kg assado), recheio de farofa temperada visível na cavidade, acompanhado de pote de maionese caseira tradicional de batata e cenoura (300g) e pote de farofa artesanal crocante com bacon (250g).")
    add_figure_with_caption(doc, IMG_FIG16, 16, "Documentação Fotográfica do Combo 1: O Clássico da Sofia", "Fotografia comercial de referência do Combo 1 gerada por IA (2026).", width_cm=14.5)

    add_p(doc, "Combo 2 – Costela Suprema no Bafo (Preço: R$ 119,90 | Rendimento: 4 pessoas): Composto por generoso corte de 1,0kg de costela bovina premium com osso, assada lentamente no bafo por 6 horas em calor indireto, exibindo crosta caramelizada e interior suculento, acompanhada de mandioca amarela na manteiga de garrafa (300g), vinagrete fresco de tomate e cebola e farofa da casa (250g).")
    add_figure_with_caption(doc, IMG_FIG17, 17, "Documentação Fotográfica do Combo 2: Costela Suprema no Bafo", "Fotografia comercial de referência do Combo 2 gerada por IA (2026).", width_cm=14.5)

    add_p(doc, "Combo 3 – Dueto Sofia (Preço: R$ 94,90 | Rendimento: 3 a 4 pessoas): Composto por exatamente meio frango assado dourado crocante com ervas + 500g de costelinha de porco macia marinada em finas ervas e glaceada, servidos em travessa com batatas rústicas douradas ao alecrim (300g) e farofa artesanal da casa (200g).")
    add_figure_with_caption(doc, IMG_FIG18, 18, "Documentação Fotográfica do Combo 3: Dueto Sofia (Frango & Costelinha)", "Fotografia comercial de referência do Combo 3 gerada por IA (2026).", width_cm=14.5)

    add_p(doc, "Combo 4 – Kit Churrasco Família (Preço: R$ 169,90 | Rendimento: 5 a 6 pessoas): Grande banquete de churrasco composto por 1 frango recheado inteiro dourado + 700g de costela bovina no bafo + 4 linguiças toscanas artesanais grelhadas na brasa + 4 pães de alho tostados na grelha, acompanhados de pote grande de maionese caseira (500g) e farofa grande artesanal (400g).")
    add_figure_with_caption(doc, IMG_FIG19, 19, "Documentação Fotográfica do Combo 4: Kit Churrasco Família", "Fotografia comercial de referência do Combo 4 gerada por IA (2026).", width_cm=14.5)

    add_p(doc, "Apresenta-se adicionalmente o conceito ilustrativo tridimensional da fachada da unidade, embalagens ecológicas seladas e estação de controle digital do CRM Casa de Assados Sofia:")
    add_figure_with_caption(doc, IMG_FIG20, 20, "Conceito Ilustrativo: Fachada, Embalagens, Produtos e Estação CRM Casa de Assados Sofia", "Desenho conceitual gerado em alta definição por IA para o plano de negócio (2026).", width_cm=15.0)

    # CONCLUSÃO
    doc.add_page_break()
    doc.add_heading("CONCLUSÃO", level=1)
    add_p(doc, "O presente Trabalho de Conclusão de Curso comprovou a plena viabilidade mercadológica, operacional, econômico-financeira e tecnológica para a implantação da Casa de Assados Sofia no bairro Umbará, em Curitiba - PR. A investigação demonstrou que a aplicação disciplinada de conceitos de administração e informática aprendidos no Colégio Excelência permite estruturar um empreendimento gastronômico de conveniência em um modelo de negócio altamente rentável, previsível e escalável.")
    add_p(doc, "A inovação do projeto decorre da união sinérgica entre a excelência artesanal na preparação de carnes tradicionais e o uso estratégico da tecnologia de gestão (CRM Casa de Assados Sofia). Essa combinação soluciona os três maiores entraves do setor: elimina o desperdício de matérias-primas por meio de pré-vendas programadas, erradica as filas de espera mediante agendamento em janelas de 15 minutos e constrói relacionamentos duradouros com os clientes da microrregião do Bairro Novo.")
    add_p(doc, "Do ponto de vista financeiro e de segurança jurídica sob o orçamento confortável de R$ 38.000,00, os números e protocolos confirmam a solidez da proposta: o ponto de equilíbrio de R$ 12.454,37 (~126 combos) situa-se confortavelmente abaixo da demanda do cenário base dos fins de semana normais (160 combos), proporcionando lucratividade líquida de 11,71% e retorno total do capital investido na curva de maturação em 11 a 12 meses. A abertura nos 28 feriados úteis projetados para o biênio 2026-2028 conforma uma alavancagem extraordinária de mais de R$ 61 mil em receita e quase R$ 20 mil em lucro líquido complementar. Ademais, a infraestrutura com exaustão industrial, a contratação sob regime intermitente (CLT 452-A), a identidade corporativa estruturada e os cardápios especializados blindam a empresa contra contingências regulatórias e trabalhistas, assegurando seu sucesso sustentável no mercado curitibano.")

    # REFERÊNCIAS
    doc.add_page_break()
    doc.add_heading("REFERÊNCIAS", level=1)
    refs = [
        "ASSOCIAÇÃO BRASILEIRA DE BARES E RESTAURANTES (ABRASEL). Panorama do Setor de Alimentação Fora do Lar no Brasil. Brasília: Abrasel, 2024.",
        "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS (ABNT). NBR 6023: Informação e documentação – Referências – Elaboração. Rio de Janeiro: ABNT, 2018.",
        "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS (ABNT). NBR 14724: Informação e documentação – Trabalhos acadêmicos – Apresentação. Rio de Janeiro: ABNT, 2011.",
        "BRASIL. Agência Nacional de Vigilância Sanitária (ANVISA). Resolução RDC nº 216, de 15 de setembro de 2004. Regulamento Técnico de Boas Práticas para Serviços de Alimentação. Brasília: Diário Oficial da União, 2004.",
        "BRASIL. Consolidação das Leis do Trabalho (CLT). Decreto-Lei nº 5.452, de 1º de maio de 1943, alterado pela Lei nº 13.467, de 13 de julho de 2017 (Reforma Trabalhista - Contrato de Trabalho Intermitente, Art. 452-A). Brasília: Presidência da República, 2017.",
        "BRASIL. Lei nº 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais (LGPD). Brasília: Presidência da República, 2018.",
        "BRASIL. Ministério do Empreendedorismo, da Microempresa e da Empresa de Pequeno Porte. Portal do Empreendedor. Disponível em: <https://www.gov.br/empresas-e-negocios/pt-br/empreendedor>. Acesso em: 15 ago. 2026.",
        "CENTRAIS DE ABASTECIMENTO DO PARANÁ S/A (CEASA/PR). Cotações de Preços e Abastecimento Regional de Hortigranjeiros. Curitiba: CEASA, 2026. Disponível em: <https://www.ceasa.pr.gov.br/>. Acesso em: 15 ago. 2026.",
        "COLÉGIO EXCELÊNCIA. Manual de Normas Técnicas para Elaboração de Trabalhos de Conclusão de Curso (TCC). Curitiba: Colégio Excelência, 2024. Disponível em: <https://colegioexcelencia.com>. Acesso em: 15 ago. 2026.",
        "CURITIBA. Prefeitura Municipal de Curitiba. Secretaria Municipal de Urbanismo. Guia de Licenciamento Comercial e Zoneamento Urbano. Curitiba: PMC, 2026. Disponível em: <https://www.curitiba.pr.gov.br/servicos/>. Acesso em: 15 ago. 2026.",
        "DOLABELA, Fernando. O segredo de Luísa: uma ideia, uma paixão e um plano de negócios: como nasce o empreendedor e se cria uma empresa. São Paulo: Editora Cultura, 2008.",
        "DORNELAS, José Carlos Assis. Empreendedorismo: transformando ideias em negócios. 8. ed. São Paulo: Empreende, 2021.",
        "FOMENTO PARANÁ. Linhas de Microcrédito Orientado para Microempresas e Empreendedores do Paraná. Curitiba: Governo do Estado do Paraná, 2026. Disponível em: <https://www.fomento.pr.gov.br/>. Acesso em: 15 ago. 2026.",
        "INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA (IBGE). Censo Demográfico 2022 e Estimativas de População de Curitiba. Rio de Janeiro: IBGE, 2024. Disponível em: <https://www.ibge.gov.br/cidades-e-estados/pr/curitiba.html>. Acesso em: 15 ago. 2026.",
        "INSTITUTO PARANAENSE DE DESENVOLVIMENTO ECONÔMICO E SOCIAL (IPARDES). Caderno Estatístico do Município de Curitiba. Curitiba: IPARDES, 2024.",
        "KOTLER, Philip; KELLER, Kevin Lane. Administração de Marketing. 15. ed. São Paulo: Pearson Education do Brasil, 2018.",
        "PEPPERS, Don; ROGERS, Martha. Managing Customer Relationships: A Strategic Framework. Hoboken: John Wiley & Sons, 2004.",
        "SCHUMPETER, Joseph Alois. Teoria do desenvolvimento econômico: uma investigação sobre lucros, capital, crédito, juro e o ciclo econômico. São Paulo: Nova Cultural, 1997.",
        "SERVIÇO BRASILEIRO DE APOIO ÀS MICRO E PEQUENAS EMPRESAS (SEBRAE). Como Elaborar um Plano de Negócio. Brasília: SEBRAE Nacional, 2013.",
        "SERVIÇO BRASILEIRO DE APOIO ÀS MICRO E PEQUENAS EMPRESAS (SEBRAE). Como Implantar Delivery na Era Digital. Curitiba: SEBRAE/PR, 2023.",
        "SWIFT, Ronald S. CRM - Customer Relationship Management: O gerenciamento do relacionamento com o cliente na era do e-business. Rio de Janeiro: Campus, 2001."
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        p.add_run(r)

    # Propriedades
    doc.core_properties.title = 'Casa de Assados Sofia - Plano de Negócio'
    doc.core_properties.author = 'Wilkin Barban Rosabal'
    doc.core_properties.subject = 'Trabalho de Conclusão de Curso (TCC) - Administração e Informática'
    doc.save(output_path)
    print(f"Portuguese thesis generated successfully: {output_path}")

if __name__ == "__main__":
    generate_portuguese_thesis(ROOT / "Borrador_Casa_de_Assados_Sofia_Portugues.docx")
